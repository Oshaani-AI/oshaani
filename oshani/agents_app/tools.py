"""Tool system for agents with API capabilities."""
import json
import logging
import warnings
import requests
import subprocess
import tempfile
import os
import re
import shutil
from typing import Dict, Any, Optional, List
from django.conf import settings
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile

logger = logging.getLogger(__name__)

# Suppress duckduckgo_search "renamed to ddgs" RuntimeWarning when using web_search
warnings.filterwarnings("ignore", message=".*renamed to ddgs.*", category=RuntimeWarning)


class Tool:
    """Base class for all tools."""
    
    def __init__(self, name: str, description: str, instructions: str = ""):
        self.name = name
        self.description = description
        self.instructions = instructions
    
    def execute(self, parameters: Dict[str, Any]) -> Dict[str, Any]:
        """Execute the tool with given parameters."""
        raise NotImplementedError("Subclasses must implement execute method")
    
    def get_schema(self) -> Dict[str, Any]:
        """Get tool schema for LLM."""
        return {
            "name": self.name,
            "description": self.description,
            "instructions": self.instructions,
            "parameters": self.get_parameters_schema()
        }
    
    def get_parameters_schema(self) -> List[Dict[str, Any]]:
        """Get parameters schema for this tool."""
        return []


class WebSearchTool(Tool):
    """Tool for searching the web using Google Custom Search with DuckDuckGo fallback."""
    
    def __init__(self):
        super().__init__(
            name="web_search",
            description="Performs a search on the web using Google Custom Search (with DuckDuckGo fallback) to retrieve structured data or summaries based on specific queries.",
            instructions="Break down complex search queries into smaller parts to find relevant information efficiently. Start by generating a plan or conducting an exploratory search to better understand the user's request."
        )
        self.api_key = getattr(settings, 'GOOGLE_SEARCH_API_KEY', None)
        self.search_engine_id = getattr(settings, 'GOOGLE_SEARCH_ENGINE_ID', None)
    
    def get_parameters_schema(self) -> List[Dict[str, Any]]:
        return [
            {
                "name": "query",
                "description": "The query for which to search the web.",
                "type": "string",
                "required": True
            }
        ]
    
    def _search_with_duckduckgo(self, query: str) -> Dict[str, Any]:
        """Search using DuckDuckGo as fallback. Supports both duckduckgo_search and ddgs packages."""
        try:
            import warnings
            DDGS = None
            with warnings.catch_warnings():
                warnings.simplefilter("ignore", RuntimeWarning)  # suppress "renamed to ddgs" warning
                try:
                    from duckduckgo_search import DDGS as _DDGS
                    DDGS = _DDGS
                except ImportError:
                    try:
                        from ddgs import DDGS as _DDGS
                        DDGS = _DDGS
                    except ImportError:
                        pass
            if DDGS is None:
                raise ImportError("Neither duckduckgo_search nor ddgs installed")
            logger.info(f"Using DuckDuckGo search for query: {query}")
            with warnings.catch_warnings():
                warnings.simplefilter("ignore", RuntimeWarning)
                ddgs = DDGS()
            results = []
            try:
                search_results = ddgs.text(query, max_results=5)
                if search_results is None:
                    pass
                else:
                    # DDGS.text() may return iterator or list; support both body/href and snippet/link keys
                    for item in list(search_results) if hasattr(search_results, '__iter__') and not isinstance(search_results, (str, bytes)) else []:
                        if not isinstance(item, dict):
                            continue
                        snippet = item.get('body') or item.get('snippet') or ''
                        link = item.get('href') or item.get('link') or ''
                        title = item.get('title', '')
                        results.append({
                            "title": title,
                            "snippet": snippet,
                            "link": link
                        })
            except Exception as search_error:
                logger.warning(f"DuckDuckGo search returned no results or error: {str(search_error)}")
            if not results:
                logger.warning(f"No results from DuckDuckGo for query: {query}")
                return {
                    "results": [],
                    "total_results": 0,
                    "source": "duckduckgo",
                    "warning": "No results found. Try a different query or check network."
                }
            return {
                "results": results,
                "total_results": len(results),
                "source": "duckduckgo"
            }
        except ImportError:
            logger.warning("Web search fallback not available: pip install duckduckgo-search (or ddgs)")
            return {"error": "Web search library not available. Install duckduckgo-search or ddgs."}
        except Exception as e:
            logger.error(f"Error in DuckDuckGo search: {str(e)}", exc_info=True)
            return {"error": f"Web search failed: {str(e)}"}
    
    def _search_with_google(self, query: str) -> Dict[str, Any]:
        """Search using Google Custom Search API."""
        try:
            url = "https://www.googleapis.com/customsearch/v1"
            params = {
                'key': self.api_key,
                'cx': self.search_engine_id,
                'q': query
            }
            response = requests.get(url, params=params, timeout=10)
            response.raise_for_status()
            data = response.json()
            
            results = []
            for item in data.get('items', [])[:5]:  # Limit to top 5 results
                results.append({
                    "title": item.get('title', ''),
                    "snippet": item.get('snippet', ''),
                    "link": item.get('link', '')
                })
            
            return {
                "results": results,
                "total_results": data.get('searchInformation', {}).get('totalResults', 0),
                "source": "google"
            }
        except Exception as e:
            logger.error(f"Error in Google search: {str(e)}", exc_info=True)
            raise  # Re-raise to trigger fallback
    
    def execute(self, parameters: Dict[str, Any]) -> Dict[str, Any]:
        query = parameters.get('query', '')
        if not query:
            return {"error": "Query parameter is required"}
        
        # Try Google Search first if API keys are configured
        if self.api_key and self.search_engine_id:
            try:
                result = self._search_with_google(query)
                logger.info(f"Google search successful for query: {query}")
                return result
            except Exception as e:
                logger.warning(f"Google search failed, falling back to DuckDuckGo: {str(e)}")
                # Fall through to DuckDuckGo
        
        # Fallback to DuckDuckGo
        logger.info("Using DuckDuckGo search (Google API keys not configured or Google search failed)")
        result = self._search_with_duckduckgo(query)
        
        # If DuckDuckGo also fails, return error
        if "error" in result:
            return result
        
        return result


class URLResolverTool(Tool):
    """Tool for retrieving detailed content from specified URLs."""
    
    def __init__(self):
        super().__init__(
            name="url_resolver",
            description="Retrieves detailed content from specified URLs. Use only when the user's message contains a URL or domain—e.g. to create a video from a URL: call url_resolver first, then use the content for text_to_video. Do not call this tool for video creation if the user did not provide a URL. Also use for extracting full webpage content from search results or other sources.",
            instructions="Use this tool only when the user has provided a URL or domain. When the user wants to create a video from a URL, call url_resolver with that URL first, then use the returned content to build the video's text_prompt. Do not call url_resolver when creating a video from text or image only (no URL in user input)."
        )
    
    def get_parameters_schema(self) -> List[Dict[str, Any]]:
        return [
            {
                "name": "url",
                "description": "The URL to retrieve content from.",
                "type": "string",
                "required": True
            }
        ]
    
    def execute(self, parameters: Dict[str, Any]) -> Dict[str, Any]:
        url = (parameters.get('url') or '').strip()
        if not url:
            return {"error": "URL parameter is required"}

        try:
            # Normalize HTTP/HTTPS URLs: add scheme if missing (e.g. "www.linkedin.com" -> "https://www.linkedin.com")
            if not url.startswith('file://') and not url.startswith('http://') and not url.startswith('https://'):
                url = 'https://' + url

            # Handle file:// URLs
            if url.startswith('file://'):
                file_path = url[7:]  # Remove 'file://' prefix
                # Handle URL-encoded paths
                import urllib.parse
                file_path = urllib.parse.unquote(file_path)
                
                # Check if file path is valid
                if not file_path or file_path.strip() == '':
                    return {
                        "error": "Invalid file:// URL: empty file path. Please provide a valid file path or use the file upload API.",
                        "suggestion": "Use the upload_file API endpoint to upload files instead of file:// URLs"
                    }
                
                # Security: Prevent path traversal attacks
                # Normalize path and check for directory traversal attempts
                normalized_path = os.path.normpath(file_path)
                if '..' in file_path or normalized_path != file_path:
                    logger.warning(f"[URLResolver] Path traversal attempt blocked: {file_path}")
                    return {
                        "error": "Security violation: Path traversal detected. Use the file upload API instead.",
                        "suggestion": "Use the upload_file API endpoint to upload files instead of file:// URLs"
                    }
                
                # Security: Restrict to allowed directories only
                from django.conf import settings
                allowed_base = getattr(settings, 'MEDIA_ROOT', None) or os.path.join(settings.BASE_DIR, 'media')
                allowed_base = os.path.abspath(allowed_base)
                
                # Only allow files within MEDIA_ROOT or the project base directory.
                # NOTE: must be BASE_DIR itself, not its parent — the parent of BASE_DIR
                # resolves to the filesystem root ('/'), which would allow reading any
                # absolute path (e.g. /etc/passwd).
                workspace_dir = os.path.abspath(settings.BASE_DIR)
                file_abs_path = os.path.abspath(file_path)
                
                # Check if file is within allowed directories
                if not (file_abs_path.startswith(allowed_base) or file_abs_path.startswith(workspace_dir)):
                    logger.warning(f"[URLResolver] File access outside allowed directories blocked: {file_path}")
                    return {
                        "error": "Security violation: File access outside allowed directories. Use the file upload API instead.",
                        "suggestion": "Use the upload_file API endpoint to upload files instead of file:// URLs"
                    }
                
                # Check if file exists
                if not os.path.exists(file_path):
                    return {
                        "error": f"File not found: {file_path}. Please ensure the file path is correct and the file exists.",
                        "file_path": file_path,
                        "suggestion": "Use the upload_file API endpoint to upload files instead of file:// URLs"
                    }
                
                # Try to read text files
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()[:5000]  # Limit to first 5000 chars
                    return {
                        "url": url,
                        "content": content,
                        "status_code": 200,
                        "content_type": "text/plain",
                        "file_path": file_path,
                        "note": "File content retrieved from local filesystem"
                    }
                except UnicodeDecodeError:
                    # Binary file - return error message
                    return {
                        "error": f"File is binary or not readable as text: {file_path}. Please use the file upload API instead.",
                        "file_path": file_path,
                        "suggestion": "Use the upload_file API endpoint to upload binary files (PDFs, images, etc.)"
                    }
                except PermissionError:
                    return {
                        "error": f"Permission denied reading file: {file_path}",
                        "file_path": file_path
                    }
                except Exception as e:
                    return {
                        "error": f"Error reading file: {str(e)}",
                        "file_path": file_path
                    }
            
            # Handle HTTP/HTTPS URLs
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            response = requests.get(url, headers=headers, timeout=30)
            response.raise_for_status()
            
            # Extract text content (simplified - in production, use BeautifulSoup or similar)
            content = response.text[:5000]  # Limit to first 5000 chars
            
            return {
                "url": url,
                "content": content,
                "status_code": response.status_code,
                "content_type": response.headers.get('Content-Type', '')
            }
        except Exception as e:
            logger.error(f"Error resolving URL: {str(e)}", exc_info=True)
            return {"error": f"Failed to retrieve URL: {str(e)}"}


class CodeExecutorTool(Tool):
    """Tool for executing code and returning results."""
    
    def __init__(self):
        super().__init__(
            name="code_executor",
            description="This tool executes code and returns the result. Used for math, data analysis, and more.",
            instructions="Use this tool for calculations, data processing, or any task that requires code execution. Always validate inputs before execution."
        )
    
    def get_parameters_schema(self) -> List[Dict[str, Any]]:
        return [
            {
                "name": "code",
                "description": "The code to execute (Python).",
                "type": "string",
                "required": True
            },
            {
                "name": "language",
                "description": "Programming language (default: python).",
                "type": "string",
                "required": False
            }
        ]
    
    def _validate_code_security(self, code: str) -> tuple[bool, str]:
        """Validate code for security risks."""
        # Block dangerous imports and functions
        dangerous_patterns = [
            (r'import\s+os\s*$', 'os module'),
            (r'from\s+os\s+import', 'os module'),
            (r'import\s+subprocess\s*$', 'subprocess module'),
            (r'from\s+subprocess\s+import', 'subprocess module'),
            (r'import\s+sys\s*$', 'sys module'),
            (r'from\s+sys\s+import', 'sys module'),
            (r'import\s+shutil\s*$', 'shutil module'),
            (r'from\s+shutil\s+import', 'shutil module'),
            (r'import\s+__builtin__', 'builtin module'),
            (r'import\s+builtins', 'builtins module'),
            (r'__import__\s*\(', '__import__ function'),
            (r'eval\s*\(', 'eval function'),
            (r'exec\s*\(', 'exec function'),
            (r'compile\s*\(', 'compile function'),
            (r'open\s*\(', 'open function'),
            (r'file\s*\(', 'file function'),
            (r'input\s*\(', 'input function'),
            (r'raw_input\s*\(', 'raw_input function'),
            (r'execfile\s*\(', 'execfile function'),
            (r'reload\s*\(', 'reload function'),
            (r'__builtins__', '__builtins__ access'),
            (r'__import__', '__import__ access'),
            (r'\.__dict__', '__dict__ access'),
            (r'\.__class__', '__class__ access'),
            (r'\.__bases__', '__bases__ access'),
            (r'\.__subclasses__', '__subclasses__ access'),
            (r'\.__globals__', '__globals__ access'),
            (r'\.__getattribute__', '__getattribute__ access'),
            (r'\.__getattr__', '__getattr__ access'),
            (r'\.__setattr__', '__setattr__ access'),
            (r'\.__delattr__', '__delattr__ access'),
            (r'\.__call__', '__call__ access'),
            (r'importlib', 'importlib module'),
            (r'imp\s+', 'imp module'),
            (r'ctypes', 'ctypes module'),
            (r'socket\s*\.', 'socket module'),
            (r'urllib', 'urllib module'),
            (r'requests\s*\.', 'requests module'),
            (r'http\s*\.', 'http module'),
            (r'ftplib', 'ftplib module'),
            (r'telnetlib', 'telnetlib module'),
            (r'smtplib', 'smtplib module'),
            (r'subprocess\s*\.', 'subprocess calls'),
            (r'os\s*\.', 'os calls'),
            (r'sys\s*\.', 'sys calls'),
            (r'shutil\s*\.', 'shutil calls'),
            (r'pickle', 'pickle module'),
            (r'marshal', 'marshal module'),
            (r'shelve', 'shelve module'),
            (r'dbhash', 'dbhash module'),
            (r'gdbm', 'gdbm module'),
            (r'anydbm', 'anydbm module'),
            (r'whichdb', 'whichdb module'),
            (r'posix', 'posix module'),
            (r'nt', 'nt module'),
            (r'pwd', 'pwd module'),
            (r'grp', 'grp module'),
            (r'termios', 'termios module'),
            (r'tty', 'tty module'),
            (r'pty', 'pty module'),
            (r'fcntl', 'fcntl module'),
            (r'pipes', 'pipes module'),
            (r'resource', 'resource module'),
            (r'pwd', 'pwd module'),
            (r'signal', 'signal module'),
            (r'threading', 'threading module'),
            (r'multiprocessing', 'multiprocessing module'),
            (r'queue', 'queue module'),
            (r'select', 'select module'),
            (r'asyncore', 'asyncore module'),
            (r'asynchat', 'asynchat module'),
            (r'xmlrpc', 'xmlrpc module'),
            (r'SimpleXMLRPCServer', 'SimpleXMLRPCServer'),
            (r'DocXMLRPCServer', 'DocXMLRPCServer'),
            (r'CGIHTTPServer', 'CGIHTTPServer'),
            (r'BaseHTTPServer', 'BaseHTTPServer'),
            (r'SimpleHTTPServer', 'SimpleHTTPServer'),
            (r'commands', 'commands module'),
            (r'dircache', 'dircache module'),
            (r'filecmp', 'filecmp module'),
            (r'glob', 'glob module'),
            (r'fnmatch', 'fnmatch module'),
            (r'linecache', 'linecache module'),
            (r'shlex', 'shlex module'),
            (r'macpath', 'macpath module'),
            (r'macurl2path', 'macurl2path module'),
            (r'ntpath', 'ntpath module'),
            (r'nturl2path', 'nturl2path module'),
            (r'posixpath', 'posixpath module'),
            (r'stat', 'stat module'),
            (r'statvfs', 'statvfs module'),
            (r'tarfile', 'tarfile module'),
            (r'zipfile', 'zipfile module'),
            (r'gzip', 'gzip module'),
            (r'bz2', 'bz2 module'),
            (r'lzma', 'lzma module'),
            (r'zlib', 'zlib module'),
            (r'hashlib', 'hashlib module'),
            (r'hmac', 'hmac module'),
            (r'secrets', 'secrets module'),
            (r'base64', 'base64 module'),
            (r'binascii', 'binascii module'),
            (r'struct', 'struct module'),
            (r'codecs', 'codecs module'),
            (r'encodings', 'encodings module'),
            (r'locale', 'locale module'),
            (r'gettext', 'gettext module'),
            (r'argparse', 'argparse module'),
            (r'getopt', 'getopt module'),
            (r'optparse', 'optparse module'),
            (r'cmd', 'cmd module'),
            (r'shlex', 'shlex module'),
            (r'configparser', 'configparser module'),
            (r'fileinput', 'fileinput module'),
            (r'linecache', 'linecache module'),
            (r'netrc', 'netrc module'),
            (r'xdrlib', 'xdrlib module'),
            (r'plistlib', 'plistlib module'),
            (r'logging', 'logging module'),
            (r'warnings', 'warnings module'),
            (r'dis', 'dis module'),
            (r'pickletools', 'pickletools module'),
            (r'pdb', 'pdb module'),
            (r'profile', 'profile module'),
            (r'pstats', 'pstats module'),
            (r'timeit', 'timeit module'),
            (r'trace', 'trace module'),
            (r'tracemalloc', 'tracemalloc module'),
            (r'gc', 'gc module'),
            (r'inspect', 'inspect module'),
            (r'ast', 'ast module'),
            (r'symtable', 'symtable module'),
            (r'symbol', 'symbol module'),
            (r'token', 'token module'),
            (r'tokenize', 'tokenize module'),
            (r'keyword', 'keyword module'),
            (r'parser', 'parser module'),
            (r'py_compile', 'py_compile module'),
            (r'compileall', 'compileall module'),
            (r'dis', 'dis module'),
            (r'pickletools', 'pickletools module'),
            (r'doctest', 'doctest module'),
            (r'unittest', 'unittest module'),
            (r'2to3', '2to3 module'),
            (r'lib2to3', 'lib2to3 module'),
            (r'test', 'test module'),
            (r'test\.', 'test module'),
            (r'pydoc', 'pydoc module'),
            (r'doctest', 'doctest module'),
            (r'unittest', 'unittest module'),
            (r'pdb', 'pdb module'),
            (r'profile', 'profile module'),
            (r'pstats', 'pstats module'),
            (r'timeit', 'timeit module'),
            (r'trace', 'trace module'),
            (r'tracemalloc', 'tracemalloc module'),
            (r'gc', 'gc module'),
            (r'inspect', 'inspect module'),
            (r'ast', 'ast module'),
            (r'symtable', 'symtable module'),
            (r'symbol', 'symbol module'),
            (r'token', 'token module'),
            (r'tokenize', 'tokenize module'),
            (r'keyword', 'keyword module'),
            (r'parser', 'parser module'),
            (r'py_compile', 'py_compile module'),
            (r'compileall', 'compileall module'),
        ]
        
        # Check code length
        if len(code) > 10000:  # Max 10KB of code
            return False, "Code exceeds maximum length (10KB)"
        
        # Check for dangerous patterns
        for pattern, description in dangerous_patterns:
            # Bare identifier patterns (e.g. 'nt', 'gc') must match whole words only.
            # Without word boundaries they false-positive on substrings, e.g. 'nt'
            # inside 'print' or 'stat' inside larger identifiers.
            search_pattern = pattern
            if re.fullmatch(r'\w+', pattern):
                search_pattern = r'\b' + pattern + r'\b'
            if re.search(search_pattern, code, re.MULTILINE | re.IGNORECASE):
                return False, f"Security violation: {description} is not allowed"
        
        # Check for nested quotes and other injection attempts
        if code.count('"""') > 2 or code.count("'''") > 2:
            return False, "Security violation: Suspicious quote nesting detected"
        
        return True, ""
    
    def execute(self, parameters: Dict[str, Any]) -> Dict[str, Any]:
        code = parameters.get('code', '')
        language = parameters.get('language', 'python')
        
        if not code:
            return {"error": "Code parameter is required"}
        
        if language != 'python':
            return {"error": f"Language {language} not supported. Only Python is supported."}
        
        # Security validation
        is_safe, error_msg = self._validate_code_security(code)
        if not is_safe:
            logger.warning(f"[CodeExecutor] Security violation blocked: {error_msg}")
            return {"error": f"Security violation: {error_msg}"}
        
        try:
            # Create a temporary isolated directory for execution
            import tempfile
            import pathlib
            
            # Create temporary directory for this execution
            temp_dir = tempfile.mkdtemp(prefix='code_exec_')
            temp_dir_path = pathlib.Path(temp_dir)
            
            # Get list of files before execution (in temp dir)
            files_before = set()
            if temp_dir_path.exists():
                for ext in ['.py', '.txt', '.json', '.csv', '.md']:
                    files_before.update(temp_dir_path.rglob(f'*{ext}'))
            
            # Set resource limits function for subprocess
            def set_limits():
                try:
                    import resource
                    # Limit CPU time to 5 seconds
                    resource.setrlimit(resource.RLIMIT_CPU, (5, 5))
                    # Limit memory to 128MB
                    resource.setrlimit(resource.RLIMIT_AS, (128 * 1024 * 1024, 128 * 1024 * 1024))
                    # Limit file size to 10MB
                    resource.setrlimit(resource.RLIMIT_FSIZE, (10 * 1024 * 1024, 10 * 1024 * 1024))
                except Exception:
                    pass  # Resource limits not available on this system
            
            # Execute code in isolated temporary directory with strict limits
            result = subprocess.run(
                ['python3', '-c', code],
                capture_output=True,
                text=True,
                timeout=5,  # Reduced timeout to 5 seconds
                cwd=str(temp_dir_path),  # Execute in isolated temp directory
                env={**os.environ, 'PYTHONPATH': '', 'PYTHONHOME': ''},  # Clear Python path
                preexec_fn=set_limits if hasattr(os, 'setrlimit') else None
            )
            
            # Get list of files after execution (in temp dir)
            files_after = set()
            created_files = []
            if temp_dir_path.exists():
                for ext in ['.py', '.txt', '.json', '.csv', '.md']:
                    files_after.update(temp_dir_path.rglob(f'*{ext}'))
                
                # Find newly created files
                new_files = files_after - files_before
                for file_path in new_files:
                    # Only include files created in the last few seconds
                    try:
                        import time
                        file_mtime = file_path.stat().st_mtime
                        current_time = time.time()
                        if current_time - file_mtime < 10:  # Created in last 10 seconds
                            created_files.append(str(file_path))
                    except Exception as e:
                        logger.warning(f"[CodeExecutor] Error checking file {file_path}: {e}")
            
            response = {
                "stdout": result.stdout[:5000] if result.stdout else "",  # Limit output to 5KB
                "stderr": result.stderr[:5000] if result.stderr else "",  # Limit error output to 5KB
                "return_code": result.returncode,
                "success": result.returncode == 0
            }
            
            # Add file_path if files were created
            if created_files:
                # Use the first created file (or could return all)
                response["file_path"] = created_files[0]
                response["created_files"] = created_files
            
            # Cleanup: Remove temporary directory and all its contents
            try:
                import shutil
                shutil.rmtree(temp_dir, ignore_errors=True)
            except Exception as e:
                logger.warning(f"[CodeExecutor] Failed to cleanup temp directory {temp_dir}: {e}")
            
            return response
        except subprocess.TimeoutExpired:
            return {"error": "Code execution timed out (5 second limit)"}
        except Exception as e:
            logger.error(f"Error executing code: {str(e)}", exc_info=True)
            return {"error": f"Execution failed: {str(e)}"}
        finally:
            # Ensure cleanup even on error
            try:
                import shutil
                if 'temp_dir' in locals() and os.path.exists(temp_dir):
                    shutil.rmtree(temp_dir, ignore_errors=True)
            except Exception:
                pass
    


class TranscriptionTool(Tool):
    """Tool for transcribing audio and video files."""
    
    def __init__(self):
        super().__init__(
            name="transcription",
            description="Transcribes audio and video files into text.",
            instructions="Use this tool when users provide audio or video files that need to be converted to text."
        )
    
    def get_parameters_schema(self) -> List[Dict[str, Any]]:
        return [
            {
                "name": "file_path",
                "description": "Path to the audio or video file to transcribe.",
                "type": "string",
                "required": True
            }
        ]
    
    def execute(self, parameters: Dict[str, Any]) -> Dict[str, Any]:
        file_path = parameters.get('file_path', '')
        if not file_path:
            return {"error": "file_path parameter is required"}
        
        try:
            # Try AWS Transcribe if available
            try:
                from .aws_utils import create_boto3_client
                create_boto3_client('transcribe')
                
                # For now, return a note that AWS Transcribe requires async job processing
                # In production, implement async transcription job handling
                return {
                    "transcription": f"Transcription initiated for file: {file_path}",
                    "note": "AWS Transcribe requires async job processing. Configure transcription job handling for full functionality.",
                    "file_path": file_path
                }
            except Exception:
                pass
            
            # Fallback: return placeholder
            return {
                "transcription": f"Transcription service not yet implemented. File: {file_path}",
                "note": "Configure a transcription service (e.g., OpenAI Whisper, AWS Transcribe) for full functionality"
            }
        except Exception as e:
            logger.error(f"Error in transcription: {str(e)}", exc_info=True)
            return {"error": f"Transcription failed: {str(e)}"}


class TextToImageTool(Tool):
    """Tool for generating images from text."""
    
    def __init__(self):
        super().__init__(
            name="text_to_image",
            description="This tool generates an image from the given text prompt.",
            instructions=(
                "When the user wants a still image, call this tool immediately with ONLY the image description in "
                "`text_prompt` (subject, style, composition, lighting). Do not put apologies, policy text, URLs, "
                'sign-offs, or "we need to…" planning inside `text_prompt`—that breaks AWS content filters. '
                'Emit only: {"tool_calls":[{"tool":"text_to_image","parameters":{"text_prompt":"...","aspect_ratio":"16:9"}}]}'
            )
        )
    
    # Amazon Titan Image (v1/v2) enforces maxLength 512 on textToImageParams.text (Bedrock schema).
    TITAN_IMAGE_TEXT_MAX_LEN = 512
    # Stability SD3 / Stable Image Core allow up to 10k chars on prompt.
    STABILITY_PROMPT_MAX_LEN = 10000

    @classmethod
    def _clip_text_for_titan(cls, text: str) -> tuple:
        """Return (text_for_titan_api, was_truncated). Titan rejects prompts over 512 characters."""
        if not text:
            return "", False
        limit = cls.TITAN_IMAGE_TEXT_MAX_LEN
        if len(text) <= limit:
            return text, False
        chunk = text[:limit]
        sp = chunk.rfind(" ", 48, limit)
        if sp > limit // 3:
            chunk = chunk[:sp].rstrip()
        return chunk[:limit], True

    @classmethod
    def _clip_text_for_stability(cls, text: str) -> str:
        if not text:
            return ""
        if len(text) <= cls.STABILITY_PROMPT_MAX_LEN:
            return text
        return text[: cls.STABILITY_PROMPT_MAX_LEN]

    @classmethod
    def _sanitize_image_prompt_for_api(cls, text: str) -> str:
        """Strip chain-of-thought, apologies, and policy blobs often leaked into tool JSON text_prompt."""
        if not text:
            return ""
        t = text.replace('\r\n', '\n').strip()
        # Normalize unicode hyphens/dashes that can appear in pasted briefs
        t = re.sub(r'[\u2010-\u2015\u2212\uFE58\uFE63\uFF0D]', '-', t)
        # Drop unclosed or closed reasoning blocks
        t = re.sub(r'<reasoning>[\s\S]*?(?:</reasoning>|$)', '', t, flags=re.IGNORECASE).strip()
        # Assistant planning / "give up and apologize" text (model puts this inside parameters)
        markers = (
            '\n\nWe need to ',
            '\n\nAccording to ',
            '\n\nThe image generation',
            '\n\nI ran the ',
            '\n\nHey there,',
            "\n\nLet's ",
            '\n\nLet us ',
            '\n\nBelow is',
            '\n\n---',
            '\n\nFeel free to',
            '\n\nMeanwhile,',
            '\n\nLooking forward',
            '\n\nPaste it into',
            '\n\nOur filters',
            '\n\nError:',
        )
        cut = len(t)
        for m in markers:
            i = t.find(m)
            if i != -1 and i < cut:
                cut = i
        t = t[:cut].strip()
        # If first line is the image brief and second line starts a plan, keep first line only
        parts = t.split('\n', 1)
        if len(parts) == 2:
            second = parts[1].lstrip()
            if second.lower().startswith(('we need', 'according to', 'the user asked', 'policy')):
                t = parts[0].strip()
        return t.strip()

    def get_parameters_schema(self) -> List[Dict[str, Any]]:
        return [
            {
                "name": "text_prompt",
                "description": (
                    "Detailed description of the image. (AWS Titan Image accepts up to 512 characters; "
                    "longer prompts are clipped for Titan and full length is used for Stability / OpenAI when those run.)"
                ),
                "type": "string",
                "required": True
            },
            {
                "name": "aspect_ratio",
                "description": "Aspect ratio for the image (e.g., '1:1', '16:9', '9:16').",
                "type": "string",
                "required": False
            }
        ]
    
    def execute(self, parameters: Dict[str, Any]) -> Dict[str, Any]:
        raw_prompt = (parameters.get('text_prompt') or '').strip()
        if not raw_prompt:
            return {"error": "text_prompt parameter is required"}

        text_prompt = self._sanitize_image_prompt_for_api(raw_prompt)
        if not text_prompt.strip():
            text_prompt = raw_prompt.split('\n')[0].strip()[: self.STABILITY_PROMPT_MAX_LEN]
        if raw_prompt != text_prompt:
            logger.info(
                "text_to_image: sanitized text_prompt %d -> %d chars (removed non-visual / leaked assistant text)",
                len(raw_prompt),
                len(text_prompt),
            )

        aspect_ratio = parameters.get('aspect_ratio', '1:1')
        try:
            # Try AWS Bedrock Titan Image Generation if available
            try:
                from botocore.exceptions import ClientError
                from .aws_utils import create_boto3_client, get_aws_region
                
                region = get_aws_region()
                bedrock_runtime = create_boto3_client('bedrock-runtime', region_name=region)
                bedrock_client = create_boto3_client('bedrock', region_name=region)
                
                # First, try to list available models to find the correct Titan Image Generator model ID
                titan_image_models = []
                try:
                    # List foundation models to find available Titan Image Generator
                    response = bedrock_client.list_foundation_models(
                        byOutputModality='IMAGE',
                        byInferenceType='ON_DEMAND'
                    )
                    for model in response.get('modelSummaries', []):
                        model_id = model.get('modelId', '')
                        if 'titan' in model_id.lower() and 'image' in model_id.lower():
                            titan_image_models.append(model_id)
                            logger.info(f"Found Titan Image model: {model_id}")
                except Exception as e:
                    logger.debug(f"Could not list foundation models: {str(e)}")
                
                # Fallback to known model IDs if listing didn't work
                # Note: Titan Image Generator model IDs (correct as of 2024)
                # Put v2:0 first since it's the most commonly available and we know the format
                if not titan_image_models:
                    titan_image_models = [
                        "amazon.titan-image-generator-v2:0",  # Most commonly available, try first
                        "amazon.titan-image-generator-v1",  # Fallback option
                        "amazon.titan-image-generator-v2",  # If available
                        "amazon.titan-image-generator-v1:0",  # Some regions may use this format
                    ]
                else:
                    # Sort discovered models to try v2:0 first (it's most commonly available)
                    # Priority: models with :0 suffix first, then v2 models
                    titan_image_models.sort(key=lambda x: (':0' not in x, 'v2' not in x))
                
                logger.info(f"Titan Image Generator models to try: {titan_image_models}")
                
                # Map aspect ratios to Titan format
                aspect_ratio_map = {
                    '1:1': '1:1',
                    '16:9': '16:9',
                    '9:16': '9:16',
                    '4:3': '4:3',
                    '3:4': '3:4'
                }
                titan_aspect = aspect_ratio_map.get(aspect_ratio, '1:1')
                titan_text, titan_truncated = self._clip_text_for_titan(text_prompt)
                if titan_truncated:
                    logger.info(
                        "Titan Image: text_prompt clipped from %d to %d chars (Bedrock maxLength %d)",
                        len(text_prompt),
                        len(titan_text),
                        self.TITAN_IMAGE_TEXT_MAX_LEN,
                    )

                # Try each model ID until one works
                last_error = None
                for model_id in titan_image_models:
                    try:
                        # Prepare request body based on model version
                        # v2:0 models don't support aspectRatio parameter
                        # Check for :0 suffix first (most specific check)
                        if ':0' in model_id:
                            # Titan Image Generator v2:0 format (no aspectRatio)
                            logger.debug(f"Using v2:0 format (no aspectRatio) for model {model_id}")
                            body = json.dumps({
                                "taskType": "TEXT_IMAGE",
                                "textToImageParams": {
                                    "text": titan_text
                                }
                            })
                        elif 'v2' in model_id:
                            # Titan Image Generator v2 format (with aspectRatio)
                            body = json.dumps({
                                "taskType": "TEXT_IMAGE",
                                "textToImageParams": {
                                    "text": titan_text,
                                    "aspectRatio": titan_aspect
                                }
                            })
                        else:
                            # Titan Image Generator v1 format
                            body = json.dumps({
                                "taskType": "TEXT_IMAGE",
                                "textToImageParams": {
                                    "text": titan_text,
                                    "aspectRatio": titan_aspect
                                }
                            })
                        
                        # Invoke Titan model
                        response = bedrock_runtime.invoke_model(
                            modelId=model_id,
                            body=body,
                            contentType="application/json",
                            accept="application/json"
                        )
                        
                        result = json.loads(response['body'].read())
                        
                        # Save image to storage
                        import base64
                        images = result.get('images', [])
                        if not images:
                            raise ValueError("No image data returned from model")
                        
                        # Handle both formats: images can be list of dicts or list of base64 strings
                        image_item = images[0]
                        if isinstance(image_item, dict):
                            image_b64 = image_item.get('image', '')
                        else:
                            image_b64 = image_item  # Direct base64 string
                        
                        if not image_b64:
                            raise ValueError("No image data in response")
                        
                        image_data = base64.b64decode(image_b64)
                        file_name = f"generated_image_{abs(hash(text_prompt))}.png"
                        saved_path = default_storage.save(f"generated_images/{file_name}", ContentFile(image_data))
                        
                        return {
                            "image_url": default_storage.url(saved_path),
                            "file_name": file_name,
                            "success": True,
                            "model_used": model_id
                        }
                    except ClientError as e:
                        error_code = e.response.get('Error', {}).get('Code', '')
                        error_message = e.response.get('Error', {}).get('Message', str(e))
                        last_error = f"{error_code}: {error_message}"
                        
                        # If it's a validation error, try next model
                        if error_code == 'ValidationException':
                            logger.debug(f"Model {model_id} not available, trying next...")
                            continue
                        elif error_code == 'UnrecognizedClientException':
                            # Authentication error - don't try other models, return error immediately
                            logger.error(f"AWS authentication error with model {model_id}: {error_message}")
                            raise
                        else:
                            # Other errors, break and return
                            raise
                    except Exception as e:
                        last_error = str(e)
                        # Try next model
                        continue
                
                # If we get here, none of the models worked
                raise Exception(f"None of the Titan image models are available. Last error: {last_error}")
                
            except ClientError as e:
                error_code = e.response.get('Error', {}).get('Code', '')
                error_message = e.response.get('Error', {}).get('Message', str(e))
                logger.warning(f"AWS Bedrock Titan image generation failed: {error_code} - {error_message}")
                
                # Provide helpful error message based on error type
                if error_code == 'ValidationException':
                    return {
                        "image_url": "#",
                        "file_name": f"generated_image_{abs(hash(text_prompt))}.png",
                        "error": f"Image generation model not available in region {region}",
                        "note": f"The Titan Image Generator model may not be enabled in your AWS account or region ({region}). Titan Image Generator is typically available in us-east-1, us-west-2, ap-southeast-1, and eu-west-1 regions.",
                        "suggestion": f"1. Go to AWS Bedrock console\n2. Enable 'Amazon Titan Image Generator' model\n3. If not available in {region}, consider using a supported region (us-east-1, us-west-2, ap-southeast-1, eu-west-1)\n4. Check model availability: https://docs.aws.amazon.com/bedrock/latest/userguide/model-ids.html"
                    }
                elif error_code == 'UnrecognizedClientException':
                    return {
                        "image_url": "#",
                        "file_name": f"generated_image_{abs(hash(text_prompt))}.png",
                        "error": f"AWS Bedrock authentication error: {error_code}",
                        "note": f"The security token included in the request is invalid or AWS credentials are not configured properly. Region: {region}",
                        "suggestion": "1. Check AWS credentials are configured (AWS_ACCESS_KEY_ID, AWS_SECRET_ACCESS_KEY)\n2. Verify AWS region is set correctly (AWS_REGION or default region)\n3. Ensure credentials have permissions to use Bedrock models\n4. Check if using temporary credentials, ensure they haven't expired"
                    }
                else:
                    return {
                        "image_url": "#",
                        "file_name": f"generated_image_{abs(hash(text_prompt))}.png",
                        "error": f"AWS Bedrock error: {error_code}",
                        "note": error_message,
                        "suggestion": f"Check AWS Bedrock console and verify model access for region {region}"
                    }
            except Exception as e:
                logger.debug(f"AWS Bedrock Titan not available: {str(e)}")
                # Try alternative Bedrock image models (Stability AI, Flux, etc.)
                try:
                    return self._try_stability_ai_bedrock(text_prompt, aspect_ratio, region)
                except Exception as stability_error:
                    logger.debug(f"Stability AI via Bedrock also failed: {str(stability_error)}")
                    # Try OpenAI DALL-E if API key is available
                    try:
                        return self._try_openai_dalle(text_prompt, aspect_ratio)
                    except Exception as dalle_error:
                        logger.debug(f"OpenAI DALL-E also failed: {str(dalle_error)}")
                        # Final fallback - return helpful error with detailed diagnostics
                        import os
                        from django.conf import settings
                        from .aws_utils import get_aws_region
                        
                        # Collect diagnostic information
                        diagnostics = []
                        try:
                            region = get_aws_region()
                        except Exception:
                            region = getattr(settings, 'AWS_REGION', 'us-east-1')
                        diagnostics.append(f"Current AWS Region: {region}")
                        
                        # Check AWS credentials
                        aws_key_set = bool(getattr(settings, 'AWS_ACCESS_KEY_ID', '') or os.getenv('AWS_ACCESS_KEY_ID'))
                        diagnostics.append(f"AWS Credentials: {'Configured' if aws_key_set else 'Not configured'}")
                        
                        # Check OpenAI API key
                        openai_key_set = bool(
                            os.getenv('OPENAI_API_KEY')
                            or getattr(settings, 'OPENAI_API_KEY', None)
                        )
                        diagnostics.append(f"OpenAI API Key: {'Configured' if openai_key_set else 'Not configured'}")
                        
                        # Build error message (surface root causes in `error` for chat UIs that only show `error`)
                        error_details = []
                        error_details.append(f"AWS Bedrock Titan: {str(e)[:200]}")
                        error_details.append(f"Stability AI via Bedrock: {str(stability_error)[:200]}")
                        error_details.append(f"OpenAI DALL-E: {str(dalle_error)[:200]}")
                        short = "; ".join(error_details)[:900]
                        user_error = (
                            "Image generation is not available: Bedrock Titan, Stability on Bedrock, and OpenAI "
                            f"DALL-E all failed ({region}). "
                            "Enable Amazon Titan Image Generator (or another image model) in the AWS Bedrock console "
                            "for this account/region, ensure the app role can bedrock:InvokeModel, or set OPENAI_API_KEY. "
                            f"Details: {short}"
                        )
                        logger.warning(
                            "text_to_image: all backends failed. diagnostics=%s details=%s",
                            diagnostics,
                            error_details,
                        )
                        
                        # Check if this might be an architectural/technical drawing request
                        architectural_keywords = ['massing', 'elevation', 'floor plan', 'building', 'architecture', 
                                                  'diagram', 'schematic', 'blueprint', 'technical', 'dimensions',
                                                  'site plan', 'section', 'facade', 'layout']
                        is_architectural = any(kw in text_prompt.lower() for kw in architectural_keywords)
                        
                        svg_suggestion = ""
                        if is_architectural:
                            svg_suggestion = (
                                "\n\n**ALTERNATIVE: For architectural/technical drawings, use the svg_diagram tool:**\n"
                                "The svg_diagram tool can generate:\n"
                                "- Massing models with labeled dimensions\n"
                                "- Front/side elevations\n"
                                "- Floor plans with area calculations\n"
                                "Example: svg_diagram(diagram_type='massing_model', title='Building A', "
                                "width='50m', height='30m', depth='40m', style='modern')"
                            )
                        
                        return {
                            "image_url": "#",
                            "file_name": f"generated_image_{abs(hash(text_prompt))}.png",
                            "error": user_error,
                            "note": "\n".join(diagnostics),
                            "details": "\n".join(error_details),
                            "suggestion": (
                                "To enable image generation, configure one of the following:\n"
                                "1. AWS Bedrock Titan Image Generator:\n"
                                "   - Enable 'Amazon Titan Image Generator' in AWS Bedrock console\n"
                                f"   - Ensure it's available in region: {region}\n"
                                "   - Supported regions: us-east-1, us-west-2, ap-southeast-1, eu-west-1\n"
                                "2. Stability AI via Bedrock:\n"
                                "   - Enable Stability AI models in AWS Bedrock console\n"
                                "3. OpenAI DALL-E:\n"
                                "   - Set OPENAI_API_KEY (environment variable or Django settings)\n"
                                "   - Example: export OPENAI_API_KEY='sk-...'"
                                f"{svg_suggestion}"
                            )
                        }
        except Exception as e:
            logger.error(f"Error in image generation: {str(e)}", exc_info=True)
            return {"error": f"Image generation failed: {str(e)}"}
    
    def _try_stability_ai_bedrock(self, text_prompt: str, aspect_ratio: str, region: str) -> Dict[str, Any]:
        """Try Stability AI image generation via AWS Bedrock (SD3.5, Stable Image Core/Ultra, legacy SDXL)."""
        import base64
        from botocore.exceptions import ClientError
        from .aws_utils import create_boto3_client

        prompt = self._clip_text_for_stability(text_prompt)
        bedrock_runtime = create_boto3_client('bedrock-runtime', region_name=region)
        bedrock_client = create_boto3_client('bedrock', region_name=region)

        # Current Stability models on Bedrock use string aspect_ratio; legacy SDXL uses width/height.
        sd_prompt_aspect_ratios = {
            '1:1', '16:9', '21:9', '2:3', '3:2', '4:5', '5:4', '9:16', '9:21',
        }
        ar = aspect_ratio if aspect_ratio in sd_prompt_aspect_ratios else '1:1'

        aspect_map_legacy = {
            '1:1': {'width': 1024, 'height': 1024},
            '16:9': {'width': 1344, 'height': 768},
            '9:16': {'width': 768, 'height': 1344},
            '4:3': {'width': 1152, 'height': 896},
            '3:4': {'width': 896, 'height': 1152},
        }
        dimensions = aspect_map_legacy.get(aspect_ratio, {'width': 1024, 'height': 1024})

        def _is_legacy_sdxl(mid: str) -> bool:
            ml = mid.lower()
            return 'stable-diffusion-xl' in ml or 'sdxl' in ml

        stability_models: List[str] = []
        try:
            response = bedrock_client.list_foundation_models(
                byOutputModality='IMAGE',
                byInferenceType='ON_DEMAND'
            )
            discovered: List[str] = []
            for model in response.get('modelSummaries', []):
                mid = model.get('modelId', '')
                if 'stability' in mid.lower() or 'stable-diffusion' in mid.lower():
                    discovered.append(mid)

            def _rank(m: str) -> tuple:
                s = m.lower()
                if 'sd3-5-large' in s:
                    return (0, m)
                if 'sd3-ultra' in s or 'stable-image-ultra' in s:
                    return (1, m)
                if 'stable-image-core' in s:
                    return (2, m)
                if 'stable-diffusion-xl' in s:
                    return (9, m)
                return (5, m)

            stability_models = sorted(set(discovered), key=_rank)
            for mid in stability_models:
                logger.info(f"Found Stability AI model: {mid}")
            # Some accounts expose image models without ON_DEMAND filter only
            if not stability_models:
                response_b = bedrock_client.list_foundation_models(byOutputModality='IMAGE')
                for model in response_b.get('modelSummaries', []):
                    mid = model.get('modelId', '')
                    if 'stability' in mid.lower() or 'stable-diffusion' in mid.lower():
                        discovered.append(mid)
                stability_models = sorted(set(discovered), key=_rank)
                for mid in stability_models:
                    logger.info(f"Found Stability AI model (broad list): {mid}")
        except Exception as e:
            logger.debug(f"Could not list Stability AI models: {str(e)}")

        # Fallback IDs: prefer current models (SDXL v1 is EOL in many accounts).
        if not stability_models:
            stability_models = [
                'stability.sd3-5-large-v1:0',
                'stability.stable-image-core-v1:0',
                'stability.sd3-ultra-v1:1',
                'stability.stable-diffusion-xl-v2:0',
            ]

        def _decode_stability_image(result: Dict[str, Any]) -> bytes:
            reasons = result.get('finish_reasons') or []
            if any(r is not None for r in reasons):
                raise ValueError(f"Image filtered or failed: {reasons}")
            imgs = result.get('images') or []
            if imgs:
                raw = imgs[0]
                if isinstance(raw, dict):
                    raw = raw.get('image') or raw.get('base64') or ''
                if not raw:
                    raise ValueError("Empty images[0] in Stability response")
                return base64.b64decode(raw)
            arts = result.get('artifacts') or []
            if arts and isinstance(arts[0], dict) and arts[0].get('base64'):
                return base64.b64decode(arts[0]['base64'])
            raise ValueError("No image data in Stability response")

        last_stability_error: Optional[str] = None
        for model_id in stability_models:
            try:
                if _is_legacy_sdxl(model_id):
                    body = json.dumps({
                        'text_prompts': [{'text': prompt}],
                        'cfg_scale': 7,
                        'steps': 30,
                        'width': dimensions['width'],
                        'height': dimensions['height'],
                    })
                else:
                    body = json.dumps({
                        'prompt': prompt,
                        'aspect_ratio': ar,
                        'output_format': 'png',
                    })

                response = bedrock_runtime.invoke_model(
                    modelId=model_id,
                    body=body,
                    contentType='application/json',
                    accept='application/json',
                )
                result = json.loads(response['body'].read())
                image_data = _decode_stability_image(result)
                file_name = f"generated_image_{abs(hash(text_prompt))}.png"
                saved_path = default_storage.save(f'generated_images/{file_name}', ContentFile(image_data))

                return {
                    'image_url': default_storage.url(saved_path),
                    'file_name': file_name,
                    'success': True,
                    'model_used': model_id,
                    'service': 'stability-ai-bedrock',
                }
            except ClientError as e:
                err = e.response.get('Error', {}) if e.response else {}
                msg = f"{err.get('Code', 'ClientError')}: {err.get('Message', str(e))}"
                last_stability_error = f"{model_id}: {msg}"
                error_code = err.get('Code', '')
                if error_code in ('ValidationException', 'ResourceNotFoundException'):
                    logger.warning("Stability model %s skipped: %s", model_id, msg[:500])
                    continue
                raise
            except Exception as e:
                last_stability_error = f"{model_id}: {str(e)}"
                logger.warning("Stability AI model %s failed: %s", model_id, str(e)[:500])
                continue

        raise Exception(
            f"No Stability AI model succeeded in {region}. Models tried: {stability_models}. "
            f"Last error: {last_stability_error or 'unknown'}"
        )
    
    def _try_openai_dalle(self, text_prompt: str, aspect_ratio: str) -> Dict[str, Any]:
        """Try OpenAI DALL-E image generation."""
        import os
        from django.conf import settings

        api_key = (
            os.environ.get('OPENAI_API_KEY')
            or getattr(settings, 'OPENAI_API_KEY', None)
            or ''
        )
        if not api_key:
            raise Exception("OPENAI_API_KEY not set (env or Django settings)")
        
        try:
            import openai
            openai.api_key = api_key
            
            # Map aspect ratios to DALL-E sizes
            size_map = {
                '1:1': '1024x1024',
                '16:9': '1792x1024',  # DALL-E 3 supports wider formats
                '9:16': '1024x1792',
                '4:3': '1024x1024',  # Closest available
                '3:4': '1024x1024',  # Closest available
            }
            size = size_map.get(aspect_ratio, '1024x1024')
            
            # Try DALL-E 3 first (better quality)
            try:
                response = openai.Image.create(
                    model="dall-e-3",
                    prompt=text_prompt,
                    size=size,
                    n=1,
                    quality="standard"
                )
                image_url = response['data'][0]['url']
                
                # Download and save image
                import requests
                img_response = requests.get(image_url)
                img_response.raise_for_status()
                
                file_name = f"generated_image_{abs(hash(text_prompt))}.png"
                saved_path = default_storage.save(f"generated_images/{file_name}", ContentFile(img_response.content))
                
                return {
                    "image_url": default_storage.url(saved_path),
                    "file_name": file_name,
                    "success": True,
                    "model_used": "dall-e-3",
                    "service": "openai-dalle"
                }
            except Exception as e:
                # Fallback to DALL-E 2
                logger.debug(f"DALL-E 3 failed, trying DALL-E 2: {str(e)}")
                response = openai.Image.create(
                    model="dall-e-2",
                    prompt=text_prompt,
                    size=size,
                    n=1
                )
                image_url = response['data'][0]['url']
                
                # Download and save image
                import requests
                img_response = requests.get(image_url)
                img_response.raise_for_status()
                
                file_name = f"generated_image_{abs(hash(text_prompt))}.png"
                saved_path = default_storage.save(f"generated_images/{file_name}", ContentFile(img_response.content))
                
                return {
                    "image_url": default_storage.url(saved_path),
                    "file_name": file_name,
                    "success": True,
                    "model_used": "dall-e-2",
                    "service": "openai-dalle"
                }
        except ImportError:
            raise Exception("OpenAI library not installed. Install with: pip install openai")
        except Exception as e:
            raise Exception(f"OpenAI DALL-E error: {str(e)}")


def _enhance_video_prompt(prompt: str, min_length: int = 60, max_chars: int = 512) -> str:
    """
    Enhance a video prompt for better quality and text/graphics clarity.
    Preserves the user's prompt at the start; only appends a short quality suffix so the
    video matches what the user asked for (no overwriting or diluting the main content).
    """
    if not prompt or not isinstance(prompt, str):
        return prompt
    user_prompt = prompt.strip()
    # Short suffixes so user content stays within API limits (e.g. Nova 512 chars)
    text_graph_keywords = ('text', 'graph', 'chart', 'label', 'diagram', 'title', 'number', 'data', 'axis', 'caption', 'heading', 'slide', 'presentation', 'infographic', 'table', 'word', 'letter', 'stat')
    has_text_graph = any(k in user_prompt.lower() for k in text_graph_keywords)
    if has_text_graph:
        clarity_suffix = " Legible text and labels, sharp, 4K."
    else:
        clarity_suffix = " High quality, sharp, 4K."
    if not user_prompt.endswith('.'):
        user_prompt = user_prompt.rstrip() + "."
    # Reserve space for suffix so we never truncate the start of the user's prompt
    max_user_chars = max(min_length, max_chars - len(clarity_suffix))
    if len(user_prompt) > max_user_chars:
        user_prompt = user_prompt[: max_user_chars - 3].rsplit(' ', 1)[0] + "..."
    result = (user_prompt + clarity_suffix).strip()
    return result[:max_chars]


class TextToVideoTool(Tool):
    """Tool for generating videos from text prompts using Amazon Bedrock."""
    
    def __init__(self):
        super().__init__(
            name="text_to_video",
            description="Generates a video from a text prompt using Amazon Bedrock (Luma Ray 2 or Nova Reel). Call url_resolver only when the user's message contains a URL or domain; then use that content to build the text_prompt. When the user asks for a video from text only (no URL), do not call url_resolver—use a detailed text_prompt directly. Always call this tool when the user asks to create or generate a video; do not refuse without calling it first.",
            instructions="When the user asks to create a video and their message contains a URL: (1) Call url_resolver with that URL to get the website content. (2) Use that content to build a detailed text_prompt, then call text_to_video. When the user asks for a video from text only (no URL in their message), do not call url_resolver; use a detailed text_prompt directly. If the user did not specify video length, use 9s or 12s. Write the prompt as a video caption: subject, action, environment, lighting, style. Generation is asynchronous and may take 1-2 minutes."
        )
    
    def get_parameters_schema(self) -> List[Dict[str, Any]]:
        return [
            {
                "name": "text_prompt",
                "description": "The video scene description—must match what the user asked for (use their exact words or a close paraphrase; never use a generic or different topic). Include: subject, action, environment, lighting, style; optionally camera motion (e.g. 'Dolly shot of...'). Example: 'Cinematic dolly shot of a cheeseburger with melting cheese on a diner table. Natural lighting, steam, 4k, photorealistic.' Avoid no/not/without. Max 512 chars for 6s videos.",
                "type": "string",
                "required": True
            },
            {
                "name": "duration",
                "description": "Video length. If the user did not specify, ask (e.g. 'How long? 9 seconds (default) or up to 2 minutes?'). Luma: '5s' or '9s'. Nova: 6s (single shot) or 12–120s (longer, in 6s steps). Default: '9s' for longer clips.",
                "type": "string",
                "required": False
            },
            {
                "name": "aspect_ratio",
                "description": "Aspect ratio for the video (e.g., '16:9', '9:16', '1:1'). Default: '16:9'.",
                "type": "string",
                "required": False
            },
            {
                "name": "resolution",
                "description": "Video resolution. Use '720p' (recommended for text/graphs clarity). Options: '720p' or '540p' for Luma; '1280x720' for Nova. Default: 720p.",
                "type": "string",
                "required": False
            },
            {
                "name": "model",
                "description": "Model: 'luma' (Luma Ray 2, better for text/graphics and detail) or 'nova' (Nova Reel). Default: 'luma'.",
                "type": "string",
                "required": False
            }
        ]
    
    def execute(self, parameters: Dict[str, Any]) -> Dict[str, Any]:
        # Accept text_prompt or common alternate keys from LLM output
        text_prompt = (
            parameters.get('text_prompt')
            or parameters.get('prompt')
            or parameters.get('text')
            or parameters.get('caption')
            or ''
        )
        if isinstance(text_prompt, str):
            text_prompt = text_prompt.strip()
        else:
            text_prompt = str(text_prompt).strip() if text_prompt else ''
        duration = parameters.get('duration', '9s')
        # Normalize duration: accept int (15) or str ("5s", "9s", "15", "30")
        if isinstance(duration, int):
            duration = str(duration)
        elif not isinstance(duration, str):
            duration = '9s'
        duration = str(duration).strip() or '9s'
        aspect_ratio = parameters.get('aspect_ratio', '16:9')
        resolution = (parameters.get('resolution') or '720p').strip() or '720p'
        if resolution not in ('720p', '540p', '1280x720'):
            resolution = '720p'
        model_preference = parameters.get('model', 'luma').lower()
        
        if not text_prompt:
            return {"error": "text_prompt is required. Do not generate video without a real prompt. Use the user's video description (or content from conversation history) as text_prompt; do not use a placeholder like \"...\"."}
        # Reject placeholder or generic prompts so we never generate without a real description
        _pl = text_prompt.lower().strip()
        _placeholders = ('.', '..', '...', '…', 'n/a', 'na', 'tbd', 'prompt', 'text', 'caption', 'video description', 'enter prompt', 'user prompt', 'description here')
        if _pl in _placeholders or (_pl.replace('.', '').replace(' ', '') == '' and len(_pl) <= 10):
            return {"error": "Do not generate video without a real prompt. text_prompt must be the actual video description (what to show in the video), not a placeholder. Use the user's words or content from the conversation (e.g. url_resolver result, script, or their request)."}
        
        original_prompt = text_prompt
        # Enhance for quality and text/graphics clarity (allow up to 4000 chars for Nova multi-shot)
        text_prompt = _enhance_video_prompt(text_prompt, min_length=60, max_chars=4000)
        logger.info(
            "text_to_video prompt: original_len=%d enhanced_len=%d original_preview=%s sent_preview=%s",
            len(original_prompt), len(text_prompt),
            (original_prompt[:200] + "..." if len(original_prompt) > 200 else original_prompt),
            (text_prompt[:200] + "..." if len(text_prompt) > 200 else text_prompt),
        )
        
        try:
            from botocore.exceptions import ClientError
            from .aws_utils import create_boto3_client, get_aws_region
            import time
            import uuid
            
            region = get_aws_region()
            bedrock_runtime = create_boto3_client('bedrock-runtime', region_name=region)
            bedrock_client = create_boto3_client('bedrock', region_name=region)
            s3_client = create_boto3_client('s3', region_name=region)
            
            # Get S3 bucket for video output (required by Bedrock async API)
            s3_bucket = os.getenv('BEDROCK_VIDEO_OUTPUT_BUCKET', None)
            if not s3_bucket:
                # Try to get default bucket from settings or use a generated name
                s3_bucket = getattr(settings, 'BEDROCK_VIDEO_OUTPUT_BUCKET', None)
                if not s3_bucket:
                    # Generate a unique bucket name based on account ID
                    try:
                        sts_client = create_boto3_client('sts', region_name=region)
                        account_id = sts_client.get_caller_identity().get('Account')
                        s3_bucket = f"bedrock-video-output-{account_id}"
                    except Exception:
                        return {
                            "error": "S3 bucket not configured",
                            "note": "Set BEDROCK_VIDEO_OUTPUT_BUCKET environment variable or configure in settings",
                            "suggestion": "Create an S3 bucket and set BEDROCK_VIDEO_OUTPUT_BUCKET environment variable"
                        }
            
            # Generate unique S3 prefix for this video. Bedrock expects s3Uri to be a bucket or
            # directory (prefix), not a full object key; it writes output as output.mp4 under that prefix.
            video_id = str(uuid.uuid4())
            s3_prefix = f"videos/{video_id}/"
            s3_uri = f"s3://{s3_bucket}/{s3_prefix}"
            s3_key = f"videos/{video_id}/output.mp4"
            
            # Try Luma Ray 2 first if preferred, otherwise try Nova Reel
            models_to_try = []
            if model_preference == 'luma':
                models_to_try = [
                    ('luma.ray-v2:0', 'luma'),
                    ('amazon.nova-reel-v1:1', 'nova')
                ]
            else:
                models_to_try = [
                    ('amazon.nova-reel-v1:1', 'nova'),
                    ('luma.ray-v2:0', 'luma')
                ]
            
            # Check available models
            available_models = []
            try:
                response = bedrock_client.list_foundation_models(
                    byOutputModality='VIDEO',
                    byInferenceType='ON_DEMAND'
                )
                for model in response.get('modelSummaries', []):
                    model_id = model.get('modelId', '')
                    available_models.append(model_id)
                    logger.info(f"Found video model: {model_id}")
            except Exception as e:
                logger.debug(f"Could not list video models: {str(e)}")
            
            # Filter models to try based on availability
            if available_models:
                models_to_try = [(m, t) for m, t in models_to_try if m in available_models]
            
            if not models_to_try:
                # Fallback to known model IDs
                models_to_try = [
                    ('luma.ray-v2:0', 'luma'),
                    ('amazon.nova-reel-v1:1', 'nova')
                ]
            
            last_error = None
            invocation_arn = None
            model_used = None
            
            # Parse requested duration for Nova (seconds, multiple of 6; 12-120 for multi-shot)
            def _parse_duration_seconds(dur_str):
                s = str(dur_str).strip().lower().rstrip('s')
                try:
                    n = int(s)
                    return max(6, min(120, n))
                except ValueError:
                    return 9
            requested_seconds = _parse_duration_seconds(duration)
            # Luma: only 5s or 9s - prefer 9s for longer clip
            luma_duration = "9s" if requested_seconds >= 9 else "5s"
            # Nova: 6s single-shot, or 12-120 in steps of 6 for multi-shot
            nova_duration_seconds = 6 if requested_seconds <= 6 else (max(12, min(120, (requested_seconds // 6) * 6)))
            
            for model_id, model_type in models_to_try:
                try:
                    if model_type == 'luma':
                        # Luma Ray 2: 5s or 9s, 720p for best quality
                        model_input = {
                            "prompt": text_prompt[:5000],
                            "aspect_ratio": aspect_ratio,
                            "duration": luma_duration,
                            "resolution": "720p",
                            "loop": False
                        }
                    else:
                        # Nova Reel: TEXT_VIDEO (6s only) or MULTI_SHOT_AUTOMATED (12-120s)
                        if resolution == '720p' or resolution == '1280x720':
                            dimension = "1280x720"
                        elif resolution == '540p':
                            dimension = "960x540"
                        else:
                            dimension = "1280x720"
                        
                        if nova_duration_seconds <= 6:
                            model_input = {
                                "taskType": "TEXT_VIDEO",
                                "textToVideoParams": {"text": text_prompt[:512]},
                                "videoGenerationConfig": {
                                    "durationSeconds": 6,
                                    "fps": 24,
                                    "dimension": dimension,
                                    "seed": 0
                                }
                            }
                        else:
                            # Multi-shot: 12-120 seconds, prompt up to 4000 chars
                            model_input = {
                                "taskType": "MULTI_SHOT_AUTOMATED",
                                "multiShotAutomatedParams": {"text": text_prompt[:4000]},
                                "videoGenerationConfig": {
                                    "durationSeconds": nova_duration_seconds,
                                    "fps": 24,
                                    "dimension": dimension,
                                    "seed": 0
                                }
                            }
                    
                    # Start async invocation
                    response = bedrock_runtime.start_async_invoke(
                        modelId=model_id,
                        modelInput=model_input,
                        outputDataConfig={
                            "s3OutputDataConfig": {
                                "s3Uri": s3_uri
                            }
                        }
                    )
                    
                    invocation_arn = response.get('invocationArn')
                    model_used = model_id
                    logger.info(f"Started video generation with {model_id}, invocation ARN: {invocation_arn}")
                    break
                    
                except ClientError as e:
                    error_code = e.response.get('Error', {}).get('Code', '')
                    error_message = e.response.get('Error', {}).get('Message', str(e))
                    last_error = f"{error_code}: {error_message}"
                    
                    if error_code == 'ValidationException':
                        logger.debug(f"Model {model_id} not available, trying next...")
                        continue
                    elif error_code == 'UnrecognizedClientException':
                        raise
                    else:
                        # Try next model
                        continue
                except Exception as e:
                    last_error = str(e)
                    continue
            
            if not invocation_arn:
                return {
                    "error": "Failed to start video generation",
                    "note": f"None of the video models are available. Last error: {last_error}",
                    "suggestion": "1. Enable video generation models in AWS Bedrock console\n2. Ensure models are available in your region\n3. Check AWS credentials and permissions"
                }
            
            # Poll for completion (timeout from settings, default 10 min)
            max_wait_time = getattr(settings, 'BEDROCK_VIDEO_GENERATION_TIMEOUT', 600)
            poll_interval = 5  # Check every 5 seconds
            start_time = time.time()
            
            while time.time() - start_time < max_wait_time:
                try:
                    status_response = bedrock_runtime.get_async_invoke(
                        invocationArn=invocation_arn
                    )
                    
                    status = status_response.get('status')
                    
                    if status == 'Completed':
                        # Download video from S3. Prefer location from response, then s3_key, then list under prefix.
                        video_data = None
                        actual_key = None
                        out_cfg = status_response.get('outputDataConfig', {}).get('s3OutputDataConfig', {})
                        response_s3_uri = (out_cfg or {}).get('s3Uri') or ''
                        if response_s3_uri and response_s3_uri.startswith('s3://'):
                            # Parse s3://bucket/key-or-prefix
                            parts = response_s3_uri.replace('s3://', '', 1).split('/', 1)
                            resp_bucket = parts[0] if parts else s3_bucket
                            resp_key_or_prefix = (parts[1] or '').rstrip('/')
                            if resp_key_or_prefix and not resp_key_or_prefix.endswith('/'):
                                try:
                                    video_obj = s3_client.get_object(Bucket=resp_bucket, Key=resp_key_or_prefix)
                                    video_data = video_obj['Body'].read()
                                    actual_key = resp_key_or_prefix
                                except ClientError:
                                    pass
                            if not video_data and resp_key_or_prefix:
                                prefix = resp_key_or_prefix + '/' if not resp_key_or_prefix.endswith('/') else resp_key_or_prefix
                                list_resp = s3_client.list_objects_v2(Bucket=resp_bucket, Prefix=prefix, MaxKeys=50)
                                contents = list_resp.get('Contents', [])
                                for obj in contents:
                                    key = (obj.get('Key') or '').strip()
                                    if key.endswith('.mp4'):
                                        video_obj = s3_client.get_object(Bucket=resp_bucket, Key=key)
                                        video_data = video_obj['Body'].read()
                                        actual_key = key
                                        break
                                if not actual_key and contents:
                                    key = (contents[0].get('Key') or '').strip()
                                    if key:
                                        video_obj = s3_client.get_object(Bucket=resp_bucket, Key=key)
                                        video_data = video_obj['Body'].read()
                                        actual_key = key
                        if not video_data:
                            try:
                                video_obj = s3_client.get_object(Bucket=s3_bucket, Key=s3_key)
                                video_data = video_obj['Body'].read()
                                actual_key = s3_key
                            except ClientError as e:
                                if e.response.get('Error', {}).get('Code') != 'NoSuchKey':
                                    raise
                                # List objects under prefix; prefer .mp4, else first object (e.g. subfolder/output.mp4)
                                list_resp = s3_client.list_objects_v2(Bucket=s3_bucket, Prefix=s3_prefix, MaxKeys=50)
                                contents = list_resp.get('Contents', [])
                                mp4_key = None
                                fallback_key = None
                                for obj in contents:
                                    key = obj.get('Key', '').strip()
                                    if not key or key.endswith('/'):
                                        continue
                                    if key.endswith('.mp4'):
                                        mp4_key = key
                                        break
                                    if fallback_key is None:
                                        fallback_key = key
                                actual_key = mp4_key or fallback_key
                                if actual_key:
                                    video_obj = s3_client.get_object(Bucket=s3_bucket, Key=actual_key)
                                    video_data = video_obj['Body'].read()
                        if video_data and actual_key:
                            file_name = f"generated_video_{video_id}.mp4"
                            saved_path = default_storage.save(f"generated_videos/{file_name}", ContentFile(video_data))
                            try:
                                s3_client.delete_object(Bucket=s3_bucket, Key=actual_key)
                            except Exception as cleanup_error:
                                logger.warning(f"Failed to clean up S3 file: {str(cleanup_error)}")
                            return {
                                "video_url": default_storage.url(saved_path),
                                "file_name": file_name,
                                "success": True,
                                "model_used": model_used,
                                "invocation_arn": invocation_arn
                            }
                        return {
                            "error": "Video generation completed but download failed",
                            "note": f"Video is available at S3: {s3_uri}. No object found under prefix.",
                            "invocation_arn": invocation_arn,
                            "s3_uri": s3_uri,
                            "download_error": "NoSuchKey and no objects listed under prefix"
                        }
                    
                    elif status == 'Failed':
                        failure_reason = status_response.get('failureReason', 'Unknown error')
                        return {
                            "error": "Video generation failed",
                            "note": failure_reason,
                            "invocation_arn": invocation_arn
                        }
                    
                    # Still in progress, wait and check again
                    time.sleep(poll_interval)
                    
                except ClientError as e:
                    error_code = e.response.get('Error', {}).get('Code', '')
                    if error_code == 'ResourceNotFoundException':
                        return {
                            "error": "Video generation job not found",
                            "invocation_arn": invocation_arn
                        }
                    raise
            
            # Timeout
            return {
                "error": "Video generation timeout",
                "note": f"Video generation is still in progress after {max_wait_time} seconds",
                "invocation_arn": invocation_arn,
                "s3_uri": s3_uri,
                "suggestion": f"Check status later using invocation ARN: {invocation_arn}"
            }
            
        except ClientError as e:
            error_code = e.response.get('Error', {}).get('Code', '')
            error_message = e.response.get('Error', {}).get('Message', str(e))
            logger.warning(f"AWS Bedrock video generation failed: {error_code} - {error_message}")
            
            if error_code == 'ValidationException':
                return {
                    "error": f"Video generation model not available in region {region}",
                    "note": "Video generation models may not be enabled in your AWS account or region.",
                    "suggestion": "1. Go to AWS Bedrock console\n2. Enable video generation models (Luma Ray 2 or Amazon Nova Reel)\n3. Ensure models are available in your region"
                }
            elif error_code == 'UnrecognizedClientException':
                return {
                    "error": f"AWS Bedrock authentication error: {error_code}",
                    "note": "AWS credentials are not configured properly.",
                    "suggestion": "1. Check AWS credentials (AWS_ACCESS_KEY_ID, AWS_SECRET_ACCESS_KEY)\n2. Verify AWS region is set correctly\n3. Ensure credentials have permissions to use Bedrock models and S3"
                }
            else:
                return {
                    "error": f"AWS Bedrock error: {error_code}",
                    "note": error_message,
                    "suggestion": f"Check AWS Bedrock console and verify model access for region {region}"
                }
        except Exception as e:
            logger.error(f"Error in video generation: {str(e)}", exc_info=True)
            return {"error": f"Video generation failed: {str(e)}"}


class ImageToVideoTool(Tool):
    """Tool for generating videos from images using Amazon Bedrock Nova Reel."""
    
    def __init__(self):
        super().__init__(
            name="image_to_video",
            description="This tool generates a video from an image using Amazon Bedrock Nova Reel model. The image serves as the starting keyframe for video generation.",
            instructions="Use this tool when users request video generation from an image. If the user did not specify video length, ask first (e.g. 'How long should the video be? I can do 6–120 seconds.') then call with the duration parameter. Provide the image file path and a text_prompt that describes the desired motion or scene. Be specific so the video matches the user's intent. Video generation is asynchronous and may take 1-2 minutes."
        )
    
    def get_parameters_schema(self) -> List[Dict[str, Any]]:
        return [
            {
                "name": "image_path",
                "description": "Path or file name of the image to use as the starting keyframe. Use the exact file_name returned by the previous text_to_image (or image generation) tool result, e.g. generated_image_12345.png. Images are stored under generated_images/.",
                "type": "string",
                "required": True
            },
            {
                "name": "text_prompt",
                "description": "Describe the motion or animation you want (caption-style). E.g. 'Camera slowly zooms in. Gentle movement, cinematic.' or 'Person walks forward, natural lighting, smooth motion.' Avoid negation (no, not, without). Helps the model produce context-aware motion.",
                "type": "string",
                "required": False
            },
            {
                "name": "duration",
                "description": "Video length in seconds (6–120). If the user did not specify, ask them before calling (e.g. 'How long? 6–120 seconds.'). Default: 6.",
                "type": "string",
                "required": False
            },
            {
                "name": "resolution",
                "description": "Video resolution. Use '720p' or '1280x720' for clearer text/graphics. Options: '720p' (default), '1280x720', or '960x540'.",
                "type": "string",
                "required": False
            }
        ]
    
    def execute(self, parameters: Dict[str, Any]) -> Dict[str, Any]:
        image_path = parameters.get('image_path', '')
        text_prompt = parameters.get('text_prompt', '')
        duration = parameters.get('duration', '6')
        # Normalize duration: accept int (15) or str ("6s", "15")
        if isinstance(duration, int):
            duration = str(duration)
        elif not isinstance(duration, str):
            duration = '6'
        resolution = parameters.get('resolution', '720p') or '720p'
        
        if not image_path:
            return {"error": "image_path parameter is required"}
        
        try:
            from botocore.exceptions import ClientError
            from .aws_utils import create_boto3_client, get_aws_region
            import time
            import uuid
            import base64
            
            region = get_aws_region()
            bedrock_runtime = create_boto3_client('bedrock-runtime', region_name=region)
            s3_client = create_boto3_client('s3', region_name=region)
            
            # Resolve image path
            resolved_path = self._resolve_file_path(image_path)
            
            # Read and encode image
            if not default_storage.exists(resolved_path):
                hint = " Use the exact file_name from the previous text_to_image or image generation tool result (e.g. generated_image_12345.png)."
                try:
                    _, filenames = default_storage.listdir("generated_images")
                    if filenames:
                        hint = f" Available images (use one as image_path): {', '.join(filenames[:10])}."
                except Exception:
                    pass
                return {"error": f"Image file not found: {image_path}.{hint}"}
            
            file_obj = default_storage.open(resolved_path, 'rb')
            image_data = file_obj.read()
            file_obj.close()
            
            # Convert image to base64
            image_b64 = base64.b64encode(image_data).decode('utf-8')
            
            # Get S3 bucket for video output
            s3_bucket = os.getenv('BEDROCK_VIDEO_OUTPUT_BUCKET', None)
            if not s3_bucket:
                s3_bucket = getattr(settings, 'BEDROCK_VIDEO_OUTPUT_BUCKET', None)
                if not s3_bucket:
                    try:
                        sts_client = create_boto3_client('sts', region_name=region)
                        account_id = sts_client.get_caller_identity().get('Account')
                        s3_bucket = f"bedrock-video-output-{account_id}"
                    except Exception:
                        return {
                            "error": "S3 bucket not configured",
                            "note": "Set BEDROCK_VIDEO_OUTPUT_BUCKET environment variable or configure in settings",
                            "suggestion": "Create an S3 bucket and set BEDROCK_VIDEO_OUTPUT_BUCKET environment variable"
                        }
            
            # Generate unique S3 prefix. Bedrock expects s3Uri to be a bucket or directory (prefix),
            # not a full object key; it writes output as output.mp4 under that prefix.
            video_id = str(uuid.uuid4())
            s3_prefix = f"videos/{video_id}/"
            s3_uri = f"s3://{s3_bucket}/{s3_prefix}"
            s3_key = f"videos/{video_id}/output.mp4"
            
            # Parse duration
            duration_seconds = 6
            if duration.endswith('s'):
                try:
                    duration_seconds = int(duration[:-1])
                except Exception:
                    duration_seconds = 6
            else:
                try:
                    duration_seconds = int(duration)
                except Exception:
                    duration_seconds = 6
            
            # Clamp duration (6-120 seconds)
            duration_seconds = max(6, min(120, duration_seconds))
            
            # Parse resolution (prefer 720p for text/graphics clarity)
            resolution = (resolution or '720p').strip() or '720p'
            if resolution == '1280x720':
                dimension = "1280x720"
            elif resolution == '960x540' or resolution == '540p':
                dimension = "960x540"
            else:
                dimension = "1280x720"
            
            # Motion prompt (enhance for text/graph clarity when relevant)
            if text_prompt:
                motion_prompt = _enhance_video_prompt(text_prompt, min_length=30)
            else:
                motion_prompt = "Subtle cinematic motion, high quality, natural movement. Sharp details, clear and crisp."
            
            invocation_arn = None
            model_used = None
            
            # Try Luma Ray 2 first (better for text/graphics and detail; supports image-to-video via keyframes)
            media_type = "image/png" if (resolved_path.lower().endswith('.png') or '.png' in resolved_path.lower()) else "image/jpeg"
            luma_duration = "9s" if duration_seconds >= 9 else "5s"
            luma_input = {
                "prompt": motion_prompt[:5000],
                "keyframes": {
                    "frame0": {
                        "type": "image",
                        "source": {
                            "type": "base64",
                            "media_type": media_type,
                            "data": image_b64
                        }
                    }
                },
                "resolution": "720p",
                "duration": luma_duration,
                "aspect_ratio": "16:9",
                "loop": False
            }
            try:
                response = bedrock_runtime.start_async_invoke(
                    modelId="luma.ray-v2:0",
                    modelInput=luma_input,
                    outputDataConfig={"s3OutputDataConfig": {"s3Uri": s3_uri}}
                )
                invocation_arn = response.get('invocationArn')
                model_used = "luma.ray-v2:0"
                logger.info(f"Started image-to-video with Luma Ray 2, invocation ARN: {invocation_arn}")
            except ClientError as e:
                error_code = e.response.get('Error', {}).get('Code', '')
                logger.info(f"Luma Ray 2 image-to-video not available ({error_code}), falling back to Nova Reel")
                invocation_arn = None
            
            # Fallback to Nova Reel if Luma failed or unavailable
            if not invocation_arn:
                model_input = {
                    "taskType": "IMAGE_VIDEO",
                    "imageToVideoParams": {
                        "image": image_b64,
                        "text": motion_prompt[:512]
                    },
                    "videoGenerationConfig": {
                        "durationSeconds": duration_seconds,
                        "fps": 24,
                        "dimension": dimension,
                        "seed": 0
                    }
                }
                try:
                    response = bedrock_runtime.start_async_invoke(
                        modelId="amazon.nova-reel-v1:1",
                        modelInput=model_input,
                        outputDataConfig={"s3OutputDataConfig": {"s3Uri": s3_uri}}
                    )
                    invocation_arn = response.get('invocationArn')
                    model_used = "amazon.nova-reel-v1:1"
                    logger.info(f"Started image-to-video with Nova Reel, invocation ARN: {invocation_arn}")
                except ClientError as e:
                    error_code = e.response.get('Error', {}).get('Code', '')
                    error_message = e.response.get('Error', {}).get('Message', str(e))
                    if error_code == 'ValidationException':
                        return {
                            "error": f"Video model not available in region {region}",
                            "note": "Enable Luma Ray 2 or Amazon Nova Reel in AWS Bedrock console.",
                            "suggestion": "1. Go to AWS Bedrock console\n2. Enable video generation models\n3. Ensure they are available in your region"
                        }
                    elif error_code == 'UnrecognizedClientException':
                        return {
                            "error": f"AWS Bedrock authentication error: {error_code}",
                            "note": "AWS credentials are not configured properly.",
                            "suggestion": "1. Check AWS credentials\n2. Verify AWS region"
                        }
                    return {"error": f"AWS Bedrock error: {error_code}", "note": error_message}
            
            # Poll for completion (timeout from settings, default 10 min)
            max_wait_time = getattr(settings, 'BEDROCK_VIDEO_GENERATION_TIMEOUT', 600)
            poll_interval = 5  # Check every 5 seconds
            start_time = time.time()
            
            while time.time() - start_time < max_wait_time:
                try:
                    status_response = bedrock_runtime.get_async_invoke(
                        invocationArn=invocation_arn
                    )
                    
                    status = status_response.get('status')
                    
                    if status == 'Completed':
                        # Download video from S3; if output.mp4 not at expected key, list prefix and use .mp4 or first object
                        video_data = None
                        actual_key = None
                        try:
                            video_obj = s3_client.get_object(Bucket=s3_bucket, Key=s3_key)
                            video_data = video_obj['Body'].read()
                            actual_key = s3_key
                        except ClientError as e:
                            if e.response.get('Error', {}).get('Code') != 'NoSuchKey':
                                raise
                            list_resp = s3_client.list_objects_v2(Bucket=s3_bucket, Prefix=s3_prefix, MaxKeys=50)
                            contents = list_resp.get('Contents', [])
                            mp4_key = None
                            fallback_key = None
                            for obj in contents:
                                key = obj.get('Key', '').strip()
                                if not key or key.endswith('/'):
                                    continue
                                if key.endswith('.mp4'):
                                    mp4_key = key
                                    break
                                if fallback_key is None:
                                    fallback_key = key
                            actual_key = mp4_key or fallback_key
                            if actual_key:
                                video_obj = s3_client.get_object(Bucket=s3_bucket, Key=actual_key)
                                video_data = video_obj['Body'].read()
                        if video_data and actual_key:
                            file_name = f"generated_video_{video_id}.mp4"
                            saved_path = default_storage.save(f"generated_videos/{file_name}", ContentFile(video_data))
                            try:
                                s3_client.delete_object(Bucket=s3_bucket, Key=actual_key)
                            except Exception as cleanup_error:
                                logger.warning(f"Failed to clean up S3 file: {str(cleanup_error)}")
                            return {
                                "video_url": default_storage.url(saved_path),
                                "file_name": file_name,
                                "success": True,
                                "model_used": model_used or "amazon.nova-reel-v1:1",
                                "invocation_arn": invocation_arn
                            }
                        return {
                            "error": "Video generation completed but download failed",
                            "note": f"Video is available at S3: {s3_uri}. No object found under prefix.",
                            "invocation_arn": invocation_arn,
                            "s3_uri": s3_uri,
                            "download_error": "NoSuchKey and no objects listed under prefix"
                        }
                    
                    elif status == 'Failed':
                        failure_reason = status_response.get('failureReason', 'Unknown error')
                        return {
                            "error": "Video generation failed",
                            "note": failure_reason,
                            "invocation_arn": invocation_arn
                        }
                    
                    # Still in progress, wait and check again
                    time.sleep(poll_interval)
                    
                except ClientError as e:
                    error_code = e.response.get('Error', {}).get('Code', '')
                    if error_code == 'ResourceNotFoundException':
                        return {
                            "error": "Video generation job not found",
                            "invocation_arn": invocation_arn
                        }
                    raise
            
            # Timeout
            return {
                "error": "Video generation timeout",
                "note": f"Video generation is still in progress after {max_wait_time} seconds",
                "invocation_arn": invocation_arn,
                "s3_uri": s3_uri,
                "suggestion": f"Check status later using invocation ARN: {invocation_arn}"
            }
            
        except Exception as e:
            logger.error(f"Error in image-to-video generation: {str(e)}", exc_info=True)
            return {"error": f"Image-to-video generation failed: {str(e)}"}
    
    def _resolve_file_path(self, file_path: str) -> str:
        """Resolve file path from various sources (ConversationFile, TrainingData, or direct path)."""
        # Check if it's a file ID (UUID format)
        import re
        uuid_pattern = re.compile(r'^[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}$', re.I)
        
        if uuid_pattern.match(file_path):
            # Try to find in ConversationFile
            try:
                from .models import ConversationFile
                conv_file = ConversationFile.objects.get(file_id=file_path)
                return conv_file.file_path.name if conv_file.file_path else conv_file.file_path
            except ConversationFile.DoesNotExist:
                pass
            
            # Try TrainingData
            try:
                from .models import TrainingData
                training_file = TrainingData.objects.get(file_id=file_path)
                if training_file.file_path:
                    return training_file.file_path.name if hasattr(training_file.file_path, 'name') else training_file.file_path
            except TrainingData.DoesNotExist:
                pass
        
        # Normalize URL-style paths so they are relative to MEDIA_ROOT (default_storage base)
        # e.g. /media/conversation_files/3/foo.png -> conversation_files/3/foo.png
        if file_path and (file_path.startswith('/media/') or file_path.startswith('media/')):
            file_path = file_path.lstrip('/').replace('media/', '', 1) if file_path.startswith('/') else file_path.replace('media/', '', 1)
        elif file_path and file_path.startswith('/'):
            # Path like /conversation_files/... (no /media prefix) - strip leading slash only
            file_path = file_path.lstrip('/')
        
        # If path is a bare filename (no directory), try generated_images/ then conversation uploads
        if file_path and '/' not in file_path:
            candidate = f"generated_images/{file_path}"
            if default_storage.exists(candidate):
                return candidate
            # User-uploaded file (e.g. aws-certification-path-and-roadmap.jpg) may be in ConversationFile
            try:
                from .models import ConversationFile
                conv_file = ConversationFile.objects.filter(file_name=file_path).first()
                if conv_file and conv_file.file_path:
                    return conv_file.file_path.name if hasattr(conv_file.file_path, 'name') else str(conv_file.file_path)
            except Exception:
                pass
        
        return file_path


class OCRTool(Tool):
    """Tool for extracting text from images and PDF files."""
    
    def __init__(self):
        super().__init__(
            name="ocr",
            description="Extracts text from images, PDFs, Word documents, Excel files, CSV files, and other document formats. Supports OCR for images and PDFs, and direct reading for text-based formats.",
            instructions="Use this tool when users provide files that need text extraction. Supports images (OCR), PDFs (text extraction), Word documents (.docx), Excel files (.xlsx, .xls), CSV files (.csv, .tsv), and text files. CSV files are automatically parsed and formatted."
        )
    
    def get_parameters_schema(self) -> List[Dict[str, Any]]:
        return [
            {
                "name": "file_path",
                "description": "Path to the image or PDF file for text extraction.",
                "type": "string",
                "required": True
            }
        ]
    
    def execute(self, parameters: Dict[str, Any]) -> Dict[str, Any]:
        file_path = parameters.get('file_path', '')
        if not file_path:
            return {"error": "file_path parameter is required"}
        
        try:
            # Resolve file path first
            resolved_path = self._resolve_file_path(file_path)
            
            # Check file extension to determine type
            file_ext = os.path.splitext(file_path)[1].lower() if file_path else ''
            is_pdf = file_ext == '.pdf'
            
            # Image file extensions
            image_extensions = ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp', '.svg']
            is_image = file_ext in image_extensions
            
            # Office document extensions
            word_extensions = ['.docx', '.doc']
            excel_extensions = ['.xlsx', '.xls']
            is_word = file_ext in word_extensions
            is_excel = file_ext in excel_extensions
            
            # Text file extensions - read directly without OCR
            text_extensions = ['.txt', '.md', '.markdown', '.yml', '.yaml', '.json', '.xml', '.csv', '.tsv', '.log', 
                             '.py', '.js', '.jsx', '.ts', '.tsx', '.html', '.htm', '.css', '.scss', '.sh', '.bash',
                             '.bat', '.conf', '.config', '.ini', '.properties', '.env', '.gitignore', '.dockerfile',
                             '.java', '.cpp', '.c', '.h', '.hpp', '.cs', '.go', '.rb', '.php', '.swift', '.kt',
                             '.scala', '.r', '.m', '.pl', '.pm', '.lua', '.sql', '.vue', '.svelte']
            
            is_text_file = file_ext in text_extensions
            
            # For PDF files, try PDF extraction first
            if is_pdf:
                return self._extract_from_pdf(file_path)
            
            # For Word documents, extract text
            if is_word:
                return self._extract_from_word(resolved_path, file_path)
            
            # For Excel documents, extract text
            if is_excel:
                return self._extract_from_excel(resolved_path, file_path)
            
            # For image files, use OCR
            if is_image:
                return self._extract_from_image(file_path)
            
            # For text files, read directly (no OCR needed)
            if is_text_file:
                return self._read_text_file(resolved_path, file_path)
            
            # For unknown file types, try multiple approaches:
            # 1. Try as image (might be image with wrong extension)
            # 2. Try as text file
            # 3. Return helpful error
            
            # First try as image
            result = self._extract_from_image(file_path)
            if result.get('success'):
                return result
            
            # If image extraction failed, try reading as text
            text_result = self._read_text_file(resolved_path, file_path)
            if text_result.get('success') or not text_result.get('error'):
                return text_result
            
            # If both failed, return helpful error
            return {
                "text": f"Could not extract text from file: {file_path}",
                "error": "File type not supported",
                "note": f"File extension '{file_ext}' is not recognized. Tried as image and text file.",
                "suggestion": f"1. Ensure the file is a valid image, PDF, Word, Excel, or text file\n2. Check if file exists: {resolved_path}\n3. For binary files, consider using appropriate tools",
                "file_type": "unknown",
                "file_extension": file_ext,
                "resolved_path": resolved_path
            }
            
        except Exception as e:
            logger.error(f"Error in OCR: {str(e)}", exc_info=True)
            return {"error": f"OCR failed: {str(e)}"}
    
    def _extract_from_word(self, resolved_path: str, original_path: str) -> Dict[str, Any]:
        """Extract text from Word documents (.docx, .doc)."""
        try:
            if not default_storage.exists(resolved_path):
                return {
                    "text": f"Word document not found: {original_path}",
                    "error": "File does not exist",
                    "resolved_path": resolved_path
                }
            
            # Try .docx first (modern format)
            if original_path.lower().endswith('.docx'):
                try:
                    from docx import Document
                    
                    file_obj = default_storage.open(resolved_path, 'rb')
                    doc = Document(file_obj)
                    file_obj.close()
                    
                    # Extract text from paragraphs
                    paragraphs = []
                    for para in doc.paragraphs:
                        if para.text.strip():
                            paragraphs.append(para.text)
                    
                    # Extract text from tables
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = []
                            for cell in row.cells:
                                if cell.text.strip():
                                    row_text.append(cell.text.strip())
                            if row_text:
                                paragraphs.append(' | '.join(row_text))
                    
                    extracted_text = '\n'.join(paragraphs)
                    
                    if extracted_text.strip():
                        return {
                            "text": extracted_text,
                            "success": True,
                            "method": "python-docx",
                            "file_type": "word",
                            "format": "docx"
                        }
                    else:
                        return {
                            "text": "Word document appears to be empty or contains no extractable text.",
                            "error": "Empty document",
                            "file_type": "word"
                        }
                        
                except ImportError:
                    return {
                        "text": f"python-docx library not installed. Cannot read .docx files.",
                        "error": "Library missing",
                        "suggestion": "Install python-docx: pip install python-docx",
                        "file_type": "word"
                    }
                except Exception as e:
                    logger.error(f"Error reading DOCX: {str(e)}", exc_info=True)
                    return {
                        "text": f"Error reading Word document '{original_path}': {str(e)}",
                        "error": "Read error",
                        "file_type": "word"
                    }
            
            # For .doc files (older format), try using textract or antiword
            elif original_path.lower().endswith('.doc'):
                # Try to read as binary and check if it's actually a .docx
                try:
                    file_obj = default_storage.open(resolved_path, 'rb')
                    header = file_obj.read(4)
                    file_obj.close()
                    
                    # Check if it's actually a .docx file (ZIP format)
                    if header == b'PK\x03\x04':
                        # It's actually a .docx file with wrong extension
                        return self._extract_from_word(resolved_path, original_path.replace('.doc', '.docx'))
                except Exception:
                    pass
                
                return {
                    "text": f"Legacy .doc format not directly supported. File: {original_path}",
                    "error": "Unsupported format",
                    "suggestion": "1. Convert .doc to .docx format\n2. Or install textract library: pip install textract\n3. Or use antiword tool for .doc files",
                    "file_type": "word",
                    "format": "doc"
                }
            
        except Exception as e:
            logger.error(f"Error in _extract_from_word: {str(e)}", exc_info=True)
            return {
                "text": f"Error processing Word document '{original_path}': {str(e)}",
                "error": "Processing error"
            }
    
    def _extract_from_excel(self, resolved_path: str, original_path: str) -> Dict[str, Any]:
        """Extract text from Excel documents (.xlsx, .xls)."""
        try:
            if not default_storage.exists(resolved_path):
                return {
                    "text": f"Excel document not found: {original_path}",
                    "error": "File does not exist",
                    "resolved_path": resolved_path
                }
            
            # Try .xlsx first (modern format)
            if original_path.lower().endswith('.xlsx'):
                try:
                    import openpyxl
                    
                    file_obj = default_storage.open(resolved_path, 'rb')
                    workbook = openpyxl.load_workbook(file_obj, data_only=True)
                    file_obj.close()
                    
                    # Extract text from all sheets
                    all_sheets_text = []
                    
                    for sheet_name in workbook.sheetnames:
                        sheet = workbook[sheet_name]
                        sheet_data = []
                        sheet_data.append(f"\n=== Sheet: {sheet_name} ===\n")
                        
                        # Get all cells with data
                        for row in sheet.iter_rows(values_only=True):
                            row_data = []
                            for cell_value in row:
                                if cell_value is not None:
                                    # Convert cell value to string
                                    cell_str = str(cell_value).strip()
                                    if cell_str:
                                        row_data.append(cell_str)
                            
                            if row_data:
                                sheet_data.append(' | '.join(row_data))
                        
                        if len(sheet_data) > 1:  # More than just the header
                            all_sheets_text.append('\n'.join(sheet_data))
                    
                    extracted_text = '\n\n'.join(all_sheets_text)
                    
                    if extracted_text.strip():
                        return {
                            "text": extracted_text,
                            "success": True,
                            "method": "openpyxl",
                            "file_type": "excel",
                            "format": "xlsx",
                            "sheets": workbook.sheetnames
                        }
                    else:
                        return {
                            "text": "Excel document appears to be empty or contains no data.",
                            "error": "Empty document",
                            "file_type": "excel"
                        }
                        
                except ImportError:
                    return {
                        "text": f"openpyxl library not installed. Cannot read .xlsx files.",
                        "error": "Library missing",
                        "suggestion": "Install openpyxl: pip install openpyxl",
                        "file_type": "excel"
                    }
                except Exception as e:
                    logger.error(f"Error reading XLSX: {str(e)}", exc_info=True)
                    return {
                        "text": f"Error reading Excel document '{original_path}': {str(e)}",
                        "error": "Read error",
                        "file_type": "excel"
                    }
            
            # For .xls files (older format), try using xlrd
            elif original_path.lower().endswith('.xls'):
                try:
                    import xlrd
                    
                    file_obj = default_storage.open(resolved_path, 'rb')
                    file_bytes = file_obj.read()
                    file_obj.close()
                    
                    workbook = xlrd.open_workbook(file_contents=file_bytes)
                    
                    # Extract text from all sheets
                    all_sheets_text = []
                    
                    for sheet_name in workbook.sheet_names():
                        sheet = workbook.sheet_by_name(sheet_name)
                        sheet_data = []
                        sheet_data.append(f"\n=== Sheet: {sheet_name} ===\n")
                        
                        # Get all cells with data
                        for row_idx in range(sheet.nrows):
                            row_data = []
                            for col_idx in range(sheet.ncols):
                                cell = sheet.cell(row_idx, col_idx)
                                cell_value = cell.value
                                if cell_value:
                                    cell_str = str(cell_value).strip()
                                    if cell_str:
                                        row_data.append(cell_str)
                            
                            if row_data:
                                sheet_data.append(' | '.join(row_data))
                        
                        if len(sheet_data) > 1:  # More than just the header
                            all_sheets_text.append('\n'.join(sheet_data))
                    
                    extracted_text = '\n\n'.join(all_sheets_text)
                    
                    if extracted_text.strip():
                        return {
                            "text": extracted_text,
                            "success": True,
                            "method": "xlrd",
                            "file_type": "excel",
                            "format": "xls",
                            "sheets": workbook.sheet_names()
                        }
                    else:
                        return {
                            "text": "Excel document appears to be empty or contains no data.",
                            "error": "Empty document",
                            "file_type": "excel"
                        }
                        
                except ImportError:
                    return {
                        "text": f"xlrd library not installed. Cannot read .xls files.",
                        "error": "Library missing",
                        "suggestion": "Install xlrd: pip install xlrd",
                        "file_type": "excel"
                    }
                except Exception as e:
                    logger.error(f"Error reading XLS: {str(e)}", exc_info=True)
                    return {
                        "text": f"Error reading Excel document '{original_path}': {str(e)}",
                        "error": "Read error",
                        "file_type": "excel"
                    }
            
        except Exception as e:
            logger.error(f"Error in _extract_from_excel: {str(e)}", exc_info=True)
            return {
                "text": f"Error processing Excel document '{original_path}': {str(e)}",
                "error": "Processing error"
            }
    
    def _read_text_file(self, resolved_path: str, original_path: str) -> Dict[str, Any]:
        """Read text content from a text file."""
        try:
            # Check file extension for special handling
            file_ext = os.path.splitext(original_path)[1].lower() if original_path else ''
            is_csv = file_ext in ['.csv', '.tsv']
            
            # Check if file exists in storage
            if default_storage.exists(resolved_path):
                try:
                    # For CSV/TSV files, parse and format properly
                    if is_csv:
                        import csv
                        # Read as binary first to handle encoding
                        file_obj = default_storage.open(resolved_path, 'rb')
                        content_bytes = file_obj.read()
                        file_obj.close()
                        
                        # Try UTF-8 first, then fallback to latin-1
                        try:
                            content_str = content_bytes.decode('utf-8')
                        except UnicodeDecodeError:
                            content_str = content_bytes.decode('latin-1', errors='ignore')
                        
                        # Parse CSV and format
                        csv_reader = csv.reader(content_str.splitlines())
                        rows = []
                        for row in csv_reader:
                            if row:  # Skip empty rows
                                rows.append(' | '.join(str(cell).strip() for cell in row))
                        
                        formatted_content = '\n'.join(rows) if rows else content_str
                        
                        return {
                            "text": formatted_content,
                            "success": True,
                            "method": "csv_parser",
                            "file_type": "csv" if file_ext == '.csv' else "tsv",
                            "encoding": "utf-8",
                            "rows": len(rows)
                        }
                    
                    # For other text files, read directly
                    # default_storage.open() doesn't support encoding parameter - read as binary and decode manually
                    file_obj = default_storage.open(resolved_path, 'rb')
                    content_bytes = file_obj.read()
                    file_obj.close()
                    # Try UTF-8 first
                    content = content_bytes.decode('utf-8')
                    
                    return {
                        "text": content,
                        "success": True,
                        "method": "direct_read",
                        "file_type": "text",
                        "encoding": "utf-8"
                    }
                except UnicodeDecodeError:
                    # Try with different encodings - read as binary and decode manually
                    file_obj = default_storage.open(resolved_path, 'rb')
                    content_bytes = file_obj.read()
                    file_obj.close()
                    for encoding in ['latin-1', 'iso-8859-1', 'cp1252']:
                        try:
                            content = content_bytes.decode(encoding, errors='ignore')
                            return {
                                "text": content,
                                "success": True,
                                "method": "direct_read",
                                "file_type": "text",
                                "encoding": encoding
                            }
                        except (UnicodeDecodeError, Exception):
                            continue
                    
                    return {
                        "text": f"Could not decode file '{original_path}' as text. File may be binary.",
                        "error": "Encoding error",
                        "suggestion": "File appears to be binary or uses an unsupported encoding."
                    }
                except Exception as e:
                    logger.error(f"Error reading text file: {str(e)}", exc_info=True)
                    return {
                        "text": f"Error reading file '{original_path}': {str(e)}",
                        "error": "Read error"
                    }
            else:
                # File not found in storage - provide detailed suggestions
                suggestions = []
                suggestions.append(f"1. Verify file path is correct: '{original_path}'")
                suggestions.append(f"2. Resolved path attempted: '{resolved_path}'")
                suggestions.append("3. Ensure file was uploaded via upload_file API or training data upload")
                
                # Try to find similar files
                try:
                    from .models import ConversationFile, TrainingData
                    base_name = os.path.basename(original_path)
                    base_name_no_ext = os.path.splitext(base_name)[0]
                    file_ext = os.path.splitext(base_name)[1].lower()
                    similar_files = []
                    
                    # Search ConversationFile
                    conv_files = ConversationFile.objects.filter(file_name__icontains=base_name)[:3]
                    for cf in conv_files:
                        similar_files.append(f"   - {cf.file_name} (ID: {cf.file_id}, Agent: {cf.agent.id})")
                    
                    # Search TrainingData
                    train_files = TrainingData.objects.filter(
                        data_type='file',
                        file_path__icontains=base_name_no_ext if base_name_no_ext else base_name
                    )[:3]
                    for tf in train_files:
                        if tf.file_path:
                            path_name = tf.file_path.name if hasattr(tf.file_path, 'name') else str(tf.file_path)
                            similar_files.append(f"   - {path_name} (Agent: {tf.agent.id})")
                    
                    if similar_files:
                        suggestions.append("\n4. Similar files found:")
                        suggestions.extend(similar_files)
                except Exception as e:
                    logger.debug(f"Error searching for similar files: {str(e)}")
                
                return {
                    "text": f"File not found: {original_path}",
                    "error": "File does not exist",
                    "resolved_path": resolved_path,
                    "suggestion": "\n".join(suggestions)
                }
        except Exception as e:
            logger.error(f"Error in _read_text_file: {str(e)}", exc_info=True)
            return {
                "text": f"Error processing text file '{original_path}': {str(e)}",
                "error": "Processing error"
            }
    
    def _resolve_file_path(self, file_path: str, agent_id: Optional[int] = None) -> str:
        """Resolve file path by looking up in ConversationFile and TrainingData models.
        
        Args:
            file_path: Original file path or filename to resolve
            agent_id: Optional agent ID to scope the search (for training data)
        
        Returns:
            Resolved file path that can be used with default_storage.
        """
        import os
        original_file_path = file_path
        base_name = os.path.basename(file_path) if file_path else file_path
        
        # First, try to resolve file_id if it looks like one
        file_id_candidate = file_path
        if '_' in file_path and '.' in file_path:
            parts = file_path.rsplit('_', 1)
            if len(parts) == 2:
                file_id_candidate = parts[0]
        
        # Try to look up by file_id or file_name
        try:
            from .models import ConversationFile, TrainingData
            
            conversation_file = None
            training_data_file = None
            
            # Try exact file_id match first
            conversation_file = ConversationFile.objects.filter(file_id=file_path).first()
            if not conversation_file and file_id_candidate != file_path:
                conversation_file = ConversationFile.objects.filter(file_id=file_id_candidate).first()
            
            # Try matching by file_name (exact match, then contains)
            if not conversation_file:
                conversation_file = ConversationFile.objects.filter(file_name=file_path).first()
            if not conversation_file:
                conversation_file = ConversationFile.objects.filter(file_name=base_name).first()
            if not conversation_file:
                conversation_file = ConversationFile.objects.filter(file_name__icontains=base_name).first()
            
            if conversation_file:
                if conversation_file.file_path:
                    if hasattr(conversation_file.file_path, 'name'):
                        file_path = conversation_file.file_path.name
                    else:
                        file_path = str(conversation_file.file_path)
                    logger.info(f"Resolved file '{original_file_path}' to ConversationFile path: {file_path}")
                elif hasattr(conversation_file, 'file_name'):
                    file_path = f"conversation_files/{conversation_file.agent.id}/{conversation_file.file_name}"
                    logger.info(f"Constructed path from ConversationFile file_name: {file_path}")
            
            # Also check TrainingData files
            if not conversation_file:
                # Build query for TrainingData
                training_data_query = TrainingData.objects.filter(data_type='file')
                
                # If agent_id provided, scope to that agent (more efficient)
                if agent_id:
                    training_data_query = training_data_query.filter(agent_id=agent_id)
                
                # Try multiple search strategies
                # 1. Exact filename match in file_path (most reliable)
                training_data_file = training_data_query.filter(
                    file_path__icontains=base_name
                ).first()
                
                # 2. Check if filename is in content metadata
                if not training_data_file:
                    training_data_file = training_data_query.filter(
                        content__icontains=base_name
                    ).first()
                
                # 3. Try with full original path
                if not training_data_file:
                    training_data_file = training_data_query.filter(
                        file_path__icontains=original_file_path
                    ).first()
                
                # 4. Try partial match (in case of encoding/character issues)
                if not training_data_file and len(base_name) > 10:
                    # Try with first part of filename (before parentheses or first 20 chars)
                    name_part = base_name.split('(')[0].strip() if '(' in base_name else base_name[:20]
                    if name_part:
                        training_data_file = training_data_query.filter(
                            file_path__icontains=name_part
                        ).first()
                
                # 5. If still not found and no agent_id, try searching all agents
                if not training_data_file and agent_id is None:
                    logger.debug(f"File not found in scoped search, trying all training data for: {base_name}")
                    training_data_file = TrainingData.objects.filter(
                        data_type='file',
                        file_path__icontains=base_name
                    ).first()
                
                if training_data_file and training_data_file.file_path:
                    if hasattr(training_data_file.file_path, 'name'):
                        file_path = training_data_file.file_path.name
                    else:
                        file_path = str(training_data_file.file_path)
                    logger.info(f"Resolved file '{original_file_path}' to TrainingData path: {file_path} (agent_id: {training_data_file.agent.id})")
                elif not training_data_file:
                    logger.debug(f"Could not find file '{original_file_path}' (base: '{base_name}') in TrainingData")
        except Exception as e:
            logger.warning(f"Could not resolve file path: {str(e)}", exc_info=True)
        
        return file_path
    
    def _extract_from_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract text from PDF files using multiple methods."""
        # Resolve file path first (look up in ConversationFile/TrainingData)
        resolved_path = self._resolve_file_path(file_path)
        
        # Check if AWS Textract is configured (needed for error messages)
        aws_configured = False
        try:
            from .aws_utils import create_boto3_client
            create_boto3_client('textract')
            aws_configured = True
        except Exception:
            pass
        
        # Method 1: Try PyPDF2 (for text-based PDFs - fastest and most common)
        try:
            import PyPDF2
            
            if default_storage.exists(resolved_path):
                file_obj = default_storage.open(resolved_path, 'rb')
                pdf_reader = PyPDF2.PdfReader(file_obj)
                
                text_pages = []
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_pages.append(page_text)
                
                file_obj.close()
                
                if text_pages:
                    extracted_text = '\n\n'.join(text_pages)
                    return {
                        "text": extracted_text,
                        "success": True,
                        "method": "pypdf2",
                        "file_type": "pdf",
                        "pages": len(text_pages)
                    }
        except ImportError:
            logger.debug("PyPDF2 not installed")
        except Exception as e:
            logger.debug(f"PyPDF2 extraction failed: {str(e)}")
        
        # Method 2: Try pdfplumber (better for complex PDFs)
        try:
            import pdfplumber
            
            if default_storage.exists(resolved_path):
                file_obj = default_storage.open(resolved_path, 'rb')
                
                text_pages = []
                with pdfplumber.open(file_obj) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text_pages.append(page_text)
                
                file_obj.close()
                
                if text_pages:
                    extracted_text = '\n\n'.join(text_pages)
                    return {
                        "text": extracted_text,
                        "success": True,
                        "method": "pdfplumber",
                        "file_type": "pdf",
                        "pages": len(text_pages)
                    }
        except ImportError:
            logger.debug("pdfplumber not installed")
        except Exception as e:
            logger.debug(f"pdfplumber extraction failed: {str(e)}")
        
        # Method 3: For scanned PDFs, convert to images and use OCR
        try:
            from pdf2image import convert_from_bytes
            import pytesseract
            
            if default_storage.exists(resolved_path):
                file_obj = default_storage.open(resolved_path, 'rb')
                file_bytes = file_obj.read()
                file_obj.close()
                
                # Convert PDF pages to images
                images = convert_from_bytes(file_bytes)
                
                text_pages = []
                for image in images:
                    page_text = pytesseract.image_to_string(image)
                    if page_text.strip():
                        text_pages.append(page_text)
                
                if text_pages:
                    extracted_text = '\n\n'.join(text_pages)
                    return {
                        "text": extracted_text,
                        "success": True,
                        "method": "pdf2image-tesseract",
                        "file_type": "pdf",
                        "pages": len(text_pages)
                    }
        except ImportError:
            logger.debug("pdf2image or pytesseract not installed")
        except Exception as e:
            logger.debug(f"PDF to image OCR failed: {str(e)}")
        
        # Method 4: Try AWS Textract (last resort - best for scanned PDFs and images, but costs money)
        try:
            from .aws_utils import create_boto3_client
            textract_client = create_boto3_client('textract')
            
            if default_storage.exists(resolved_path):
                file_obj = default_storage.open(resolved_path, 'rb')
                file_bytes = file_obj.read()
                file_obj.close()
                
                # Textract can handle PDFs directly
                response = textract_client.detect_document_text(
                    Document={'Bytes': file_bytes}
                )
                
                # Extract text blocks
                text_blocks = []
                for block in response.get('Blocks', []):
                    if block.get('BlockType') == 'LINE':
                        text_blocks.append(block.get('Text', ''))
                
                extracted_text = '\n'.join(text_blocks)
                
                if extracted_text.strip():
                    return {
                        "text": extracted_text,
                        "success": True,
                        "method": "aws-textract",
                        "file_type": "pdf"
                    }
        except Exception as e:
            logger.debug(f"AWS Textract for PDF not available: {str(e)}")
        
        # Check if file exists before saying libraries aren't configured
        file_exists = default_storage.exists(resolved_path) if resolved_path else False
        
        # Check which libraries are available
        has_pypdf2 = False
        has_pdfplumber = False
        try:
            import PyPDF2
            has_pypdf2 = True
        except ImportError:
            pass
        
        try:
            import pdfplumber
            has_pdfplumber = True
        except ImportError:
            pass
        
        if not file_exists:
            # Try to provide helpful suggestions
            suggestions = [
                f"1. Verify file path is correct: '{file_path}'",
                f"2. Resolved path attempted: '{resolved_path}'",
                "3. Ensure file was uploaded via upload_file API or training data upload",
                "4. Check if file exists in ConversationFile or TrainingData models"
            ]
            
            # Try to find similar files
            try:
                from .models import TrainingData, ConversationFile
                import os
                base_name = os.path.basename(file_path)
                
                # Check for similar training data files
                similar_training = TrainingData.objects.filter(
                    data_type='file',
                    file_path__icontains=base_name[:20] if len(base_name) > 20 else base_name
                )[:5]
                
                if similar_training.exists():
                    suggestions.append("\n5. Similar training data files found:")
                    for td in similar_training:
                        if td.file_path:
                            path_name = td.file_path.name if hasattr(td.file_path, 'name') else str(td.file_path)
                            suggestions.append(f"   - Agent {td.agent.id}: {path_name}")
                
                # Check for similar conversation files
                similar_conv = ConversationFile.objects.filter(
                    file_name__icontains=base_name[:20] if len(base_name) > 20 else base_name
                )[:5]
                
                if similar_conv.exists():
                    suggestions.append("\n6. Similar conversation files found:")
                    for cf in similar_conv:
                        suggestions.append(f"   - Agent {cf.agent.id}: {cf.file_name} (ID: {cf.file_id})")
            except Exception as e:
                logger.debug(f"Could not search for similar files: {str(e)}")
            
            return {
                "text": f"PDF file not found: {file_path}",
                "error": "File does not exist in storage",
                "note": f"Verify the file exists. Available libraries: PyPDF2={has_pypdf2}, pdfplumber={has_pdfplumber}, AWS Textract={'Available' if aws_configured else 'Not configured'}",
                "suggestion": "\n".join(suggestions)
            }
        
        # File exists but extraction failed - provide library status
        missing_libs = []
        if not has_pypdf2:
            missing_libs.append("PyPDF2")
        if not has_pdfplumber:
            missing_libs.append("pdfplumber")
        
        error_msg = f"PDF text extraction failed for: {resolved_path}"
        if missing_libs:
            error_msg += f". Missing libraries: {', '.join(missing_libs)}"
        
        suggestion_parts = []
        if missing_libs:
            suggestion_parts.append(f"1. Install missing libraries: pip install {' '.join(missing_libs)}")
        suggestion_parts.append("2. For scanned PDFs, use AWS Textract (already configured) or install pdf2image + pytesseract")
        suggestion_parts.append("3. Check if PDF is corrupted or password-protected")
        
        return {
            "text": error_msg,
            "note": f"File exists but text extraction failed. Available: PyPDF2={has_pypdf2}, pdfplumber={has_pdfplumber}, AWS Textract={'Available' if aws_configured else 'Not configured'}",
            "suggestion": "\n".join(suggestion_parts)
        }
    
    def _extract_from_image(self, file_path: str) -> Dict[str, Any]:
        """Extract text from image files using OCR."""
        # Resolve file path first (look up in ConversationFile/TrainingData)
        resolved_path = self._resolve_file_path(file_path)
        
        # AWS Textract is disabled - using EasyOCR as primary method
        textract_error = "AWS Textract disabled - using EasyOCR instead"
        
        # Method 1: Try EasyOCR (Python-only, no system dependencies) - PRIMARY METHOD
        easyocr_error = None
        try:
            import easyocr
            
            if default_storage.exists(resolved_path):
                file_obj = default_storage.open(resolved_path, 'rb')
                file_bytes = file_obj.read()
                file_obj.close()
                
                # Initialize EasyOCR reader (downloads models on first use)
                reader = easyocr.Reader(['en'], gpu=False)  # Use CPU, no GPU needed
                
                # Read image from bytes
                import numpy as np
                from PIL import Image
                import io
                
                image = Image.open(io.BytesIO(file_bytes))
                image_array = np.array(image)
                
                # Extract text
                results = reader.readtext(image_array)
                text_lines = [result[1] for result in results]  # Extract text from results
                extracted_text = '\n'.join(text_lines)
                
                if extracted_text.strip():
                    return {
                        "text": extracted_text,
                        "success": True,
                        "method": "easyocr",
                        "file_type": "image"
                    }
        except ImportError:
            easyocr_error = "easyocr not installed (pip install easyocr)"
            logger.debug("EasyOCR not available")
        except Exception as e:
            error_msg = str(e)
            easyocr_error = f"EasyOCR error: {error_msg[:200]}"
            logger.debug(f"EasyOCR failed: {error_msg}")
        
        # Method 2: Try Tesseract OCR (open source, requires system installation)
        tesseract_error = None
        try:
            import pytesseract
            from PIL import Image
            
            if default_storage.exists(resolved_path):
                file_obj = default_storage.open(resolved_path, 'rb')
                image = Image.open(file_obj)
                text = pytesseract.image_to_string(image)
                file_obj.close()
                
                if text.strip():
                    return {
                        "text": text,
                        "success": True,
                        "method": "tesseract",
                        "file_type": "image"
                    }
        except ImportError:
            tesseract_error = "pytesseract not installed (pip install pytesseract)"
            logger.debug("Tesseract Python package not available")
        except Exception as e:
            error_msg = str(e)
            if 'tesseract' in error_msg.lower() or 'not found' in error_msg.lower():
                tesseract_error = "Tesseract OCR engine not installed. Install: sudo yum install tesseract (Amazon Linux) or sudo apt-get install tesseract-ocr (Ubuntu)"
            else:
                tesseract_error = f"Tesseract error: {error_msg[:200]}"
            logger.debug(f"Tesseract OCR failed: {error_msg}")
        
        # Final fallback - provide detailed diagnostics
        diagnostics = []
        diagnostics.append(f"File: {file_path}")
        diagnostics.append(f"Resolved path: {resolved_path}")
        diagnostics.append(f"File exists: {default_storage.exists(resolved_path) if resolved_path else 'Unknown'}")
        
        if textract_error:
            diagnostics.append(f"AWS Textract: {textract_error} (Disabled)")
        else:
            diagnostics.append("AWS Textract: Disabled - using EasyOCR instead")
        
        if easyocr_error:
            diagnostics.append(f"EasyOCR: {easyocr_error}")
        else:
            diagnostics.append("EasyOCR: Not attempted or unknown error")
        
        if tesseract_error:
            diagnostics.append(f"Tesseract OCR: {tesseract_error}")
        else:
            diagnostics.append("Tesseract OCR: Not attempted or unknown error")
        
        suggestion_parts = [
            "To enable OCR for images:",
            "",
            "Option 1: EasyOCR (Free, Python-only - Already Installed!)",
            "  - EasyOCR is installed and ready to use",
            "  - No system dependencies needed",
            "  - Models will download automatically on first use",
            "",
            "Option 2: AWS Textract (Best quality, requires AWS permissions)",
            "  - Ensure IAM role has 'textract:DetectDocumentText' permission",
            "  - AWS Textract uses existing AWS credentials",
            "",
            "Option 3: Tesseract OCR (Open source fallback)",
            "  - Install Python package: pip install pytesseract Pillow",
            "  - Install Tesseract engine:",
            "    * Amazon Linux: Requires compilation from source",
            "    * Ubuntu/Debian: sudo apt-get install tesseract-ocr",
            "    * macOS: brew install tesseract",
        ]
        
        return {
            "text": f"OCR service is not currently configured to extract text from the image you provided ({os.path.basename(file_path)}).",
            "error": "OCR service not configured",
            "note": "\n".join(diagnostics),
            "suggestion": "\n".join(suggestion_parts),
            "success": False
        }


class SummarizationTool(Tool):
    """Tool for summarizing documents."""
    
    def __init__(self):
        super().__init__(
            name="summarization",
            description="This tool is used to summarize files (e.g., articles, pdfs, docx, etc.).",
            instructions="Use this tool when users request summaries of documents or long text."
        )
    
    def get_parameters_schema(self) -> List[Dict[str, Any]]:
        return [
            {
                "name": "file_path",
                "description": "Path to the file to summarize.",
                "type": "string",
                "required": False
            },
            {
                "name": "text_content",
                "description": "Direct text content to summarize (alternative to file_path).",
                "type": "string",
                "required": False
            },
            {
                "name": "max_length",
                "description": "Maximum length of the summary.",
                "type": "integer",
                "required": False
            }
        ]
    
    def execute(self, parameters: Dict[str, Any]) -> Dict[str, Any]:
        file_path = parameters.get('file_path', '')
        max_length = parameters.get('max_length', 500)
        text_content = parameters.get('text_content', '')  # Allow direct text input
        
        if not file_path and not text_content:
            return {"error": "Either file_path or text_content parameter is required"}
        
        try:
            # Read file content if file_path provided
            content = text_content
            if file_path and not content:
                # Use OCR tool's comprehensive file reading (handles all file types)
                try:
                    ocr_tool = OCRTool()
                    result = ocr_tool.execute({"file_path": file_path})
                    
                    # Extract text from OCR tool result
                    if result.get('success'):
                        content = result.get('text', '')
                    elif result.get('text'):
                        # Even if not marked as success, use the text if available
                        content = result.get('text', '')
                    else:
                        # OCR tool couldn't read the file
                        error_msg = result.get('error', 'Unknown error')
                        note = result.get('note', '')
                        suggestion = result.get('suggestion', '')
                        return {
                            "error": f"Could not read file: {file_path}",
                            "details": error_msg,
                            "note": note,
                            "suggestion": suggestion
                        }
                except Exception as e:
                    logger.error(f"Error using OCR tool for file reading: {str(e)}", exc_info=True)
                    # Fallback to original method
                    content = self._read_file_content(file_path)
                    if not content or (content.startswith('[') and 'Error' in content):
                        return {"error": f"Could not read file: {file_path}", "details": content}
            
            if not content:
                return {"error": "No content to summarize"}
            
            # Use LLM for summarization via Bedrock
            try:
                from .aws_integration import get_bedrock_client
                from .models import AIModel
                
                client = get_bedrock_client()
                
                # Get a default Claude model for summarization
                # Try to find an available Claude model from the database
                default_model = "anthropic.claude-3-sonnet-20240229-v1:0"
                try:
                    claude_model = AIModel.objects.filter(
                        provider='bedrock',
                        model_id__icontains='claude',
                        is_available=True
                    ).first()
                    if claude_model:
                        default_model = claude_model.model_id
                except Exception:
                    pass  # Use default if database lookup fails
                
                # Truncate content if too long (keep first 8000 chars for context)
                content_to_summarize = content[:8000] if len(content) > 8000 else content
                
                prompt = f"""Please provide a concise summary of the following text. 
The summary should be approximately {max_length} characters or less.

Text to summarize:
{content_to_summarize}

Summary:"""
                
                result = client.invoke_agent(
                    agent_id="summarization_tool",  # Tool identifier, not a real agent
                    query=prompt,
                    model=default_model,
                    system_prompt="You are a helpful assistant that creates concise, accurate summaries.",
                    model_provider="bedrock"
                )
                
                summary = result.get('response', '').strip()
                
                if summary:
                    return {
                        "summary": summary,
                        "original_length": len(content),
                        "summary_length": len(summary),
                        "success": True,
                        "model_used": default_model
                    }
                else:
                    # Empty response - fallback to truncation
                    raise Exception("Empty response from LLM")
                    
            except Exception as e:
                logger.debug(f"LLM summarization failed: {str(e)}")
                # Fallback: simple truncation
                summary = content[:max_length] + "..." if len(content) > max_length else content
                return {
                    "summary": summary,
                    "note": f"Used simple truncation. LLM summarization failed: {str(e)}. Configure LLM for better summarization.",
                    "original_length": len(content),
                    "summary_length": len(summary),
                    "fallback": True
                }
        except Exception as e:
            logger.error(f"Error in summarization: {str(e)}", exc_info=True)
            return {"error": f"Summarization failed: {str(e)}"}
    
    def _read_file_content(self, file_path: str) -> str:
        """Read content from various file types, handling file_id lookups and different formats."""
        original_file_path = file_path
        
        # First, try to resolve file_id if it looks like one
        # Handle cases like "aeZN6x5WEJ_generated.docx" - extract file_id part
        file_id_candidate = file_path
        if '_' in file_path and '.' in file_path:
            # Might be "file_id_suffix.ext" format
            parts = file_path.rsplit('_', 1)
            if len(parts) == 2:
                file_id_candidate = parts[0]  # Try the part before last underscore
        
        # Try to look up by file_id or file_name
        try:
            from .models import ConversationFile, TrainingData
            conversation_file = None
            training_data_file = None
            
            # Try exact file_id match first
            conversation_file = ConversationFile.objects.filter(file_id=file_path).first()
            if not conversation_file and file_id_candidate != file_path:
                # Try with extracted file_id (e.g., "aeZN6x5WEJ" from "aeZN6x5WEJ_generated.docx")
                conversation_file = ConversationFile.objects.filter(file_id=file_id_candidate).first()
            
            # Try matching by file_name (exact match first, then contains)
            if not conversation_file:
                # Exact match
                conversation_file = ConversationFile.objects.filter(file_name=file_path).first()
            if not conversation_file:
                # Partial match (in case of path differences)
                conversation_file = ConversationFile.objects.filter(file_name__icontains=os.path.basename(file_path)).first()
            
            if conversation_file:
                # Get the actual file path
                if conversation_file.file_path:
                    if hasattr(conversation_file.file_path, 'name'):
                        file_path = conversation_file.file_path.name
                    else:
                        file_path = str(conversation_file.file_path)
                    logger.info(f"Resolved file '{original_file_path}' to path: {file_path}")
                elif hasattr(conversation_file, 'file_name'):
                    # Try to construct path from file_name
                    file_path = f"conversation_files/{conversation_file.agent.id}/{conversation_file.file_name}"
                    logger.info(f"Constructed path from file_name: {file_path}")
            
            # Also check TrainingData files
            if not conversation_file:
                # Try by file_path field
                training_data_file = TrainingData.objects.filter(file_path__icontains=os.path.basename(file_path)).first()
                if not training_data_file:
                    # Try by file name in content or metadata
                    training_data_file = TrainingData.objects.filter(
                        content__icontains=os.path.basename(file_path)
                    ).first()
                
                if training_data_file and training_data_file.file_path:
                    if hasattr(training_data_file.file_path, 'name'):
                        file_path = training_data_file.file_path.name
                    else:
                        file_path = str(training_data_file.file_path)
                    logger.info(f"Found file in TrainingData: {file_path}")
        except Exception as e:
            logger.debug(f"Could not resolve file: {str(e)}")
        
        # Check file extension
        file_ext = os.path.splitext(file_path)[1].lower() if file_path else ''
        
        # Try default_storage first
        if default_storage.exists(file_path):
            try:
                # Handle different file types
                if file_ext == '.docx':
                    # Word document - need special handling
                    try:
                        from docx import Document
                        file_obj = default_storage.open(file_path, 'rb')
                        doc = Document(file_obj)
                        content = []
                        for paragraph in doc.paragraphs:
                            if paragraph.text.strip():
                                content.append(paragraph.text)
                        file_obj.close()
                        return '\n'.join(content)
                    except ImportError:
                        return f"[Word document: {os.path.basename(file_path)} - python-docx not installed. Install with: pip install python-docx]"
                    except Exception as e:
                        logger.error(f"Error reading DOCX: {str(e)}")
                        return f"[Word document: {os.path.basename(file_path)} - Error: {str(e)}]"
                
                elif file_ext == '.pdf':
                    # PDF - use OCR tool's extraction method
                    try:
                        ocr_tool = OCRTool()
                        result = ocr_tool._extract_from_pdf(file_path)
                        if result.get('success') and result.get('text'):
                            return result['text']
                        else:
                            return f"[PDF file: {os.path.basename(file_path)} - Could not extract text. {result.get('note', '')}]"
                    except Exception as e:
                        logger.error(f"Error reading PDF: {str(e)}")
                        return f"[PDF file: {os.path.basename(file_path)} - Error: {str(e)}]"
                
                elif file_ext in ['.csv', '.tsv']:
                    # CSV/TSV files - read and format
                    try:
                        import csv
                        file_obj = default_storage.open(file_path, 'rb')
                        content_bytes = file_obj.read()
                        file_obj.close()
                        
                        # Try UTF-8 first, then fallback to latin-1
                        try:
                            content_str = content_bytes.decode('utf-8')
                        except UnicodeDecodeError:
                            content_str = content_bytes.decode('latin-1', errors='ignore')
                        
                        # Parse CSV and format
                        csv_reader = csv.reader(content_str.splitlines())
                        rows = []
                        for row in csv_reader:
                            if row:  # Skip empty rows
                                rows.append(', '.join(str(cell) for cell in row))
                        
                        return '\n'.join(rows) if rows else content_str
                    except Exception as e:
                        logger.error(f"Error reading CSV: {str(e)}")
                        return f"[CSV file: {os.path.basename(file_path)} - Error: {str(e)}]"
                
                else:
                    # Try to read as text
                    try:
                        # default_storage.open() doesn't support encoding parameter
                        # Read as binary and decode manually
                        file_obj = default_storage.open(file_path, 'rb')
                        content_bytes = file_obj.read()
                        file_obj.close()
                        
                        # Try UTF-8 first
                        try:
                            content = content_bytes.decode('utf-8')
                        except UnicodeDecodeError:
                            # Try other encodings
                            try:
                                content = content_bytes.decode('latin-1', errors='ignore')
                            except Exception:
                                return f"[Binary file: {os.path.basename(file_path)} - Cannot decode as text. Use OCR tool for images/PDFs]"
                        
                        return content
                    except Exception as e:
                        return f"[File: {os.path.basename(file_path)} - Error: {str(e)}]"
            except Exception as e:
                logger.error(f"Error reading from storage: {str(e)}")
                return f"[File: {os.path.basename(file_path)} - Storage error: {str(e)}]"
        
        # Try absolute path
        if os.path.exists(file_path):
            try:
                if file_ext == '.docx':
                    try:
                        from docx import Document
                        doc = Document(file_path)
                        content = []
                        for paragraph in doc.paragraphs:
                            if paragraph.text.strip():
                                content.append(paragraph.text)
                        return '\n'.join(content)
                    except ImportError:
                        return f"[Word document: {os.path.basename(file_path)} - python-docx not installed]"
                    except Exception as e:
                        return f"[Word document: {os.path.basename(file_path)} - Error: {str(e)}]"
                else:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        return f.read()
            except Exception as e:
                return f"[File: {os.path.basename(file_path)} - Error: {str(e)}]"
        
            # File not found - provide helpful error message with search results
            error_msg = f"File not found: {file_path}"
            suggestions = []
            similar_files = []
            
            # Try to find similar files in database with multiple strategies
            try:
                from .models import ConversationFile, TrainingData
                base_name = os.path.basename(file_path)
                base_name_no_ext = os.path.splitext(base_name)[0]  # Name without extension
                file_ext = os.path.splitext(base_name)[1].lower()  # Extension
                
                # Strategy 1: Search by exact filename (case-insensitive)
                conv_files = ConversationFile.objects.filter(file_name__iexact=base_name)[:5]
                for cf in conv_files:
                    similar_files.append(f"  - ConversationFile (exact match): {cf.file_name} (ID: {cf.file_id}, Agent: {cf.agent.id})")
                
                # Strategy 2: Search by partial filename (contains)
                if not similar_files:
                    conv_files = ConversationFile.objects.filter(file_name__icontains=base_name)[:5]
                    for cf in conv_files:
                        similar_files.append(f"  - ConversationFile (partial match): {cf.file_name} (ID: {cf.file_id}, Agent: {cf.agent.id})")
                
                # Strategy 3: Search by filename without extension
                if not similar_files and base_name_no_ext:
                    conv_files = ConversationFile.objects.filter(file_name__icontains=base_name_no_ext)[:5]
                    for cf in conv_files:
                        similar_files.append(f"  - ConversationFile (name match): {cf.file_name} (ID: {cf.file_id}, Agent: {cf.agent.id})")
                
                # Strategy 4: Search by extension only (if provided)
                if not similar_files and file_ext:
                    conv_files = ConversationFile.objects.filter(file_name__iendswith=file_ext)[:10]
                    for cf in conv_files[:5]:  # Limit to 5
                        similar_files.append(f"  - ConversationFile (same extension .{file_ext[1:]}): {cf.file_name} (ID: {cf.file_id}, Agent: {cf.agent.id})")
                
                # Strategy 5: Search TrainingData files
                train_files = TrainingData.objects.filter(
                    data_type='file',
                    file_path__icontains=base_name
                )[:5]
                for tf in train_files:
                    if tf.file_path:
                        path_name = tf.file_path.name if hasattr(tf.file_path, 'name') else str(tf.file_path)
                        similar_files.append(f"  - TrainingData: {path_name} (Agent: {tf.agent.id})")
                
                # Strategy 6: Search TrainingData by name without extension
                if not train_files.exists() and base_name_no_ext:
                    train_files = TrainingData.objects.filter(
                        data_type='file',
                        file_path__icontains=base_name_no_ext
                    )[:5]
                    for tf in train_files:
                        if tf.file_path:
                            path_name = tf.file_path.name if hasattr(tf.file_path, 'name') else str(tf.file_path)
                            similar_files.append(f"  - TrainingData (name match): {path_name} (Agent: {tf.agent.id})")
                
                if similar_files:
                    suggestions.append(f"Similar files found in database:\n" + "\n".join(similar_files))
            except Exception as e:
                logger.debug(f"Error searching for similar files: {str(e)}")
        
        # Check if it might be a file_id
        if len(file_path) < 50 and '_' not in file_path:
            suggestions.append(f"\n1. If '{file_path}' is a file_id, ensure the file was uploaded via the upload_file API")
            suggestions.append(f"2. Check if the file exists in ConversationFile model with file_id='{file_path}'")
        
        # Check common storage locations
        suggestions.append(f"\n3. Verify file exists in storage. Common locations:")
        suggestions.append(f"   - conversation_files/<agent_id>/<filename>")
        suggestions.append(f"   - chat_files/<agent_id>/<filename>")
        suggestions.append(f"   - training_data/<filename>")
        
        # File type specific suggestions
        file_ext = os.path.splitext(file_path)[1].lower()
        if file_ext == '.docx':
            suggestions.append(f"\n4. For .docx files, ensure python-docx is installed: pip install python-docx")
        elif file_ext == '.pdf':
            suggestions.append(f"\n4. For .pdf files, ensure PyPDF2 is installed: pip install PyPDF2")
        
        return f"[{error_msg}.\n" + "\n".join(suggestions) + "]"


class QuestionAnsweringTool(Tool):
    """Tool for answering questions based on documents."""
    
    def __init__(self):
        super().__init__(
            name="question_answering",
            description="This tool is used to answer questions based on text files (e.g., articles, pdfs, docx, etc.).",
            instructions="Use this tool when users ask questions about document content."
        )
    
    def get_parameters_schema(self) -> List[Dict[str, Any]]:
        return [
            {
                "name": "file_path",
                "description": "Path to the file to query.",
                "type": "string",
                "required": False
            },
            {
                "name": "text_content",
                "description": "Direct text content to query (alternative to file_path).",
                "type": "string",
                "required": False
            },
            {
                "name": "question",
                "description": "The question to answer based on the file content.",
                "type": "string",
                "required": True
            }
        ]
    
    def execute(self, parameters: Dict[str, Any]) -> Dict[str, Any]:
        file_path = parameters.get('file_path', '')
        question = parameters.get('question', '')
        text_content = parameters.get('text_content', '')  # Allow direct text input
        
        # Validate required parameters
        if not question:
            return {"error": "question parameter is required"}
        
        if not file_path and not text_content:
            return {"error": "Either file_path or text_content parameter is required"}
        
        try:
            # Read file content if file_path provided
            content = text_content.strip() if text_content else ''
            
            if file_path and not content:
                # Use OCR tool's comprehensive file reading (handles all file types)
                try:
                    ocr_tool = OCRTool()
                    result = ocr_tool.execute({"file_path": file_path})
                    
                    # Extract text from OCR tool result
                    if result.get('success'):
                        content = result.get('text', '')
                    elif result.get('text'):
                        content = result.get('text', '')
                    else:
                        error_msg = result.get('error', 'Unknown error')
                        return {"error": f"Could not read file: {file_path}", "details": error_msg}
                except Exception as e:
                    logger.error(f"Error using OCR tool for file reading: {str(e)}", exc_info=True)
                    # Fallback to SummarizationTool's method
                    summarization_tool = SummarizationTool()
                    content = summarization_tool._read_file_content(file_path)
                    if not content or (content.startswith('[') and 'Error' in content):
                        return {"error": f"Could not read file: {file_path}", "details": content}
            
            if not content:
                return {"error": "No content to answer questions about"}
            
            # Use LLM for question answering via Bedrock
            try:
                from .aws_integration import get_bedrock_client
                client = get_bedrock_client()
                
                # Optimize content length for token efficiency (reduce from 8000 to 6000 chars)
                # This leaves more room for the question and response
                max_content_length = 6000
                if len(content) > max_content_length:
                    # Try to truncate at a sentence boundary
                    truncated = content[:max_content_length]
                    last_period = truncated.rfind('.')
                    if last_period > max_content_length * 0.8:  # If we can keep at least 80%
                        content_for_qa = truncated[:last_period + 1]
                    else:
                        content_for_qa = truncated + "..."
                    logger.debug(f"Truncated content from {len(content)} to {len(content_for_qa)} chars")
                else:
                    content_for_qa = content
                
                # Optimize prompt for token usage
                prompt = f"Document:\n{content_for_qa}\n\nQ: {question}\nA:"
                
                result = client.invoke_agent(
                    agent_id="qa_tool",
                    query=prompt,
                    model="anthropic.claude-3-sonnet-20240229-v1:0",  # Use a Claude model
                    system_prompt="Answer questions accurately based on the provided document.",
                    model_provider="bedrock"
                )
                
                answer = result.get('response', '').strip()
                
                if not answer:
                    return {
                        "error": "No answer received from question answering service",
                        "question": question
                    }
                
                return {
                    "answer": answer,
                    "question": question,
                    "success": True
                }
            except Exception as e:
                logger.error(f"LLM question answering failed: {str(e)}", exc_info=True)
                return {
                    "error": f"Question answering service error: {str(e)}",
                    "question": question
                }
        except Exception as e:
            logger.error(f"Error in question answering: {str(e)}", exc_info=True)
            return {"error": f"Question answering failed: {str(e)}"}


class TranslationTool(Tool):
    """Tool for translating documents."""
    
    def __init__(self):
        super().__init__(
            name="translation",
            description="This tool is used to translate full text files (e.g., articles, pdfs, docx, etc.).",
            instructions="Use this tool when users request document translation."
        )
    
    def get_parameters_schema(self) -> List[Dict[str, Any]]:
        return [
            {
                "name": "file_path",
                "description": "Path to the file to translate.",
                "type": "string",
                "required": False
            },
            {
                "name": "text_content",
                "description": "Direct text content to translate (alternative to file_path).",
                "type": "string",
                "required": False
            },
            {
                "name": "target_language",
                "description": "Target language code (e.g., 'es', 'fr', 'de').",
                "type": "string",
                "required": True
            },
            {
                "name": "source_language",
                "description": "Source language code (optional, auto-detect if not provided).",
                "type": "string",
                "required": False
            }
        ]
    
    def execute(self, parameters: Dict[str, Any]) -> Dict[str, Any]:
        file_path = parameters.get('file_path', '')
        target_language = parameters.get('target_language', '')
        source_language = parameters.get('source_language', 'auto')
        text_content = parameters.get('text_content', '')  # Allow direct text input
        
        if not target_language:
            return {"error": "target_language parameter is required"}
        
        if not file_path and not text_content:
            return {"error": "Either file_path or text_content parameter is required"}
        
        try:
            # Read file content if file_path provided
            content = text_content
            if file_path and not content:
                # Use OCR tool's comprehensive file reading (handles all file types)
                try:
                    ocr_tool = OCRTool()
                    result = ocr_tool.execute({"file_path": file_path})
                    
                    # Extract text from OCR tool result
                    if result.get('success'):
                        content = result.get('text', '')
                    elif result.get('text'):
                        content = result.get('text', '')
                    else:
                        error_msg = result.get('error', 'Unknown error')
                        return {"error": f"Could not read file: {file_path}", "details": error_msg}
                except Exception as e:
                    logger.error(f"Error using OCR tool for file reading: {str(e)}", exc_info=True)
                    # Fallback to SummarizationTool's method
                    summarization_tool = SummarizationTool()
                    content = summarization_tool._read_file_content(file_path)
                    if not content or (content.startswith('[') and 'Error' in content):
                        return {"error": f"Could not read file: {file_path}", "details": content}
            
            if not content:
                return {"error": "No content to translate"}
            
            # Try AWS Translate if available
            try:
                from .aws_utils import create_boto3_client
                translate_client = create_boto3_client('translate')
                
                # Truncate if too long (AWS Translate has limits)
                content_to_translate = content[:10000] if len(content) > 10000 else content
                
                translate_params = {
                    'Text': content_to_translate,
                    'SourceLanguageCode': source_language if source_language != 'auto' else 'auto',
                    'TargetLanguageCode': target_language
                }
                
                response = translate_client.translate_text(**translate_params)
                
                return {
                    "translated_text": response['TranslatedText'],
                    "source_language": response.get('SourceLanguageCode', source_language),
                    "target_language": target_language,
                    "success": True
                }
            except Exception as e:
                logger.debug(f"AWS Translate not available: {str(e)}")
                # Fallback: Use LLM for translation
                try:
                    from .aws_integration import get_bedrock_client
                    client = get_bedrock_client()
                    
                    # Truncate content if too long
                    content_to_translate = content[:8000] if len(content) > 8000 else content
                    
                    prompt = f"""Translate the following text to {target_language}. 
Provide only the translated text, no explanations.

Text to translate:
{content_to_translate}

Translation:"""
                    
                    result = client.invoke_agent(
                        agent_id="translation_tool",
                        query=prompt,
                        model="anthropic.claude-3-sonnet-20240229-v1:0",
                        system_prompt="You are a professional translator. Translate accurately and preserve meaning.",
                        model_provider="bedrock"
                    )
                    
                    translated_text = result.get('response', '').strip()
                    
                    return {
                        "translated_text": translated_text,
                        "source_language": source_language,
                        "target_language": target_language,
                        "success": True,
                        "method": "llm"
                    }
                except Exception as e2:
                    logger.debug(f"LLM translation failed: {str(e2)}")
                    return {
                        "translated_text": f"Translation service not yet configured. Error: {str(e2)}",
                        "note": "Configure AWS Translate or LLM for translation functionality"
                    }
        except Exception as e:
            logger.error(f"Error in translation: {str(e)}", exc_info=True)
            return {"error": f"Translation failed: {str(e)}"}


class ReadFileTool(Tool):
    """Tool for reading file contents from various sources."""
    
    def __init__(self):
        super().__init__(
            name="read_file",
            description="Reads and returns the content of files. Supports text files, PDFs, Word documents, Excel files, CSV files, images, and more.",
            instructions="Use this tool when users ask to read, view, or access file contents. The tool can handle text files (.txt, .md, .json, etc.), PDFs, Word documents (.docx), Excel files (.xlsx), CSV files (.csv, .tsv), and other formats. CSV files are automatically parsed and formatted for easy reading."
        )
    
    def get_parameters_schema(self) -> List[Dict[str, Any]]:
        return [
            {
                "name": "file_path",
                "description": "Path to the file to read. Can be a filename, file_id, or full path. The tool will automatically search in ConversationFile and TrainingData models.",
                "type": "string",
                "required": True
            }
        ]
    
    def execute(self, parameters: Dict[str, Any]) -> Dict[str, Any]:
        file_path = parameters.get('file_path', '')
        if not file_path:
            return {"error": "file_path parameter is required"}
        
        try:
            # Use OCR tool's comprehensive file reading capabilities
            # It handles all file types including text, PDF, Word, Excel, images, etc.
            ocr_tool = OCRTool()
            result = ocr_tool.execute({"file_path": file_path})
            
            # Format the response for read_file tool
            if result.get('success'):
                return {
                    "content": result.get('text', ''),
                    "success": True,
                    "file_path": file_path,
                    "method": result.get('method', 'unknown'),
                    "file_type": result.get('file_type', 'unknown')
                }
            elif result.get('text'):
                # Even if not marked as success, return the text if available
                return {
                    "content": result.get('text', ''),
                    "success": True,
                    "file_path": file_path,
                    "note": result.get('note', ''),
                    "method": result.get('method', 'unknown')
                }
            else:
                # Return error information
                return {
                    "error": result.get('error', 'Unknown error'),
                    "file_path": file_path,
                    "note": result.get('note', ''),
                    "suggestion": result.get('suggestion', ''),
                    "resolved_path": result.get('resolved_path', '')
                }
        except Exception as e:
            logger.error(f"Error in read_file tool: {str(e)}", exc_info=True)
            return {
                "error": f"Failed to read file: {str(e)}",
                "file_path": file_path
            }


class WriteFileTool(Tool):
    """Tool for creating and saving files with content."""
    
    # Placeholder patterns that indicate the LLM didn't provide actual content
    PLACEHOLDER_PATTERNS = [
        'output', '[content goes here]', '[content]', '[text]', '[your content]',
        'placeholder', 'sample content', 'example content', 'content here',
        '[insert content]', '[add content]', 'todo', '[todo]', '...'
    ]
    
    def __init__(self):
        super().__init__(
            name="write_file",
            description="Creates a new file with the specified content and saves it for download. Supports text files, code files, markdown, JSON, CSV, HTML, and PDF files. Use this tool when users request to save content, create documents, write files, or generate downloadable files.",
            instructions="""Use this tool to create files with text content, code, documents, PDFs, or any other content that should be saved as a downloadable file.

PDF FILES:
- To create a PDF file, use file_name ending with .pdf (e.g., 'report.pdf', 'document.pdf')
- The content can be plain text or markdown-formatted text
- Markdown formatting is supported: # for headings, ## for subheadings, - or * for lists, **bold**, *italic*
- Optional parameters: pdf_title (document title), pdf_author (author name)
- Example: write_file(file_name='report.pdf', content='# Report Title\\n\\nContent...', pdf_title='Monthly Report')

CRITICAL REQUIREMENTS:
1. The 'content' parameter MUST contain the ACTUAL, COMPLETE content to be written to the file - NOT placeholders like 'Output', '[content goes here]', etc.
2. Content must be at least 20 characters of meaningful text.
3. If you need to create a file, generate the full content FIRST, then call this tool with that content.

IMPORTANT - After creating a file:
- The download link is AUTOMATICALLY displayed by the system below your message
- Do NOT generate fake download links like "📥 Download ready: filename.txt"
- Do NOT include URLs or emoji-based download text
- Simply say "I've created the file for you" or "The file has been saved" - the system handles the rest
- NEVER write markdown links like [Download](url) for created files

Example of WRONG response after creating a file:
"📥 Download ready: report.txt" or "[Download report.txt](https://...)"

Example of CORRECT response after creating a file:
"I've created the report for you. You can download it using the link that appears below."

Example of WRONG usage: {"file_name": "essay.txt", "content": "Output"}
Example of CORRECT usage: {"file_name": "essay.txt", "content": "This is the full text of the essay about climate change..."}"""
        )
    
    def get_parameters_schema(self) -> List[Dict[str, Any]]:
        return [
            {
                "name": "file_name",
                "description": "The name of the file to create (e.g., 'essay.txt', 'report.md', 'script.py', 'document.pdf'). Include the file extension.",
                "type": "string",
                "required": True
            },
            {
                "name": "content",
                "description": "The COMPLETE content to write to the file. Must be the actual content, NOT a placeholder. Minimum 20 characters required. For PDF files, this can be plain text or markdown-formatted text.",
                "type": "string",
                "required": True
            },
            {
                "name": "file_type",
                "description": "Optional MIME type of the file (e.g., 'text/plain', 'text/markdown', 'text/html', 'application/json', 'application/pdf'). If not provided, will be guessed from file extension.",
                "type": "string",
                "required": False
            },
            {
                "name": "pdf_title",
                "description": "Optional title for PDF documents (only used when creating PDF files).",
                "type": "string",
                "required": False
            },
            {
                "name": "pdf_author",
                "description": "Optional author name for PDF documents (only used when creating PDF files).",
                "type": "string",
                "required": False
            }
        ]
    
    def _is_placeholder_content(self, content: str) -> bool:
        """Check if content appears to be placeholder text."""
        if not content:
            return True
        
        content_lower = content.strip().lower()
        
        # Check against known placeholder patterns
        for pattern in self.PLACEHOLDER_PATTERNS:
            if content_lower == pattern or content_lower.startswith(pattern + '\n'):
                return True
        
        # Check for bracket-enclosed placeholders like [anything]
        if content_lower.startswith('[') and content_lower.endswith(']') and len(content_lower) < 50:
            return True
        
        return False
    
    def execute(self, parameters: Dict[str, Any]) -> Dict[str, Any]:
        file_name = parameters.get('file_name', '')
        content = parameters.get('content', '')
        
        if not file_name:
            return {"error": "file_name parameter is required"}
        
        if content is None:
            return {"error": "content parameter is required"}
        
        # Validate content is not empty or too short
        content_str = str(content).strip()
        if len(content_str) < 20:
            return {
                "error": f"Content is too short ({len(content_str)} characters). Please provide meaningful content of at least 20 characters. Do not use placeholders like 'Output' or '[content goes here]'.",
                "hint": "Generate the actual content you want to save, then call write_file with that complete content."
            }
        
        # Check for placeholder content
        if self._is_placeholder_content(content_str):
            return {
                "error": f"Content appears to be a placeholder ('{content_str[:50]}...'). Please provide the actual content to write to the file.",
                "hint": "Generate the complete content first, then call write_file with the full text."
            }
        
        try:
            import os
            
            # Check if this is a PDF file
            file_ext = os.path.splitext(file_name)[1].lower()
            is_pdf = file_ext == '.pdf' or parameters.get('file_type', '').lower() == 'application/pdf'
            
            # Generate PDF if needed
            if is_pdf:
                pdf_content = self._generate_pdf(content_str, parameters)
                if isinstance(pdf_content, dict) and 'error' in pdf_content:
                    return pdf_content  # Return error if PDF generation failed
                
                # Return PDF as binary content
                result = {
                    "file_content": pdf_content,  # Binary PDF content
                    "file_name": file_name,
                    "file_type": "application/pdf",
                    "success": True,
                    "message": f"PDF file '{file_name}' created successfully. The download link will appear automatically below your message - do NOT generate a fake download link or emoji-based download text."
                }
                return result
            
            # For non-PDF files, return text content as before
            result = {
                "file_content": content,  # This triggers automatic file registration
                "file_name": file_name,
                "success": True,
                "message": f"File '{file_name}' created successfully. The download link will appear automatically below your message - do NOT generate a fake download link or emoji-based download text."
            }
            
            # Add file_type if provided
            if parameters.get('file_type'):
                result['file_type'] = parameters.get('file_type')
            
            # Add file extension if not present in file_name
            if not os.path.splitext(file_name)[1]:
                # Guess extension from file_type or default to .txt
                file_type = parameters.get('file_type', '')
                if 'markdown' in file_type or 'md' in file_type:
                    result['file_extension'] = '.md'
                elif 'html' in file_type:
                    result['file_extension'] = '.html'
                elif 'json' in file_type:
                    result['file_extension'] = '.json'
                elif 'python' in file_type or 'py' in file_type:
                    result['file_extension'] = '.py'
                else:
                    result['file_extension'] = '.txt'
            
            return result
        except Exception as e:
            logger.error(f"Error in write_file tool: {str(e)}", exc_info=True)
            return {"error": f"Failed to create file: {str(e)}"}
    
    def _generate_pdf(self, content: str, parameters: Dict[str, Any]) -> bytes:
        """
        Generate PDF from text content.
        
        Args:
            content: Text content to convert to PDF
            parameters: Tool parameters (may contain pdf_title, pdf_author)
        
        Returns:
            bytes: PDF file content as bytes
        """
        try:
            # Try reportlab first (better for programmatic PDF generation)
            try:
                from reportlab.lib.pagesizes import A4
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib.units import inch
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                from reportlab.lib.enums import TA_LEFT, TA_CENTER
                from io import BytesIO
                import re
                
                # Create PDF in memory
                buffer = BytesIO()
                
                # Use A4 size (standard) or letter size
                page_size = A4
                
                # Create document
                doc = SimpleDocTemplate(
                    buffer,
                    pagesize=page_size,
                    rightMargin=72,
                    leftMargin=72,
                    topMargin=72,
                    bottomMargin=72
                )
                
                # Get styles
                styles = getSampleStyleSheet()
                
                # Create custom styles
                title_style = ParagraphStyle(
                    'CustomTitle',
                    parent=styles['Heading1'],
                    fontSize=18,
                    textColor='#000000',
                    spaceAfter=12,
                    alignment=TA_CENTER
                )
                
                heading_style = ParagraphStyle(
                    'CustomHeading',
                    parent=styles['Heading2'],
                    fontSize=14,
                    textColor='#333333',
                    spaceAfter=10,
                    spaceBefore=12
                )
                
                normal_style = ParagraphStyle(
                    'CustomNormal',
                    parent=styles['Normal'],
                    fontSize=11,
                    textColor='#000000',
                    spaceAfter=6,
                    alignment=TA_LEFT,
                    leading=14
                )
                
                # Build PDF content
                story = []
                
                # Add title if provided
                pdf_title = parameters.get('pdf_title', '')
                if pdf_title:
                    story.append(Paragraph(pdf_title, title_style))
                    story.append(Spacer(1, 0.3*inch))
                
                # Process content - handle markdown-style formatting
                lines = content.split('\n')
                current_paragraph = []
                
                for line in lines:
                    line = line.strip()
                    
                    if not line:
                        # Empty line - add current paragraph and spacer
                        if current_paragraph:
                            para_text = ' '.join(current_paragraph)
                            story.append(Paragraph(para_text, normal_style))
                            current_paragraph = []
                        story.append(Spacer(1, 0.1*inch))
                        continue
                    
                    # Check for markdown-style headings
                    if line.startswith('# '):
                        # H1 heading
                        if current_paragraph:
                            para_text = ' '.join(current_paragraph)
                            story.append(Paragraph(para_text, normal_style))
                            current_paragraph = []
                        story.append(Paragraph(line[2:], title_style))
                        story.append(Spacer(1, 0.2*inch))
                    elif line.startswith('## '):
                        # H2 heading
                        if current_paragraph:
                            para_text = ' '.join(current_paragraph)
                            story.append(Paragraph(para_text, normal_style))
                            current_paragraph = []
                        story.append(Paragraph(line[3:], heading_style))
                        story.append(Spacer(1, 0.15*inch))
                    elif line.startswith('### '):
                        # H3 heading
                        if current_paragraph:
                            para_text = ' '.join(current_paragraph)
                            story.append(Paragraph(para_text, normal_style))
                            current_paragraph = []
                        h3_style = ParagraphStyle(
                            'CustomH3',
                            parent=styles['Heading3'],
                            fontSize=12,
                            textColor='#444444',
                            spaceAfter=8,
                            spaceBefore=10
                        )
                        story.append(Paragraph(line[4:], h3_style))
                        story.append(Spacer(1, 0.1*inch))
                    elif line.startswith('- ') or line.startswith('* '):
                        # Bullet point
                        if current_paragraph:
                            para_text = ' '.join(current_paragraph)
                            story.append(Paragraph(para_text, normal_style))
                            current_paragraph = []
                        bullet_text = '• ' + line[2:]  # Replace - or * with bullet
                        story.append(Paragraph(bullet_text, normal_style))
                    elif line.startswith('1. ') or re.match(r'^\d+\.\s', line):
                        # Numbered list
                        if current_paragraph:
                            para_text = ' '.join(current_paragraph)
                            story.append(Paragraph(para_text, normal_style))
                            current_paragraph = []
                        story.append(Paragraph(line, normal_style))
                    else:
                        # Regular text - accumulate into paragraph
                        current_paragraph.append(line)
                
                # Add remaining paragraph
                if current_paragraph:
                    para_text = ' '.join(current_paragraph)
                    story.append(Paragraph(para_text, normal_style))
                
                # Build PDF
                doc.build(
                    story,
                    onFirstPage=self._add_pdf_header_footer(parameters.get('pdf_title', '')),
                    onLaterPages=self._add_pdf_header_footer(parameters.get('pdf_title', ''))
                )
                
                # Get PDF bytes
                pdf_bytes = buffer.getvalue()
                buffer.close()
                
                logger.info(f"Generated PDF with {len(pdf_bytes)} bytes using reportlab")
                return pdf_bytes
                
            except ImportError:
                # Fallback to weasyprint if reportlab not available
                logger.info("reportlab not available, trying weasyprint for PDF generation")
                try:
                    from weasyprint import HTML
                    from io import BytesIO
                    
                    # Convert text to HTML for PDF generation
                    # Basic markdown to HTML conversion
                    html_content = self._text_to_html(content)
                    
                    # Create PDF from HTML
                    pdf_bytes = HTML(string=html_content).write_pdf()
                    
                    logger.info(f"Generated PDF with {len(pdf_bytes)} bytes using weasyprint")
                    return pdf_bytes
                    
                except ImportError:
                    return {
                        "error": "PDF generation libraries not available. Install reportlab (pip install reportlab) or ensure weasyprint is installed.",
                        "suggestion": "pip install reportlab"
                    }
                except Exception as e:
                    logger.error(f"WeasyPrint PDF generation failed: {str(e)}", exc_info=True)
                    return {
                        "error": f"PDF generation failed: {str(e)}",
                        "suggestion": "Try installing reportlab: pip install reportlab"
                    }
            except Exception as e:
                logger.error(f"ReportLab PDF generation failed: {str(e)}", exc_info=True)
                # Try weasyprint as fallback
                try:
                    from weasyprint import HTML
                    html_content = self._text_to_html(content)
                    pdf_bytes = HTML(string=html_content).write_pdf()
                    logger.info(f"Generated PDF with {len(pdf_bytes)} bytes using weasyprint (fallback)")
                    return pdf_bytes
                except Exception as fallback_error:
                    return {
                        "error": f"PDF generation failed: {str(e)}",
                        "fallback_error": str(fallback_error),
                        "suggestion": "Ensure reportlab or weasyprint is installed: pip install reportlab"
                    }
        except Exception as e:
            logger.error(f"Error generating PDF: {str(e)}", exc_info=True)
            return {
                "error": f"PDF generation failed: {str(e)}",
                "suggestion": "Install PDF generation library: pip install reportlab"
            }
    
    def _text_to_html(self, content: str) -> str:
        """Convert plain text to HTML for PDF generation."""
        import re
        
        # Basic markdown to HTML conversion
        html = content
        
        # Headers
        html = re.sub(r'^# (.+)$', r'<h1>\1</h1>', html, flags=re.MULTILINE)
        html = re.sub(r'^## (.+)$', r'<h2>\1</h2>', html, flags=re.MULTILINE)
        html = re.sub(r'^### (.+)$', r'<h3>\1</h3>', html, flags=re.MULTILINE)
        
        # Bold and italic
        html = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', html)
        html = re.sub(r'\*(.+?)\*', r'<em>\1</em>', html)
        
        # Lists
        html = re.sub(r'^- (.+)$', r'<li>\1</li>', html, flags=re.MULTILINE)
        html = re.sub(r'(\n<li>.*</li>)', r'<ul>\1</ul>', html, flags=re.DOTALL)
        
        # Line breaks
        html = html.replace('\n\n', '</p><p>')
        html = html.replace('\n', '<br>')
        
        # Wrap in HTML structure
        html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <style>
        body {{
            font-family: Arial, sans-serif;
            line-height: 1.6;
            margin: 40px;
            color: #333;
        }}
        h1 {{ font-size: 24px; margin-top: 20px; margin-bottom: 10px; }}
        h2 {{ font-size: 20px; margin-top: 18px; margin-bottom: 8px; }}
        h3 {{ font-size: 16px; margin-top: 16px; margin-bottom: 6px; }}
        p {{ margin-bottom: 10px; }}
        ul {{ margin-left: 20px; }}
        li {{ margin-bottom: 5px; }}
    </style>
</head>
<body>
    <p>{html}</p>
</body>
</html>"""
        
        return html
    
    def _add_pdf_header_footer(self, canvas, doc, title: str):
        """Add header and footer to PDF pages."""
        from reportlab.lib.pagesizes import A4
        
        # Save state
        canvas.saveState()
        
        # Header
        canvas.setFont('Helvetica', 9)
        if title:
            canvas.drawString(72, A4[1] - 50, title)
        
        # Footer - page number
        page_num = canvas.getPageNumber()
        canvas.drawString(A4[0] - 100, 50, f"Page {page_num}")
        
        # Restore state
        canvas.restoreState()


class CustomTool(Tool):
    """Custom tool that can be configured by users."""
    
    def __init__(self, tool_config: Dict[str, Any]):
        self.tool_config = tool_config
        super().__init__(
            name=tool_config.get('function_name', ''),
            description=tool_config.get('description', ''),
            instructions=tool_config.get('instructions', '')
        )
        self.url = tool_config.get('url', '')
        self.method = tool_config.get('method', 'POST')
        self.headers = tool_config.get('headers', {})
        self.parameters_config = tool_config.get('parameters', [])
    
    def get_parameters_schema(self) -> List[Dict[str, Any]]:
        return self.parameters_config
    
    def execute(self, parameters: Dict[str, Any]) -> Dict[str, Any]:
        """Execute custom tool by making HTTP request."""
        if not self.url:
            return {"error": "Tool URL not configured"}
        
        try:
            # Build request URL with parameter substitution
            url = self.url
            for param_name, param_value in parameters.items():
                url = url.replace(f"$({param_name})", str(param_value))
            
            # Build request body
            body = {}
            for param_config in self.parameters_config:
                param_name = param_config.get('name', '')
                if param_name in parameters:
                    body[param_name] = parameters[param_name]
            
            # Make request
            if self.method.upper() == 'GET':
                response = requests.get(url, headers=self.headers, params=body, timeout=30)
            else:
                response = requests.request(
                    self.method.upper(),
                    url,
                    headers=self.headers,
                    json=body if body else None,
                    timeout=30
                )
            
            response.raise_for_status()
            
            return {
                "status_code": response.status_code,
                "data": response.json() if response.content else {},
                "success": True
            }
        except Exception as e:
            logger.error(f"Error executing custom tool: {str(e)}", exc_info=True)
            return {"error": f"Tool execution failed: {str(e)}", "success": False}


class SVGDiagramTool(Tool):
    """Tool for generating SVG diagrams including architectural massing models, floor plans, and technical drawings."""
    
    def __init__(self):
        super().__init__(
            name="svg_diagram",
            description="Generate SVG vector diagrams for buildings and architecture: massing models, floor plans, elevations, site plans. Use for ANY building/architecture/construction context (physical structures, dimensions, layouts). Pass user's description in 'description' param.",
            instructions="Use this tool to create architectural diagrams, massing models, floor plans, or technical drawings. ALWAYS pass the user's description in the 'description' parameter - it is used to infer style, type, and display on the diagram. Also provide dimensions (width, height, depth). The tool infers diagram_type from keywords: 'floor plan'/'layout' -> floor_plan, 'elevation'/'facade' -> elevation, 'courtyard'/'tower' -> adds features."
        )
    
    def get_parameters_schema(self) -> List[Dict[str, Any]]:
        return [
            {
                "name": "description",
                "description": "User's description of what to create (e.g., 'modern office building with courtyard', 'residential house floor plan'). Use this to infer style, type, and features.",
                "type": "string",
                "required": False
            },
            {
                "name": "diagram_type",
                "description": "Type of diagram: 'massing_model' (3D-like building mass), 'floor_plan', 'elevation' (front/side view), 'site_plan'. Can be inferred from description if not provided.",
                "type": "string",
                "required": False
            },
            {
                "name": "title",
                "description": "Title for the diagram",
                "type": "string",
                "required": True
            },
            {
                "name": "width",
                "description": "Building/space width in meters or feet (e.g., '50m' or '150ft')",
                "type": "string",
                "required": True
            },
            {
                "name": "height",
                "description": "Building/space height in meters or feet",
                "type": "string",
                "required": True
            },
            {
                "name": "depth",
                "description": "Building/space depth in meters or feet (optional, for 3D massing)",
                "type": "string",
                "required": False
            },
            {
                "name": "features",
                "description": "JSON array of features: [{\"type\": \"podium\"|\"tower\"|\"setback\"|\"courtyard\"|\"wing\", \"dimensions\": {...}, \"position\": {...}}]",
                "type": "string",
                "required": False
            },
            {
                "name": "style",
                "description": "Visual style: 'modern' (blue tones), 'classic' (warm tones), 'minimal' (grayscale), 'blueprint' (white on blue)",
                "type": "string",
                "required": False
            },
            {
                "name": "aspect_ratio",
                "description": "Output aspect ratio: '16:9', '4:3', '1:1', or 'A4'",
                "type": "string",
                "required": False
            }
        ]
    
    def _infer_from_description(self, description: str) -> Dict[str, Any]:
        """Infer diagram parameters from user description."""
        if not description:
            return {}
        desc_lower = description.lower()
        inferred = {}
        # Infer style
        if any(w in desc_lower for w in ['modern', 'contemporary', 'office', 'commercial']):
            inferred['style'] = 'modern'
        elif any(w in desc_lower for w in ['classic', 'traditional', 'vintage', 'heritage']):
            inferred['style'] = 'classic'
        elif any(w in desc_lower for w in ['minimal', 'simple', 'plain']):
            inferred['style'] = 'minimal'
        elif any(w in desc_lower for w in ['blueprint', 'technical', 'draft']):
            inferred['style'] = 'blueprint'
        # Infer diagram type
        if any(w in desc_lower for w in ['floor plan', 'floorplan', 'layout', 'rooms']):
            inferred['diagram_type'] = 'floor_plan'
        elif any(w in desc_lower for w in ['elevation', 'front view', 'facade', 'side view']):
            inferred['diagram_type'] = 'elevation'
        elif any(w in desc_lower for w in ['site plan', 'siteplan', 'site']):
            inferred['diagram_type'] = 'site_plan'
        # Infer features from description
        features = []
        if any(w in desc_lower for w in ['courtyard', 'atrium', 'inner court']):
            features.append({'type': 'courtyard', 'position': {'x': 0.3, 'y': 0.3}, 'dimensions': {'w': 0.3, 'd': 0.3}})
        if any(w in desc_lower for w in ['tower', 'high-rise', 'skyscraper']):
            features.append({'type': 'tower', 'position': {'x': 0.6, 'y': 0}, 'dimensions': {'w': 0.25, 'h': 0.6}})
        if any(w in desc_lower for w in ['podium', 'base', 'plinth']):
            features.append({'type': 'podium', 'position': {'x': 0, 'y': 0.7}, 'dimensions': {'h': 0.3}})
        if features:
            inferred['features'] = features
        return inferred
    
    def execute(self, parameters: Dict[str, Any]) -> Dict[str, Any]:
        description = parameters.get('description', '')
        inferred = self._infer_from_description(description)
        
        diagram_type = parameters.get('diagram_type') or inferred.get('diagram_type', 'massing_model')
        title = parameters.get('title', 'Architectural Diagram')
        width_str = parameters.get('width', '50m')
        height_str = parameters.get('height', '30m')
        depth_str = parameters.get('depth', '40m')
        features_str = parameters.get('features', '[]')
        style = parameters.get('style') or inferred.get('style', 'modern')
        aspect_ratio = parameters.get('aspect_ratio', '16:9')
        
        try:
            # Parse dimensions
            def parse_dimension(dim_str):
                """Extract numeric value and unit from dimension string."""
                import re
                match = re.match(r'([\d.]+)\s*(m|ft|meters?|feet)?', str(dim_str).lower())
                if match:
                    return float(match.group(1)), match.group(2) or 'm'
                return float(dim_str), 'm'
            
            width_val, width_unit = parse_dimension(width_str)
            height_val, height_unit = parse_dimension(height_str)
            depth_val, depth_unit = parse_dimension(depth_str)
            
            # Parse features (merge inferred from description with explicit params)
            try:
                features = json.loads(features_str) if features_str else []
            except Exception:
                features = []
            if inferred.get('features'):
                features = inferred['features'] + features
            
            # Set canvas dimensions based on aspect ratio
            aspect_map = {
                '16:9': (1920, 1080),
                '4:3': (1600, 1200),
                '1:1': (1200, 1200),
                'A4': (1190, 1684)
            }
            canvas_width, canvas_height = aspect_map.get(aspect_ratio, (1920, 1080))
            
            # Style configurations
            styles = {
                'modern': {
                    'bg': '#f8fafc',
                    'primary': '#3b82f6',
                    'secondary': '#60a5fa',
                    'accent': '#1e40af',
                    'text': '#1e293b',
                    'grid': '#e2e8f0',
                    'dimension': '#64748b'
                },
                'classic': {
                    'bg': '#faf7f2',
                    'primary': '#b45309',
                    'secondary': '#d97706',
                    'accent': '#92400e',
                    'text': '#44403c',
                    'grid': '#e7e5e4',
                    'dimension': '#78716c'
                },
                'minimal': {
                    'bg': '#ffffff',
                    'primary': '#374151',
                    'secondary': '#6b7280',
                    'accent': '#111827',
                    'text': '#111827',
                    'grid': '#f3f4f6',
                    'dimension': '#4b5563'
                },
                'blueprint': {
                    'bg': '#1e3a5f',
                    'primary': '#ffffff',
                    'secondary': '#93c5fd',
                    'accent': '#60a5fa',
                    'text': '#ffffff',
                    'grid': '#2563eb',
                    'dimension': '#bfdbfe'
                }
            }
            colors = styles.get(style, styles['modern'])
            
            # Truncate description for display (max 80 chars)
            desc_display = (description[:77] + '...') if description and len(description) > 80 else (description or '')
            
            # Generate SVG based on diagram type
            if diagram_type == 'massing_model':
                svg_content = self._generate_massing_model(
                    canvas_width, canvas_height, title, desc_display,
                    width_val, height_val, depth_val,
                    width_unit, features, colors
                )
            elif diagram_type == 'elevation':
                svg_content = self._generate_elevation(
                    canvas_width, canvas_height, title, desc_display,
                    width_val, height_val, width_unit, features, colors
                )
            elif diagram_type == 'floor_plan':
                svg_content = self._generate_floor_plan(
                    canvas_width, canvas_height, title, desc_display,
                    width_val, depth_val, width_unit, features, colors
                )
            else:
                svg_content = self._generate_massing_model(
                    canvas_width, canvas_height, title, desc_display,
                    width_val, height_val, depth_val,
                    width_unit, features, colors
                )
            
            # Save SVG to storage
            import hashlib
            file_hash = hashlib.md5(f"{title}{width_str}{height_str}{depth_str}".encode()).hexdigest()[:12]
            file_name = f"diagram_{file_hash}.svg"
            saved_path = default_storage.save(f"generated_diagrams/{file_name}", ContentFile(svg_content.encode('utf-8')))
            
            return {
                "image_url": default_storage.url(saved_path),
                "file_name": file_name,
                "file_type": "image/svg+xml",
                "success": True,
                "diagram_type": diagram_type,
                "dimensions": {
                    "width": f"{width_val}{width_unit}",
                    "height": f"{height_val}{height_unit}",
                    "depth": f"{depth_val}{depth_unit}"
                },
                "note": "SVG diagram generated successfully. This is a scalable vector graphic that can be viewed in any browser and exported to other formats."
            }
            
        except Exception as e:
            logger.error(f"SVG diagram generation error: {str(e)}", exc_info=True)
            return {"error": f"Failed to generate SVG diagram: {str(e)}"}
    
    def _generate_massing_model(self, canvas_w, canvas_h, title, description, width, height, depth, unit, features, colors):
        """Generate isometric massing model SVG."""
        from html import escape
        desc_safe = escape(description) if description else ''
        desc_line = f'<text x="{canvas_w/2}" y="115" text-anchor="middle" class="label" fill="{colors["dimension"]}" font-size="16" font-style="italic">{desc_safe}</text>' if desc_safe else ''
        # Calculate drawing area (with margins for dimensions)
        margin = 150
        draw_w = canvas_w - 2 * margin
        draw_h = canvas_h - 2 * margin - 80  # Extra space for title
        
        # Scale factor to fit building in drawing area
        scale = min(draw_w / (width + depth * 0.5), draw_h / (height + depth * 0.3)) * 0.6
        
        # Isometric projection angles
        iso_x = 0.866  # cos(30°)
        iso_y = 0.5    # sin(30°)
        
        # Building base position (center of canvas)
        base_x = canvas_w / 2 - (width * scale * iso_x) / 2
        base_y = canvas_h / 2 + (height * scale) / 3
        
        # Calculate isometric corners
        def iso_point(x, y, z):
            """Convert 3D coordinates to isometric 2D."""
            px = base_x + (x - z) * iso_x * scale
            py = base_y - y * scale + (x + z) * iso_y * scale / 2
            return px, py
        
        # Building corners
        p000 = iso_point(0, 0, 0)
        p100 = iso_point(width, 0, 0)
        p010 = iso_point(0, height, 0)
        p110 = iso_point(width, height, 0)
        p101 = iso_point(width, 0, depth)
        p011 = iso_point(0, height, depth)
        p111 = iso_point(width, height, depth)
        
        svg = f'''<?xml version="1.0" encoding="UTF-8"?>
<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 {canvas_w} {canvas_h}" width="{canvas_w}" height="{canvas_h}">
  <defs>
    <linearGradient id="frontGrad" x1="0%" y1="0%" x2="0%" y2="100%">
      <stop offset="0%" style="stop-color:{colors['primary']};stop-opacity:1" />
      <stop offset="100%" style="stop-color:{colors['accent']};stop-opacity:1" />
    </linearGradient>
    <linearGradient id="sideGrad" x1="0%" y1="0%" x2="100%" y2="0%">
      <stop offset="0%" style="stop-color:{colors['secondary']};stop-opacity:1" />
      <stop offset="100%" style="stop-color:{colors['primary']};stop-opacity:1" />
    </linearGradient>
    <linearGradient id="topGrad" x1="0%" y1="100%" x2="0%" y2="0%">
      <stop offset="0%" style="stop-color:{colors['secondary']};stop-opacity:0.9" />
      <stop offset="100%" style="stop-color:{colors['secondary']};stop-opacity:0.7" />
    </linearGradient>
    <filter id="shadow" x="-20%" y="-20%" width="140%" height="140%">
      <feDropShadow dx="5" dy="10" stdDeviation="8" flood-color="#000" flood-opacity="0.15"/>
    </filter>
    <marker id="arrow" markerWidth="10" markerHeight="10" refX="9" refY="3" orient="auto">
      <path d="M0,0 L0,6 L9,3 z" fill="{colors['dimension']}"/>
    </marker>
    <style>
      .title {{ font-family: 'Segoe UI', Arial, sans-serif; font-weight: 600; }}
      .dim-text {{ font-family: 'Segoe UI', Arial, sans-serif; font-weight: 500; }}
      .label {{ font-family: 'Segoe UI', Arial, sans-serif; }}
    </style>
  </defs>
  
  <!-- Background -->
  <rect width="{canvas_w}" height="{canvas_h}" fill="{colors['bg']}"/>
  
  <!-- Grid pattern -->
  <g stroke="{colors['grid']}" stroke-width="1" opacity="0.5">
    {"".join([f'<line x1="0" y1="{y}" x2="{canvas_w}" y2="{y}"/>' for y in range(0, canvas_h, 50)])}
    {"".join([f'<line x1="{x}" y1="0" x2="{x}" y2="{canvas_h}"/>' for x in range(0, canvas_w, 50)])}
  </g>
  
  <!-- Title -->
  <text x="{canvas_w/2}" y="50" text-anchor="middle" class="title" fill="{colors['text']}" font-size="36">
    {title}
  </text>
  <text x="{canvas_w/2}" y="85" text-anchor="middle" class="label" fill="{colors['dimension']}" font-size="20">
    CONCEPTUAL MASSING MODEL - ISOMETRIC VIEW
  </text>
  {desc_line}
  
  <!-- Building mass with shadow -->
  <g filter="url(#shadow)">
    <!-- Right side face -->
    <polygon points="{p100[0]},{p100[1]} {p110[0]},{p110[1]} {p111[0]},{p111[1]} {p101[0]},{p101[1]}" 
             fill="url(#sideGrad)" stroke="{colors['accent']}" stroke-width="2"/>
    
    <!-- Front face -->
    <polygon points="{p000[0]},{p000[1]} {p100[0]},{p100[1]} {p110[0]},{p110[1]} {p010[0]},{p010[1]}" 
             fill="url(#frontGrad)" stroke="{colors['accent']}" stroke-width="2"/>
    
    <!-- Top face -->
    <polygon points="{p010[0]},{p010[1]} {p110[0]},{p110[1]} {p111[0]},{p111[1]} {p011[0]},{p011[1]}" 
             fill="url(#topGrad)" stroke="{colors['accent']}" stroke-width="2"/>
  </g>
  
  <!-- Dimension lines and labels -->
  <!-- Width dimension (bottom front) -->
  <g class="dimension" stroke="{colors['dimension']}" stroke-width="2">
    <line x1="{p000[0]}" y1="{p000[1] + 40}" x2="{p100[0]}" y2="{p100[1] + 40}" marker-start="url(#arrow)" marker-end="url(#arrow)"/>
    <line x1="{p000[0]}" y1="{p000[1]}" x2="{p000[0]}" y2="{p000[1] + 45}" stroke-dasharray="5,5"/>
    <line x1="{p100[0]}" y1="{p100[1]}" x2="{p100[0]}" y2="{p100[1] + 45}" stroke-dasharray="5,5"/>
  </g>
  <text x="{(p000[0] + p100[0])/2}" y="{p000[1] + 70}" text-anchor="middle" class="dim-text" fill="{colors['dimension']}" font-size="24">
    {width:.1f}{unit} WIDTH
  </text>
  
  <!-- Height dimension (left front) -->
  <g class="dimension" stroke="{colors['dimension']}" stroke-width="2">
    <line x1="{p000[0] - 50}" y1="{p000[1]}" x2="{p000[0] - 50}" y2="{p010[1]}"/>
    <line x1="{p000[0] - 55}" y1="{p000[1]}" x2="{p000[0] - 45}" y2="{p000[1]}"/>
    <line x1="{p000[0] - 55}" y1="{p010[1]}" x2="{p000[0] - 45}" y2="{p010[1]}"/>
  </g>
  <text x="{p000[0] - 70}" y="{(p000[1] + p010[1])/2}" text-anchor="middle" class="dim-text" fill="{colors['dimension']}" font-size="24" transform="rotate(-90, {p000[0] - 70}, {(p000[1] + p010[1])/2})">
    {height:.1f}{unit} HEIGHT
  </text>
  
  <!-- Depth dimension (bottom right) -->
  <g class="dimension" stroke="{colors['dimension']}" stroke-width="2">
    <line x1="{p100[0] + 30}" y1="{p100[1] + 15}" x2="{p101[0] + 30}" y2="{p101[1] + 15}"/>
    <line x1="{p100[0]}" y1="{p100[1]}" x2="{p100[0] + 35}" y2="{p100[1] + 17}" stroke-dasharray="5,5"/>
    <line x1="{p101[0]}" y1="{p101[1]}" x2="{p101[0] + 35}" y2="{p101[1] + 17}" stroke-dasharray="5,5"/>
  </g>
  <text x="{(p100[0] + p101[0])/2 + 60}" y="{(p100[1] + p101[1])/2 + 40}" text-anchor="middle" class="dim-text" fill="{colors['dimension']}" font-size="24">
    {depth:.1f}{unit} DEPTH
  </text>
  
  <!-- Scale indicator -->
  <g transform="translate(100, {canvas_h - 80})">
    <rect x="0" y="0" width="200" height="30" fill="{colors['bg']}" stroke="{colors['dimension']}" stroke-width="1" rx="5"/>
    <text x="100" y="22" text-anchor="middle" class="label" fill="{colors['dimension']}" font-size="14">
      Scale: Conceptual | Not to Scale
    </text>
  </g>
  
  <!-- North arrow -->
  <g transform="translate({canvas_w - 100}, {canvas_h - 100})">
    <circle cx="0" cy="0" r="30" fill="none" stroke="{colors['dimension']}" stroke-width="2"/>
    <polygon points="0,-25 -8,10 0,0 8,10" fill="{colors['accent']}"/>
    <text x="0" y="-35" text-anchor="middle" class="label" fill="{colors['dimension']}" font-size="14">N</text>
  </g>
  
  <!-- Legend/Info box -->
  <g transform="translate({canvas_w - 300}, 120)">
    <rect x="0" y="0" width="250" height="120" fill="{colors['bg']}" stroke="{colors['grid']}" stroke-width="2" rx="8" opacity="0.95"/>
    <text x="125" y="30" text-anchor="middle" class="label" fill="{colors['text']}" font-size="16" font-weight="600">DIMENSIONS</text>
    <line x1="20" y1="45" x2="230" y2="45" stroke="{colors['grid']}" stroke-width="1"/>
    <text x="30" y="70" class="label" fill="{colors['dimension']}" font-size="14">Width: {width:.1f}{unit}</text>
    <text x="30" y="90" class="label" fill="{colors['dimension']}" font-size="14">Height: {height:.1f}{unit}</text>
    <text x="30" y="110" class="label" fill="{colors['dimension']}" font-size="14">Depth: {depth:.1f}{unit}</text>
  </g>
</svg>'''
        
        return svg
    
    def _generate_elevation(self, canvas_w, canvas_h, title, description, width, height, unit, features, colors):
        """Generate front elevation SVG."""
        from html import escape
        desc_safe = escape(description) if description else ''
        desc_line = f'<text x="{canvas_w/2}" y="115" text-anchor="middle" class="label" fill="{colors["dimension"]}" font-size="16" font-style="italic">{desc_safe}</text>' if desc_safe else ''
        margin = 150
        draw_w = canvas_w - 2 * margin
        draw_h = canvas_h - 2 * margin - 80
        
        # Scale to fit
        scale = min(draw_w / width, draw_h / height) * 0.7
        
        # Building position (centered)
        bldg_w = width * scale
        bldg_h = height * scale
        bldg_x = (canvas_w - bldg_w) / 2
        bldg_y = canvas_h - margin - bldg_h - 50
        
        # Calculate floor lines (assume ~3.5m per floor)
        floor_height = 3.5
        num_floors = max(1, int(height / floor_height))
        floor_h_px = bldg_h / num_floors
        
        svg = f'''<?xml version="1.0" encoding="UTF-8"?>
<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 {canvas_w} {canvas_h}" width="{canvas_w}" height="{canvas_h}">
  <defs>
    <linearGradient id="bldgGrad" x1="0%" y1="0%" x2="0%" y2="100%">
      <stop offset="0%" style="stop-color:{colors['secondary']};stop-opacity:1" />
      <stop offset="100%" style="stop-color:{colors['primary']};stop-opacity:1" />
    </linearGradient>
    <pattern id="windows" patternUnits="userSpaceOnUse" width="60" height="{floor_h_px}">
      <rect x="10" y="10" width="40" height="{floor_h_px - 20}" fill="{colors['accent']}" opacity="0.3" rx="2"/>
    </pattern>
    <style>
      .title {{ font-family: 'Segoe UI', Arial, sans-serif; font-weight: 600; }}
      .dim-text {{ font-family: 'Segoe UI', Arial, sans-serif; font-weight: 500; }}
      .label {{ font-family: 'Segoe UI', Arial, sans-serif; }}
    </style>
  </defs>
  
  <!-- Background -->
  <rect width="{canvas_w}" height="{canvas_h}" fill="{colors['bg']}"/>
  
  <!-- Title -->
  <text x="{canvas_w/2}" y="50" text-anchor="middle" class="title" fill="{colors['text']}" font-size="36">
    {title}
  </text>
  <text x="{canvas_w/2}" y="85" text-anchor="middle" class="label" fill="{colors['dimension']}" font-size="20">
    FRONT ELEVATION
  </text>
  {desc_line}
  
  <!-- Ground line -->
  <line x1="{margin - 50}" y1="{bldg_y + bldg_h}" x2="{canvas_w - margin + 50}" y2="{bldg_y + bldg_h}" 
        stroke="{colors['accent']}" stroke-width="3"/>
  <line x1="{margin - 50}" y1="{bldg_y + bldg_h + 5}" x2="{canvas_w - margin + 50}" y2="{bldg_y + bldg_h + 5}" 
        stroke="{colors['dimension']}" stroke-width="1" stroke-dasharray="10,5"/>
  
  <!-- Building main mass -->
  <rect x="{bldg_x}" y="{bldg_y}" width="{bldg_w}" height="{bldg_h}" 
        fill="url(#bldgGrad)" stroke="{colors['accent']}" stroke-width="3"/>
  
  <!-- Window pattern overlay -->
  <rect x="{bldg_x}" y="{bldg_y}" width="{bldg_w}" height="{bldg_h}" 
        fill="url(#windows)"/>
  
  <!-- Floor lines -->
  {"".join([f'<line x1="{bldg_x}" y1="{bldg_y + i * floor_h_px}" x2="{bldg_x + bldg_w}" y2="{bldg_y + i * floor_h_px}" stroke="{colors["accent"]}" stroke-width="1" opacity="0.5"/>' for i in range(1, num_floors)])}
  
  <!-- Floor labels -->
  {"".join([f'<text x="{bldg_x - 40}" y="{bldg_y + bldg_h - i * floor_h_px - floor_h_px/2 + 5}" text-anchor="end" class="label" fill="{colors["dimension"]}" font-size="12">L{i}</text>' for i in range(num_floors)])}
  
  <!-- Width dimension -->
  <g stroke="{colors['dimension']}" stroke-width="2">
    <line x1="{bldg_x}" y1="{bldg_y + bldg_h + 50}" x2="{bldg_x + bldg_w}" y2="{bldg_y + bldg_h + 50}"/>
    <line x1="{bldg_x}" y1="{bldg_y + bldg_h}" x2="{bldg_x}" y2="{bldg_y + bldg_h + 55}" stroke-dasharray="5,5"/>
    <line x1="{bldg_x + bldg_w}" y1="{bldg_y + bldg_h}" x2="{bldg_x + bldg_w}" y2="{bldg_y + bldg_h + 55}" stroke-dasharray="5,5"/>
  </g>
  <text x="{bldg_x + bldg_w/2}" y="{bldg_y + bldg_h + 80}" text-anchor="middle" class="dim-text" fill="{colors['dimension']}" font-size="24">
    {width:.1f}{unit}
  </text>
  
  <!-- Height dimension -->
  <g stroke="{colors['dimension']}" stroke-width="2">
    <line x1="{bldg_x + bldg_w + 50}" y1="{bldg_y}" x2="{bldg_x + bldg_w + 50}" y2="{bldg_y + bldg_h}"/>
    <line x1="{bldg_x + bldg_w}" y1="{bldg_y}" x2="{bldg_x + bldg_w + 55}" y2="{bldg_y}" stroke-dasharray="5,5"/>
    <line x1="{bldg_x + bldg_w}" y1="{bldg_y + bldg_h}" x2="{bldg_x + bldg_w + 55}" y2="{bldg_y + bldg_h}" stroke-dasharray="5,5"/>
  </g>
  <text x="{bldg_x + bldg_w + 80}" y="{bldg_y + bldg_h/2}" text-anchor="middle" class="dim-text" fill="{colors['dimension']}" font-size="24" 
        transform="rotate(-90, {bldg_x + bldg_w + 80}, {bldg_y + bldg_h/2})">
    {height:.1f}{unit}
  </text>
  
  <!-- Info box -->
  <g transform="translate({canvas_w - 280}, 120)">
    <rect x="0" y="0" width="230" height="100" fill="{colors['bg']}" stroke="{colors['grid']}" stroke-width="2" rx="8" opacity="0.95"/>
    <text x="115" y="28" text-anchor="middle" class="label" fill="{colors['text']}" font-size="16" font-weight="600">BUILDING INFO</text>
    <line x1="15" y1="40" x2="215" y2="40" stroke="{colors['grid']}" stroke-width="1"/>
    <text x="25" y="62" class="label" fill="{colors['dimension']}" font-size="13">Width: {width:.1f}{unit}</text>
    <text x="25" y="82" class="label" fill="{colors['dimension']}" font-size="13">Height: {height:.1f}{unit}</text>
    <text x="130" y="62" class="label" fill="{colors['dimension']}" font-size="13">Floors: {num_floors}</text>
  </g>
</svg>'''
        
        return svg
    
    def _generate_floor_plan(self, canvas_w, canvas_h, title, description, width, depth, unit, features, colors):
        """Generate floor plan SVG."""
        from html import escape
        desc_safe = escape(description) if description else ''
        desc_line = f'<text x="{canvas_w/2}" y="115" text-anchor="middle" class="label" fill="{colors["dimension"]}" font-size="16" font-style="italic">{desc_safe}</text>' if desc_safe else ''
        margin = 150
        draw_w = canvas_w - 2 * margin
        draw_h = canvas_h - 2 * margin - 80
        
        # Scale to fit
        scale = min(draw_w / width, draw_h / depth) * 0.7
        
        # Building position (centered)
        bldg_w = width * scale
        bldg_d = depth * scale
        bldg_x = (canvas_w - bldg_w) / 2
        bldg_y = (canvas_h - bldg_d) / 2 + 40
        
        svg = f'''<?xml version="1.0" encoding="UTF-8"?>
<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 {canvas_w} {canvas_h}" width="{canvas_w}" height="{canvas_h}">
  <defs>
    <pattern id="hatch" patternUnits="userSpaceOnUse" width="10" height="10">
      <path d="M-1,1 l2,-2 M0,10 l10,-10 M9,11 l2,-2" stroke="{colors['dimension']}" stroke-width="1" opacity="0.3"/>
    </pattern>
    <style>
      .title {{ font-family: 'Segoe UI', Arial, sans-serif; font-weight: 600; }}
      .dim-text {{ font-family: 'Segoe UI', Arial, sans-serif; font-weight: 500; }}
      .label {{ font-family: 'Segoe UI', Arial, sans-serif; }}
    </style>
  </defs>
  
  <!-- Background -->
  <rect width="{canvas_w}" height="{canvas_h}" fill="{colors['bg']}"/>
  
  <!-- Grid -->
  <g stroke="{colors['grid']}" stroke-width="1" opacity="0.3">
    {"".join([f'<line x1="0" y1="{y}" x2="{canvas_w}" y2="{y}"/>' for y in range(0, canvas_h, 50)])}
    {"".join([f'<line x1="{x}" y1="0" x2="{x}" y2="{canvas_h}"/>' for x in range(0, canvas_w, 50)])}
  </g>
  
  <!-- Title -->
  <text x="{canvas_w/2}" y="50" text-anchor="middle" class="title" fill="{colors['text']}" font-size="36">
    {title}
  </text>
  <text x="{canvas_w/2}" y="85" text-anchor="middle" class="label" fill="{colors['dimension']}" font-size="20">
    FLOOR PLAN
  </text>
  {desc_line}
  
  <!-- Floor plate outline (thick walls) -->
  <rect x="{bldg_x}" y="{bldg_y}" width="{bldg_w}" height="{bldg_d}" 
        fill="#ffffff" stroke="{colors['accent']}" stroke-width="8"/>
  
  <!-- Interior hatching -->
  <rect x="{bldg_x + 15}" y="{bldg_y + 15}" width="{bldg_w - 30}" height="{bldg_d - 30}" 
        fill="url(#hatch)"/>
  
  <!-- Structural grid (columns) -->
  {self._generate_column_grid(bldg_x, bldg_y, bldg_w, bldg_d, colors)}
  
  <!-- Width dimension (bottom) -->
  <g stroke="{colors['dimension']}" stroke-width="2">
    <line x1="{bldg_x}" y1="{bldg_y + bldg_d + 50}" x2="{bldg_x + bldg_w}" y2="{bldg_y + bldg_d + 50}"/>
    <line x1="{bldg_x}" y1="{bldg_y + bldg_d}" x2="{bldg_x}" y2="{bldg_y + bldg_d + 55}" stroke-dasharray="5,5"/>
    <line x1="{bldg_x + bldg_w}" y1="{bldg_y + bldg_d}" x2="{bldg_x + bldg_w}" y2="{bldg_y + bldg_d + 55}" stroke-dasharray="5,5"/>
  </g>
  <text x="{bldg_x + bldg_w/2}" y="{bldg_y + bldg_d + 80}" text-anchor="middle" class="dim-text" fill="{colors['dimension']}" font-size="24">
    {width:.1f}{unit}
  </text>
  
  <!-- Depth dimension (right) -->
  <g stroke="{colors['dimension']}" stroke-width="2">
    <line x1="{bldg_x + bldg_w + 50}" y1="{bldg_y}" x2="{bldg_x + bldg_w + 50}" y2="{bldg_y + bldg_d}"/>
    <line x1="{bldg_x + bldg_w}" y1="{bldg_y}" x2="{bldg_x + bldg_w + 55}" y2="{bldg_y}" stroke-dasharray="5,5"/>
    <line x1="{bldg_x + bldg_w}" y1="{bldg_y + bldg_d}" x2="{bldg_x + bldg_w + 55}" y2="{bldg_y + bldg_d}" stroke-dasharray="5,5"/>
  </g>
  <text x="{bldg_x + bldg_w + 80}" y="{bldg_y + bldg_d/2}" text-anchor="middle" class="dim-text" fill="{colors['dimension']}" font-size="24"
        transform="rotate(-90, {bldg_x + bldg_w + 80}, {bldg_y + bldg_d/2})">
    {depth:.1f}{unit}
  </text>
  
  <!-- North arrow -->
  <g transform="translate(100, {canvas_h - 100})">
    <circle cx="0" cy="0" r="30" fill="none" stroke="{colors['dimension']}" stroke-width="2"/>
    <polygon points="0,-25 -8,10 0,0 8,10" fill="{colors['accent']}"/>
    <text x="0" y="-35" text-anchor="middle" class="label" fill="{colors['dimension']}" font-size="14">N</text>
  </g>
  
  <!-- Area calculation -->
  <g transform="translate({canvas_w - 280}, 120)">
    <rect x="0" y="0" width="230" height="120" fill="{colors['bg']}" stroke="{colors['grid']}" stroke-width="2" rx="8" opacity="0.95"/>
    <text x="115" y="28" text-anchor="middle" class="label" fill="{colors['text']}" font-size="16" font-weight="600">FLOOR AREA</text>
    <line x1="15" y1="40" x2="215" y2="40" stroke="{colors['grid']}" stroke-width="1"/>
    <text x="25" y="65" class="label" fill="{colors['dimension']}" font-size="13">Width: {width:.1f}{unit}</text>
    <text x="25" y="85" class="label" fill="{colors['dimension']}" font-size="13">Depth: {depth:.1f}{unit}</text>
    <text x="25" y="108" class="label" fill="{colors['accent']}" font-size="14" font-weight="600">Area: {width * depth:.1f}{unit}²</text>
  </g>
</svg>'''
        
        return svg
    
    def _generate_column_grid(self, x, y, w, h, colors):
        """Generate structural column grid markers."""
        cols = 5
        rows = 4
        col_spacing = w / cols
        row_spacing = h / rows
        
        markers = []
        for i in range(cols + 1):
            for j in range(rows + 1):
                cx = x + i * col_spacing
                cy = y + j * row_spacing
                markers.append(f'<circle cx="{cx}" cy="{cy}" r="6" fill="{colors["accent"]}" opacity="0.6"/>')
        
        return '\n  '.join(markers)


class TextToSpeechTool(Tool):
    """Tool for generating speech audio from text using AWS Polly."""

    def __init__(self):
        super().__init__(
            name="text_to_speech",
            description="Generates speech audio from text using AWS Polly. Use for narration in demo videos with sound. When creating a demo video with sound: use url_resolver only if the user provided a URL; then call this tool with the script/summary text, then combine_video_audio to mux the audio with the generated video.",
            instructions="Use when the user wants a video with narration or sound. Call with the narration script (e.g. from summarization of website content). Keep text under ~3000 characters for a single clip; for longer text, use a shortened summary. Returns an audio file (.mp3) to use with combine_video_audio."
        )

    def get_parameters_schema(self) -> List[Dict[str, Any]]:
        return [
            {
                "name": "text",
                "description": "The text to convert to speech (narration script). Use concise, natural phrasing. Max ~3000 characters per call.",
                "type": "string",
                "required": True
            },
            {
                "name": "voice_id",
                "description": "Polly voice. Examples: 'Joanna', 'Matthew' (US English), 'Amy', 'Brian' (UK). Default: 'Joanna'.",
                "type": "string",
                "required": False
            },
            {
                "name": "engine",
                "description": "Polly engine: 'standard' or 'neural' (better quality). Default: 'neural' if available.",
                "type": "string",
                "required": False
            }
        ]

    def execute(self, parameters: Dict[str, Any]) -> Dict[str, Any]:
        text = (parameters.get('text') or '').strip()
        if not text:
            return {"error": "text parameter is required"}

        voice_id = (parameters.get('voice_id') or 'Joanna').strip()
        engine = (parameters.get('engine') or 'neural').strip().lower()
        if engine not in ('standard', 'neural'):
            engine = 'neural'

        try:
            from .aws_utils import create_boto3_client, get_aws_region
            import uuid

            region = get_aws_region()
            polly = create_boto3_client('polly', region_name=region)

            # Polly has a 3000 character limit per request. Use synchronous synthesize_speech
            # only (no start_speech_synthesis_task) so the tool returns only after audio is ready.
            text = text[:3000]

            try:
                response = polly.synthesize_speech(
                    Text=text,
                    OutputFormat='mp3',
                    VoiceId=voice_id,
                    Engine='neural' if engine == 'neural' else 'standard'
                )
            except Exception as e:
                if 'neural' in str(e).lower() or 'Engine' in str(e):
                    try:
                        response = polly.synthesize_speech(
                            Text=text,
                            OutputFormat='mp3',
                            VoiceId=voice_id,
                            Engine='standard'
                        )
                    except Exception as e2:
                        logger.warning(f"Polly synthesize_speech failed: {e2}")
                        return {"error": f"Text-to-speech failed: {str(e2)}"}
                else:
                    logger.warning(f"Polly synthesize_speech failed: {e}")
                    return {"error": f"Text-to-speech failed: {str(e)}"}

            audio_stream = response.get('AudioStream')
            if not audio_stream:
                return {"error": "No audio stream returned from Polly"}

            # Read entire stream synchronously (block until all bytes received)
            audio_data = audio_stream.read()
            if not audio_data or len(audio_data) < 100:
                return {"error": "No audio data received from Polly"}

            file_name = f"generated_audio_{uuid.uuid4().hex[:12]}.mp3"
            saved_path = default_storage.save(
                f"generated_audio/{file_name}",
                ContentFile(audio_data)
            )
            # Ensure file is committed and readable before returning (sync behavior for combine_video_audio)
            if not default_storage.exists(saved_path):
                return {"error": "Failed to save audio file to storage"}
            try:
                size = default_storage.size(saved_path)
                if size is None or size < 100:
                    return {"error": "Saved audio file is empty or too small"}
            except Exception:
                pass  # size() not supported by all backends
            audio_url = default_storage.url(saved_path)

            return {
                "audio_url": audio_url,
                "file_name": file_name,
                "saved_path": saved_path,
                "success": True,
                "voice_id": voice_id,
                "text_length": len(text)
            }
        except Exception as e:
            logger.error(f"Error in text_to_speech: {str(e)}", exc_info=True)
            return {"error": f"Text-to-speech failed: {str(e)}"}


class CombineVideoAudioTool(Tool):
    """Tool for combining a silent video with an audio file into a single video with sound."""

    def __init__(self):
        super().__init__(
            name="combine_video_audio",
            description="Combines a silent video (e.g. from text_to_video or image_to_video) with an audio file (e.g. from text_to_speech) into one video with sound. Use after generating the video and the narration audio for demo videos with sound.",
            instructions="Use when the user wants a video with sound. Call with the video file name or path from text_to_video/image_to_video and the audio file name or path from text_to_speech. Returns the final video URL and file name."
        )

    def get_parameters_schema(self) -> List[Dict[str, Any]]:
        return [
            {
                "name": "video_path",
                "description": "Video file name or path (e.g. generated_video_xxx.mp4 from text_to_video, or path like generated_videos/generated_video_xxx.mp4).",
                "type": "string",
                "required": True
            },
            {
                "name": "audio_path",
                "description": "Audio file name or path (e.g. generated_audio_xxx.mp3 from text_to_speech, or path like generated_audio/generated_audio_xxx.mp3).",
                "type": "string",
                "required": True
            }
        ]

    def _resolve_media_path(self, file_path: str, kind: str) -> Optional[str]:
        """Resolve to a storage key (path) for video or audio. Returns None if not found."""
        if not file_path or not file_path.strip():
            return None
        file_path = file_path.strip()

        uuid_pattern = re.compile(r'^[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}$', re.I)
        if uuid_pattern.match(file_path):
            try:
                from .models import ConversationFile
                conv_file = ConversationFile.objects.get(file_id=file_path)
                if conv_file.file_path:
                    return conv_file.file_path.name if hasattr(conv_file.file_path, 'name') else str(conv_file.file_path)
            except Exception:
                pass

        if file_path.startswith('/media/') or file_path.startswith('media/'):
            file_path = file_path.lstrip('/').replace('media/', '', 1) if file_path.startswith('/') else file_path.replace('media/', '', 1)
        elif file_path.startswith('/'):
            file_path = file_path.lstrip('/')

        if '/' not in file_path:
            for prefix in ('generated_videos/', 'generated_audio/'):
                candidate = f"{prefix}{file_path}"
                if default_storage.exists(candidate):
                    return candidate
            try:
                from .models import ConversationFile
                conv_file = ConversationFile.objects.filter(file_name=file_path).first()
                if conv_file and conv_file.file_path:
                    return conv_file.file_path.name if hasattr(conv_file.file_path, 'name') else str(conv_file.file_path)
            except Exception:
                pass

        return file_path if default_storage.exists(file_path) else None

    def execute(self, parameters: Dict[str, Any]) -> Dict[str, Any]:
        # Accept canonical or common alternate keys from LLM output
        video_path = (
            parameters.get('video_path')
            or parameters.get('video_file')
            or parameters.get('video')
            or ''
        )
        audio_path = (
            parameters.get('audio_path')
            or parameters.get('audio_file')
            or parameters.get('audio')
            or ''
        )
        if isinstance(video_path, str):
            video_path = video_path.strip()
        else:
            video_path = str(video_path).strip() if video_path else ''
        if isinstance(audio_path, str):
            audio_path = audio_path.strip()
        else:
            audio_path = str(audio_path).strip() if audio_path else ''
        if not video_path or not audio_path:
            return {"error": "video_path and audio_path are required (video: video_path/video_file/video; audio: audio_path/audio_file/audio)"}

        video_key = self._resolve_media_path(video_path, 'video')
        audio_key = self._resolve_media_path(audio_path, 'audio')
        if not video_key:
            return {"error": f"Video file not found: {video_path}"}
        if not audio_key:
            return {"error": f"Audio file not found: {audio_path}"}

        try:
            import uuid
            # Read from storage to temp files (works with S3 or local)
            with default_storage.open(video_key, 'rb') as vf:
                video_data = vf.read()
            with default_storage.open(audio_key, 'rb') as af:
                audio_data = af.read()

            if not audio_data or len(audio_data) < 100:
                return {"error": "Audio file is empty or too small; ensure text_to_speech ran successfully and returned audio."}
            if not video_data or len(video_data) < 100:
                return {"error": "Video file is empty or too small; ensure text_to_video or image_to_video completed successfully."}

            with tempfile.NamedTemporaryFile(suffix='.mp4', delete=False) as vtmp:
                vtmp.write(video_data)
                vtmp.flush()
                video_temp = vtmp.name
            with tempfile.NamedTemporaryFile(suffix='.mp3', delete=False) as atmp:
                atmp.write(audio_data)
                atmp.flush()
                audio_temp = atmp.name
            out_temp = tempfile.NamedTemporaryFile(suffix='.mp4', delete=False)
            out_temp.close()
            out_path = out_temp.name

            try:
                # Resolve ffmpeg: setting FFMPEG_PATH, then PATH, then /usr/bin/ffmpeg
                ffmpeg_bin = getattr(settings, 'FFMPEG_PATH', None) or ''
                if not ffmpeg_bin or not os.path.isfile(ffmpeg_bin):
                    ffmpeg_bin = shutil.which('ffmpeg') or (('/usr/bin/ffmpeg') if os.path.isfile('/usr/bin/ffmpeg') else None) or 'ffmpeg'
                # ffmpeg: video + audio -> one mp4. Map video from input 0, audio from input 1 so we
                # always use the narration track (video from Bedrock may have a silent audio stream).
                # Encode audio as AAC 44.1kHz for broad compatibility; -shortest trims to shortest stream.
                cmd = [
                    ffmpeg_bin, '-y',
                    '-i', video_temp,
                    '-i', audio_temp,
                    '-map', '0:v:0',
                    '-map', '1:a:0',
                    '-c:v', 'copy',
                    '-c:a', 'aac',
                    '-b:a', '128k',
                    '-ar', '44100',
                    '-shortest',
                    out_path
                ]
                proc = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
                if proc.returncode != 0:
                    logger.warning(f"ffmpeg stderr: {proc.stderr}")
                    return {
                        "error": "Failed to combine video and audio",
                        "note": proc.stderr or str(proc.returncode)
                    }

                # Verify output has an audio stream (avoid delivering silent video)
                probe_bin = (ffmpeg_bin.replace('ffmpeg', 'ffprobe') if 'ffmpeg' in ffmpeg_bin else 'ffprobe')
                if not os.path.isfile(probe_bin):
                    probe_bin = shutil.which('ffprobe') or '/usr/bin/ffprobe'
                if os.path.isfile(probe_bin):
                    probe_proc = subprocess.run(
                        [probe_bin, '-v', 'error', '-select_streams', 'a', '-show_entries', 'stream=codec_type', '-of', 'csv=p=0', out_path],
                        capture_output=True, text=True, timeout=10
                    )
                    if probe_proc.returncode != 0 or 'audio' not in (probe_proc.stdout or '').lower():
                        logger.warning(f"ffprobe did not find audio in output: {probe_proc.stdout!r} {probe_proc.stderr!r}")
                        return {"error": "Combined file has no audio track; please retry or check that text_to_speech produced valid audio."}

                with open(out_path, 'rb') as f:
                    out_data = f.read()
            finally:
                for p in (video_temp, audio_temp, out_path):
                    try:
                        if os.path.exists(p):
                            os.unlink(p)
                    except Exception:
                        pass

            file_name = f"generated_video_{uuid.uuid4()}.mp4"
            saved_path = default_storage.save(
                f"generated_videos/{file_name}",
                ContentFile(out_data)
            )
            video_url = default_storage.url(saved_path)

            return {
                "video_url": video_url,
                "file_name": file_name,
                "success": True,
                "note": "Video with sound ready for download."
            }
        except subprocess.TimeoutExpired:
            return {"error": "Combine video/audio timed out (120s)"}
        except FileNotFoundError:
            return {
                "error": "ffmpeg not found",
                "suggestion": (
                    "Install ffmpeg on the server. "
                    "Amazon Linux 2023: ffmpeg is not in dnf repos; use a static build: "
                    "curl -sL https://johnvansickle.com/ffmpeg/releases/ffmpeg-release-arm64-static.tar.xz | tar -xJ -C /tmp && sudo mv /tmp/ffmpeg-*-arm64-static/ffmpeg /usr/local/bin/ffmpeg. "
                    "Ubuntu/Debian: sudo apt-get install -y ffmpeg. "
                    "Or set FFMPEG_PATH in .env to the full path of the ffmpeg binary."
                )
            }
        except Exception as e:
            logger.error(f"Error in combine_video_audio: {str(e)}", exc_info=True)
            return {"error": f"Combine video/audio failed: {str(e)}"}


class ToolManager:
    """Manages all available tools for agents."""
    
    def __init__(self):
        self.tools = {}
        self._register_default_tools()
    
    def _register_default_tools(self):
        """Register default tools."""
        default_tools = [
            ReadFileTool(),
            WriteFileTool(),
            WebSearchTool(),
            URLResolverTool(),
            CodeExecutorTool(),
            TranscriptionTool(),
            TextToImageTool(),
            TextToVideoTool(),
            ImageToVideoTool(),
            TextToSpeechTool(),
            CombineVideoAudioTool(),
            SVGDiagramTool(),
            OCRTool(),
            SummarizationTool(),
            QuestionAnsweringTool(),
            TranslationTool(),
        ]
        
        for tool in default_tools:
            self.tools[tool.name] = tool
    
    def register_custom_tool(self, tool: CustomTool):
        """Register a custom tool."""
        if len([t for t in self.tools.values() if isinstance(t, CustomTool)]) >= 10:
            raise ValueError("Maximum 10 custom tools allowed")
        self.tools[tool.name] = tool
    
    def get_tool(self, name: str) -> Optional[Tool]:
        """Get a tool by name."""
        return self.tools.get(name)
    
    def get_all_tools(self) -> Dict[str, Tool]:
        """Get all registered tools."""
        return self.tools
    
    def get_tools_schema(self) -> List[Dict[str, Any]]:
        """Get schema for all tools (for LLM)."""
        return [tool.get_schema() for tool in self.tools.values()]



